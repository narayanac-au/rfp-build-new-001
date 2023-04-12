import openai as ap
from dataclasses import dataclass
from email.charset import BASE64
from urllib import response
from django.shortcuts import render, redirect
from docx import Document
from RFP.models import *
from RFP.models import Question as QP
from RFP.models import UserQuery as UQ
from PyPDF2 import PdfFileReader, PdfFileWriter, PdfFileMerger
from docx2pdf import convert
from docx import Document as CX
from django_pandas.io import read_frame
from RFP.Question_Similarity import Model_Buiding_approach
from RFP.Question_Extraction import Extraction
import pandas as pd
from .forms import UserQueryForm, ImageForm, SelectDropQueryForm
from .models import Image
from django.contrib.auth.models import User, auth
from django.contrib import messages
from sentence_transformers import SentenceTransformer, util
import os
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse
import json
from .models import KPMGgeo as kg
from .models import First_Page_upload as FPU
from .models import Product
from .forms import ProjectForm

from django.db.models import Q
# Loading the Embedding Pretrained Model


def load_model():
    # Initializing the embedding model
    embedder = SentenceTransformer("./RFP/bert-base-uncased")
    return embedder


model_name = "bert-base-uncased"
embedder = load_model()


# first page
def index_view(request):
    # showname=request.session['showname']

    # docobj1 = Document.objects.filter(
    # id__in=[1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 18, 19, 20, 21, 22, 23]).update(selected=" ")
    #     UQL=UQ.objects.all()
    #     UQL.delete()
    #     ud=SelectDropQuery.objects.all()
    #     ud.delete()
    # Dp=DropQuery.objects.filter(user=showname)
    # Dp.delete()
    return render(request, 'index.html')

# second page details


def infodetails_view(request):
    if request.method == "GET":
        geo = KPMGgeo.objects.all()

        return render(request, 'info.html', {"geo": geo})


@csrf_exempt
def geoadd_view(request):
    if request.method == "POST":
        id = request.POST['geo']
        add = list(KPMGadd.objects.filter(KPMGgeo=id).values())
        print("ADD FROM GEOADD_VIEW FUNCTION IS ", add)
        return JsonResponse({"add": add}, safe=False)
# third page select standard RFP


def doc_content_view(request):

    if request.method == "POST":
        KPMGaddress1 = request.POST.get("address1")
        clientaddress_line1 = request.POST.get("address_line1")
        clientaddress_line2 = request.POST.get("address_line2")
        clientPostal_Code = request.POST.get("Postal_Code")
        client_name = request.POST.get("clientname")
        showname = request.POST.get("showname")
        industry = request.POST.get("industry")
        country = request.POST.get("countries")
        address = request.POST.get("address")
        KPMGgeol = request.POST.get("geo")
        KPMGgeoo = kg.objects.filter(id=KPMGgeol)
        KPMGgeo = str(KPMGgeoo[0])

        request.session['client_name'] = client_name
        request.session['industry'] = industry
        request.session['country'] = country
        request.session['showname'] = showname
        request.session['address'] = address
        #formm = ImageForm(data=request.POST, files=request.FILES)

        # formm.save()

        client_name = request.session['client_name']
        industry_data = request.session['industry']
        country_data = request.session['country']
        infos = info.objects.update_or_create(clientfullname=client_name, clientshortname=showname,
                                    clientindustry=industry, clientgeo=country, clientaddress_line1=clientaddress_line1, clientaddress_line2=clientaddress_line2, clientPostal_Code=clientPostal_Code, KPMGaddress1=KPMGaddress1, KPMGgeo=KPMGgeo)
        # infos.save()
        import os
        path = 'static/media/'+client_name+'/'+country+'/'+industry
        try:
            os.makedirs(path)

        except:
            pass
        """
        p = Users.objects.filter(user=client_name)
        # Doc = Document.objects.filter(
        #     country=country, industry=industry, user__in=p)
        Doc = Document.objects.filter(
            country=country, industry=industry)[:3]
        # Doc = list(Doc)
        print(Doc, 'top3_doc')
        UserDoc = Document.objects.filter(
            country=country, industry=industry, user__in=p).exclude(id__in=Doc)
        print(UserDoc, 'user_doc')
        default_user_doc = list(set(Doc) | set(UserDoc))

        print(default_user_doc, 'default docs')

        Poc = Document.objects.filter(
            country=country, industry=industry).exclude(id__in=Doc).exclude(id__in=UserDoc)

        """
        user = Users.objects.filter(user=client_name)

        Doc = RfpSection.objects.filter(
            Q(industry=industry_data) | Q(country=country_data)
        )[:3]
        
        print(Doc, country, industry, 'doccc')

        UserDoc = RfpSection.objects.filter(
            Q(industry=industry_data) | Q(country=country_data), Q(user__in=user)).exclude(id__in=Doc)
        
        print(UserDoc, 'user_doc')
        default_user_doc = list(set(Doc) | set(UserDoc))

        print(default_user_doc, 'default docs')

        Poc = RfpSection.objects.filter(
            Q(industry=industry_data) | Q(country=country_data)).exclude(id__in=Doc).exclude(id__in=UserDoc)
        
        print(default_user_doc, 'default ,  user doc from rfpsection ')

        
        que = askques.objects.filter(user=client_name, selected="on")

        if que:
            que
        else:
            que = None
        from .models import Image
        industry = request.POST.get("industry")
        img = Image.objects.filter(caption=industry)
        print(Image, 'imggg')

        print(Doc, 'Doc')
        print()
        print()
        print('-----')
        print(Poc, 'poc')
        return render(request, 'doc_content.html', {'UserDoc': default_user_doc, 'Poc': Poc, "showname": showname, "country": country, "industry": industry,  "que": que, 'Image': img})
    if request.method == "GET":

        client_name = request.session['client_name']
        industry_data = request.session['industry']
        country_data = request.session['country']
        showname = request.session['showname']

        user = Users.objects.filter(user=client_name)

        Doc = RfpSection.objects.filter(
            Q(industry=industry_data) | Q(country=country_data)
        )[:3]
        
        # print(Doc, country, industry, 'doccc')

        UserDoc = RfpSection.objects.filter(
            Q(industry=industry_data) | Q(country=country_data), Q(user__in=user)).exclude(id__in=Doc)
        
        print(UserDoc, 'user_doc')
        default_user_doc = list(set(Doc) | set(UserDoc))

        print(default_user_doc, 'default docs')

        Poc = RfpSection.objects.filter(
            Q(industry=industry_data) | Q(country=country_data)).exclude(id__in=Doc).exclude(id__in=UserDoc)
        
        print(default_user_doc, 'default ,  user doc from rfpsection ')
        
        que = askques.objects.filter(user=client_name, selected="on")
        if que:
            que
        else:
            que = None
        from .models import Image

        #Image = Image.objects.filter(caption=industry)
        try:
            SelectedImage = Image.objects.filter(
                caption=industry_data, user__in=p)
        except:
            SelectedImage = None
        if SelectedImage:
            img = Image.objects.filter(caption=industry_data).exclude(user__in=user)

        else:
            img = Image.objects.filter(caption=industry_data)

        print(Doc, 'Doc')
        print()
        print()
        print('-----')
        print(Poc, 'poc')
        return render(request, 'doc_content.html', {'UserDoc': default_user_doc, 'Poc': Poc, "showname": showname, "country": country_data, "industry": industry_data, "que": que, 'Image': img, 'SelectedImage': SelectedImage})

# fourth page type question RFP


def firstpage_view(request):

    if request.method == "POST":
        client_name = request.session['client_name']
        industry = request.session['industry']
        country = request.session['country']
        radio = request.POST.get("flexRadioDefault")
        Tick1 = request.POST.get("1")
        Tick2 = request.POST.get("2")
        Tick3 = request.POST.get("3")
        Tick4 = request.POST.get("4")
        Tick5 = request.POST.get("5")
        Tick6 = request.POST.get("6")
        Tick7 = request.POST.get("7")
        Tick8 = request.POST.get("8")
        Tick9 = request.POST.get("9")
        Tick10 = request.POST.get("10")
        Tick11 = request.POST.get("11")
        Tick12 = request.POST.get("12")
        Tick13 = request.POST.get("13")
        Tick14 = request.POST.get("14")
        Tick15 = request.POST.get("15")
        Tick16 = request.POST.get("16")
        Tick17 = request.POST.get("17")
        Tick18 = request.POST.get("18")
        Tick19 = request.POST.get("19")
        Tick20 = request.POST.get("20")
        Tick21 = request.POST.get("21")
        Tick22 = request.POST.get("22")
        Tick23 = request.POST.get("23")
        Quest = request.POST.get("Quest")

        # from .models import Image
        # p = Users.objects.filter(user=client_name)
        # print("NAMEEEEEEEEEEEEEEEEEEEEEEE")
        # print(p)
        # print("NAMEEEEEEEEEEEEEEEEEEEEEEE")
        # print("clear")
        # p.Image_set.all()
        # print("clear")
        # print("IMAAAAAAAAAAAAAGGEEE")
        # print(p.Image_set.all())
        # print("IMAAAAAAAAAAAAAGGEEE")

        check_user = Users.objects.filter(user=client_name)

        if check_user:
            d = check_user[0]

        else:
            d = Users(user=client_name)
            d.save()

        if radio:

            #u = Users.objects.get(user=client_name)

            Imagee = Image.objects.get(
                id=radio)

            l = Image.objects.filter(id=radio)
            t = Image.objects.filter(id=radio).update(selected="on")

            if client_name in l:
                pass
            else:
                try:
                    c = Imagee.user.add(d)
                except:
                    pass

                l = Imagee.user.all()

        if Quest:
            ck = askques.objects.filter(user=client_name)
            if not ck:
                askquestion = askques(user=client_name, selected="on")
                askquestion.save()
            if ck:
                askquestion = askques.objects.filter(
                    user=client_name).update(selected="on")
        else:
            check = askques.objects.filter(user=client_name)
            if check:
                askque = askques.objects.filter(
                    user=client_name).update(selected=" ")
            else:
                askque = askques.objects.create(user=client_name, selected=" ")

        if Tick1:
            print('inside tick 111')

            docobj1 = Document.objects.filter(id=1).update(selected="on")
            docobj1 = Document.objects.get(id=1).user.add(d)

        else:
            print('inside tick 111 elseeee')
            try:
                docobj1 = Document.objects.get(id=1, user=d)
                c = docobj1.user.remove(d)
            except:
                pass

        if Tick2:

            docobj2 = Document.objects.filter(id=2).update(selected="on")
            docobj2 = Document.objects.get(id=2).user.add(d)

        else:
            try:
                docobj2 = Document.objects.get(id=2, user=d)
                c = docobj2.user.remove(d)
            except:
                pass

        if Tick3:

            docobj3 = Document.objects.filter(id=3).update(selected="on")
            docobj3 = Document.objects.get(id=3).user.add(d)

        else:
            try:
                docobj3 = Document.objects.get(id=3, user=d)
                c = docobj3.user.remove(d)
            except:
                pass

        if Tick4:

            docobj4 = Document.objects.filter(id=4).update(selected="on")
            docobj4 = Document.objects.get(id=4).user.add(d)

        else:
            try:
                docobj4 = Document.objects.get(id=4, user=d)
                c = docobj4.user.remove(d)
            except:
                pass

        if Tick5:

            docobj5 = Document.objects.filter(id=5).update(selected="on")
            docobj5 = Document.objects.get(id=5).user.add(d)

        else:
            try:
                docobj5 = Document.objects.get(id=5, user=d)
                c = docobj5.user.remove(d)
            except:
                pass

        if Tick6:

            docobj6 = Document.objects.filter(id=6).update(selected="on")
            docobj6 = Document.objects.get(id=6).user.add(d)

        else:
            try:
                docobj6 = Document.objects.get(id=6, user=d)
                c = docobj6.user.remove(d)
            except:
                pass

        if Tick7:

            docobj7 = Document.objects.filter(id=7).update(selected="on")
            docobj7 = Document.objects.get(id=7).user.add(d)

        else:
            try:
                docobj7 = Document.objects.get(id=7, user=d)
                c = docobj7.user.remove(d)
            except:
                pass

        if Tick8:

            docobj8 = Document.objects.filter(id=8).update(selected="on")
            docobj8 = Document.objects.get(id=8).user.add(d)

        else:
            try:
                docobj8 = Document.objects.get(id=8, user=d)
                c = docobj8.user.remove(d)
            except:
                pass

        if Tick9:

            docobj9 = Document.objects.filter(id=9).update(selected="on")
            docobj9 = Document.objects.get(id=9).user.add(d)

        else:
            try:
                docobj9 = Document.objects.get(id=9, user=d)
                c = docobj9.user.remove(d)
            except:
                pass

        if Tick10:

            docobj10 = Document.objects.filter(id=10).update(selected="on")
            docobj10 = Document.objects.get(id=10).user.add(d)

        else:
            try:
                docobj10 = Document.objects.get(id=10, user=d)
                c = docobj10.user.remove(d)
            except:
                pass

        if Tick11:

            docobj11 = Document.objects.filter(id=11).update(selected="on")
            docobj11 = Document.objects.get(id=11).user.add(d)

        else:
            try:
                docobj11 = Document.objects.get(id=11, user=d)
                c = docobj11.user.remove(d)
            except:
                pass

        if Tick12:

            docobj12 = Document.objects.filter(id=12).update(selected="on")
            docobj12 = Document.objects.get(id=12).user.add(d)

        else:
            try:
                docobj12 = Document.objects.get(id=12, user=d)
                c = docobj12.user.remove(d)
            except:
                pass

        if Tick13:

            docobj13 = Document.objects.filter(id=13).update(selected="on")
            docobj13 = Document.objects.get(id=13).user.add(d)

        else:
            try:
                docobj13 = Document.objects.get(id=13, user=d)
                c = docobj13.user.remove(d)
            except:
                pass

        if Tick14:

            docobj14 = Document.objects.filter(id=14).update(selected="on")
            docobj14 = Document.objects.get(id=14).user.add(d)

        else:
            try:
                docobj14 = Document.objects.get(id=14, user=d)
                c = docobj14.user.remove(d)
            except:
                pass

        if Tick15:

            docobj15 = Document.objects.filter(id=15).update(selected="on")
            docobj15 = Document.objects.get(id=15).user.add(d)

        else:
            try:
                docobj15 = Document.objects.get(id=15, user=d)
                c = docobj15.user.remove(d)
            except:
                pass

        if Tick16:

            docobj16 = Document.objects.filter(id=16).update(selected="on")
            docobj16 = Document.objects.get(id=16).user.add(d)

        else:
            try:
                docobj16 = Document.objects.get(id=16, user=d)
                c = docobj16.user.remove(d)
            except:
                pass

        if Tick17:

            docobj17 = Document.objects.filter(id=17).update(selected="on")
            docobj17 = Document.objects.get(id=17).user.add(d)

        else:
            try:
                docobj17 = Document.objects.get(id=17, user=d)
                c = docobj17.user.remove(d)
            except:
                pass

        if Tick18:

            docobj18 = Document.objects.filter(id=18).update(selected="on")
            docobj18 = Document.objects.get(id=18).user.add(d)

        else:
            try:
                docobj18 = Document.objects.get(id=18, user=d)
                c = docobj18.user.remove(d)
            except:
                pass

        if Tick19:

            docobj19 = Document.objects.filter(id=19).update(selected="on")
            docobj19 = Document.objects.get(id=19).user.add(d)

        else:
            try:
                docobj19 = Document.objects.get(id=19, user=d)
                c = docobj19.user.remove(d)
            except:
                pass

        if Tick20:

            docobj20 = Document.objects.filter(id=20).update(selected="on")
            docobj20 = Document.objects.get(id=20).user.add(d)

        else:
            try:
                docobj20 = Document.objects.get(id=20, user=d)
                c = docobj20.user.remove(d)
            except:
                pass

        if Tick21:

            docobj21 = Document.objects.filter(id=21).update(selected="on")
            docobj21 = Document.objects.get(id=21).user.add(d)

        else:
            try:
                docobj21 = Document.objects.get(id=21, user=d)
                c = docobj21.user.remove(d)
            except:
                pass

        if Tick22:

            docobj22 = Document.objects.filter(id=22).update(selected="on")
            docobj22 = Document.objects.get(id=22).user.add(d)

        else:
            try:
                docobj22 = Document.objects.get(id=22, user=d)
                c = docobj22.user.remove(d)
            except:
                pass

        if Tick23:

            docobj23 = Document.objects.filter(id=23).update(selected="on")
            docobj23 = Document.objects.get(id=23).user.add(d)

        else:
            try:
                docobj23 = Document.objects.get(id=23, user=d)
                c = docobj23.user.remove(d)
            except:
                pass

        client_name = request.session['client_name']
        industry = request.session['industry']
        country = request.session['country']
        showname = request.session['showname']
        data = Question.objects.filter(country=country, industry=industry)
        # data=Question.objects.all()
        # print("!!!!!!!!!!!!!!!!!!!!!!!!!")
        # print(Quest)
        # print("!!!!!!!!!!!!!!!!!!!!!!!!!")

        # print("final output",index_list,"_",cos_sim_list)

        # display seleted index
        p = Users.objects.filter(user=client_name)
        Doc1 = Document.objects.filter(selected="on", user__in=p)
        Userq = UQ.objects.filter(user=client_name)
        c = Userq.exists()

        User = askques.objects.filter(user=client_name, selected="on")
        Quest = User.exists()

        # display seleted index
        return render(request, 'firstpage.html', {'Data': data, "showname": showname, "country": country, "industry": industry, "Doc": Doc1, "Quest": Quest, "c": c})
    if request.method == "GET":
        client_name = request.session['client_name']
        industry = request.session['industry']
        country = request.session['country']
        showname = request.session['showname']
        data = Question.objects.filter(country=country, industry=industry)
        p = Users.objects.filter(user=client_name)
        Doc1 = RfpSection.objects.filter(user__in=p)
        Userq = UQ.objects.filter(user=client_name)
        c = Userq.exists()

        User = askques.objects.filter(user=client_name, selected="on")
        Quest = User.exists()

        # display seleted index
        return render(request, 'firstpage.html', {'Data': data, "showname": showname, "country": country, "industry": industry, "Doc": Doc1, "Quest": Quest, "c": c})
 # fifth page mcq options


def secondpage_view(request):

    if request.method == "GET":
        showname = request.session['showname']
        industry = request.session['industry']
        country = request.session['country']
        Query = request.GET.get("Query")
        data = Question.objects.filter(country=country, industry=industry)
        # df = read_frame(data)
        # df.to_excel("Data.xlsx", index=False)
        # df = pd.read_excel("Question_Pairing_Final.xlsx",
        #                    sheet_name="Tablib Dataset")
        # Time taking for passing multiple queries

        #---------------------------------------------------Search Engine Functionality(Start)-------------------------------------
        industry= ["Agnostic"]
        country=  ["AU"]
        # section=  ["Relevant Experience","Project Team Structure"]
        section = []

        industry_text = "_".join(industry)
        print("Industry text:-",industry_text)
        country_text = "_".join(country)
        print("Country text:-",country_text)
        df = pd.read_excel("database.xlsx")

        # Creating an model
        model_build = Model_Buiding_approach(industry_text, country_text, model_name=model_name)

        print("Before Filtered DataFrame:- ", df.shape)
        # Filter the data based on Country and Industry
        df = model_build.dataframe_filter_return(df, industry, country, section)



        print("Filtered DataFrame:- ", df.shape)
        # Training Dataset Vector
        if os.path.isfile("RFP/Embedding Models/corpus_embedding"+'-'+model_name+'-'+country_text+"-"+industry_text+".pt"):
            print("Need not to Run Training File as Model is Present")
        else:
            print("File is not present. Initializing the training command.")
            execution_time = model_build.training_dataset_vector(df, embedder = embedder)
            print("Execution of Training Time:-", execution_time)

        # Input the query
        query = "Has your organisation been involved in any business-related litigation that may affect the performance of these Services in the past five years?"
        # query=Query

        index_list, cos_sim_list = model_build.query_resolved(query=query, df=df, embedder=embedder)

        print("Index List:-", index_list)
        print("Cosine Similarity List", cos_sim_list)
        #---------------------------------------------------Search Engine Functionality(End)-------------------------------------
        # Input the query
        # query = "Has your organisation been involved in any business-related litigation that may affect the performance of these Services in the past five years?"
        # query = Query

        # request.session['query'] = query
        # non = ""
        # index_list, cos_sim_list = model_build.query_resolved(
        #     query=query, df=df, embedder=embedder)

        data = Question.objects.filter(id__in=index_list)
        # data =df.iloc[index_list]
        data0 = ""
        data1 = ""
        data2 = ""
        id0 = ""
        id1 = ""
        id2 = ""
        if len(index_list) == 1:
            data0 = data[0].Model_Image_Link_BLOB
            id0 = data[0].id
        elif len(index_list) == 2:
            data0 = data[0].Model_Image_Link_BLOB
            id0 = data[0].id
            data1 = data[1].Model_Image_Link_BLOB
            id1 = data[1].id
        elif len(index_list) == 3:
            data0 = data[0].Model_Image_Link_BLOB
            id0 = data[0].id
            data1 = data[1].Model_Image_Link_BLOB
            id1 = data[1].id
            data2 = data[2].Model_Image_Link_BLOB
            id2 = data[2].id
        else:
            non = "No matching question found"
            # data0=data[0].Questions
            # id0=data[0].id
            # data1=data[1].Questions
            # id1=data[1].id
            # data2=data[2].Questions
           # id2=data[2].id
        l = Question.objects.get(id=170)

        f = l.File

        return render(request, 'mcq.html', {'non': non, 'data': data, 'index_list': index_list, "data0": data0, "data1": data1, "data2": data2, "id0": id0, "id1": id1, "id2": id2, "Query": Query, "showname": showname, "country": country, "industry": industry, "f": f})

 # sixth page preview


def pre_view(request):
    client_name = request.session['client_name']
    if request.method == "GET":
        showname = request.session['showname']
        industry = request.session['industry']
        country = request.session['country']
        UQL = UQ.objects.exclude(sentapproval="on").filter(user=client_name)
        return render(request, 'preview.html', {'Question': Question, "showname": showname, "country": country, "industry": industry, "UQL": UQL})
    if request.method == "POST":

        industry = request.session['industry']
        country = request.session['country']
        showname = request.session['showname']
        id0 = request.POST.get("K")
        id1 = request.POST.get("M")
        id2 = request.POST.get("L")
        p = [id0, id1, id2]
        answer = Question.objects.filter(id__in=p)
        final_list = filter(None, p)
        print_pdf = list(final_list)

        # request.session['print_pdf']=print_pdf

        answer = Question.objects.filter(id__in=print_pdf)

        query = request.session['query']

        print("%%%%%%%%%%%%%%%%%%")
        print("%%%%%%%%%%%%%%%%%%")
        print(query)
        print("##################")
        print("##################")
        UQ.answer1 = " "
        UQ.answer2 = " "
        UQL = UQ.objects.exclude(sentapproval="on").filter(user=client_name)

        Questionobj = Question.objects.filter(id__in=p).update(Tick1="on")
        return render(request, 'preview.html', {"answer": answer, "showname": showname, "country": country, "industry": industry, "UQL": UQL, "query": query})
 # preview page print standard pdf


def printpdf(request):
    if request.method == "GET":
        client_name = request.session['client_name']
        industry = request.session['industry']
        country = request.session['country']
        # Once index selected  starts---------------------------------------------------------------------
        p = Users.objects.filter(user=client_name)

        files = Document.objects.filter(user__in=p)
        print('$$$$$$$$$$$$$$$$$$')
        print(files)
        print('$$$$$$$$$$$$$$$$$$')
        for file in files:
            file = files[0].File

        from docxcompose.composer import Composer
        from docx import Document as Document_compose
        path = 'static/media/'+client_name+'/'+country+'/'+industry

        def combine_all_docx(filename_master, files_list):
            number_of_sections = len(files_list)
            master = Document_compose(filename_master)
            composer = Composer(master)
            for i in range(0, number_of_sections):
                doc_temp = Document_compose(files_list[i])
                composer.append(doc_temp)
            composer.save("./"+path+"/combined_file.docx")
        # For Example
        from docx import Document as DF
        docu = DF()

        docu.save("./"+path+"/C.docx")
        filename_master = "./"+path+"./C.docx"
        files_list = []
        for file in files:
            files_list.append(file.File)
        print('$$$$$$$$$$$$$$$$$$')
        print(files_list)
        print('$$$$$$$$$$$$$$$$$$')

        # Calling the function
        combine_all_docx(filename_master, files_list)
        # print standard starts----------------------------start-----------------------------------------

        # print standard starts--------------------------------end-------------------------------------

        # print fpdf file for question answer--------------------start-----------------------------------
        client_name = request.session['client_name']
        industry = request.session['industry']
        country = request.session['country']
        import os
        path = 'static/media/'+client_name+'/'+country+'/'+industry

        # download combined file--------------start------------    content_type="application/pdf///vnd.openxmlformats-officedocument.wordprocessingml.document"
        import os
        from django.conf import settings
        from django.http import HttpResponse, Http404

        client_name = request.session['client_name']
        industry = request.session['industry']
        country = request.session['country']
        import os
        path = 'static/media/'+client_name+'/'+country+'/'+industry
        with open("./"+path+"/combined_file.docx", 'rb') as f:
            data = f.read()
        response = HttpResponse(
            data, content_type="vnd.openxmlformats-officedocument.wordprocessingml.document")
        response['Content-Disposition'] = 'attachment; filename="RFP.docx"'

        return response
        # download combined file--------------start------------------------content_type="application/pdf--vnd.openxmlformats-officedocument.wordprocessingml.document"
        Question = request.GET.get("Question")

        return render(request, 'preview.html', {'Question': Question})


# preview page erase data and again start RFP
def start_new_rfp(request):
    #     showname=request.session['showname']
    docobj1 = Document.objects.filter(
        id__in=[1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 18, 19, 20, 21, 22, 23]).update(selected=" ")
#     UQL=UQ.objects.filter(user=showname)
#     UQL.delete()
#     ud=SelectDropQuery.objects.all()
#     ud.delete()
    return redirect('/RFP/index/')

# type question page print standard pdf


# upload files
def pic_upload_view(request):
    client_name = request.session['client_name']
    country = request.session['country']
    pic = request.FILES.get('file')
    prod = Product(user=client_name, picup=pic, clientgeo=country)
    prod.save()
    return render(request, 'preview.html', {'Question': Question})


# editans_view
def editans_view(request, id):
    industry = request.session['industry']
    country = request.session['country']
    showname = request.session['showname']
    id = id
    if request.method == "POST":

        pt = UQ.objects.get(pk=id)
        fm = UserQueryForm(request.POST, instance=pt)
        if fm.is_valid():
            fm.save()
    if request.method == "GET":
        pt = UQ.objects.get(pk=id)
        fm = UserQueryForm(instance=pt)

    return render(request, "editans.html", {"form": fm, "showname": showname, "country": country, "industry": industry, "id": id})

# pic_down_view


def pic_down_view(request):
    client_name = request.session['client_name']
    country = request.session['country']
    pic = request.FILES.get('file')
    First = First_Page_upload(user=client_name, picdown=pic, clientgeo=country)
    First.save()
    return render(request, 'preview.html', {'Question': Question})

# file_injection_view-Direct download


def file_injection_view(request):
    client_name = request.session['client_name']
    country = request.session['country']
    try:
        import os
        files = os.listdir("./media/First_page_upload/Drop")
        print(files)

        print("---------------------------------Download Injection file using Python--------------------------------")
        ## from pathlib import Path
        ## path1 = './media/pickupdir'
        ## name_list = os.listdir(path1)
        ## full_list = [os.path.join(path1, i) for i in name_list]
        ## time_sorted_list = sorted(full_list, key=os.path.getmtime)

        # lst_sorted = [str1.replace('\\', '/') for str1 in time_sorted_list]
        FP = Product.objects.filter(user=client_name, clientgeo=country)
        l = len(FP)
        if l == 1:
            FP = Product.objects.get(user=client_name, clientgeo=country)
            s = FP.picup
        else:
            s = FP[l-1].picup
        # filename1 = "C:/Users/shubhamjain35/Desktop/P Copy-801-showcase/KPM/media/pickupdir/Resolution_Life_Finance_Systems_RFP_Final.docx"
        # file = filename1.split("/")[-1]
        # print("Final File Path============", file)
        filename1 = s
        file = s
        filename1 = "media/"+str(filename1)
        print("Filename============",filename1)
        # tokens = ('is', 'does', 'do', 'what', 'when', 'where',
        #           'who', 'why', 'what', 'how', 'please', 'describe')

        # from docx import Document
        # import tqdm

        # document = Document(filename1)
        # #print("Questions in RFP "+filename1)
        # questions = []
        # sentences = []
        # number = 0
        # paragraphs = [para.text for para in document.paragraphs if para.text]

        # for index, paragraph in enumerate(paragraphs):
        #     text = paragraph.strip().lower()
        #     sentences.append(paragraph)
        #     if text.startswith(tokens) or text.endswith("?"):
        #         #         print(len(text))
        #         if len(text) > 100:
        #             number += 1
        #             print(str(number)+". ", paragraph)
        #             print(len(text))
        #             print()
        #             questions.append(paragraph)
        quest = Extraction()
        # filename=r"01 JHG_FMS SI Partner RFP Main Document final.docx"
        print("---"*10, "Start "+filename1+"---"*10)
        df_questions, list_questions = quest.question_search(filename1)
        print("Questions in RFP "+filename1)
        print("Number of questions detected from the document is :", len(list_questions))
        print("++++++"*20)
        num = 0
        list_questions = [text.strip() for text in list_questions]
        for question in list_questions:
            num+=1
            print(str(num)+". ", question)
            print("------"*20)
        print("---"*10, "End "+file+"---"*10)

        # showname = request.session['showname']
        # industry = request.session['industry']
        # country = request.session['country']
        # Query = request.GET.get("Query")
        # data = Question.objects.filter(country=country, industry=industry)
        # df = read_frame(data)
        # print(df.columns)
        # df.to_excel("Data.xlsx", index=False)

        # df = pd.read_excel("Question_Pairing_Final.xlsx",
        #                    sheet_name="Tablib Dataset")

         #---------------------------------------------------Search Engine Functionality(Start)-------------------------------------
        industry1= []
        country1=  []
        section1 = []
        # section=  ["Relevant Experience","Project Team Structure"]
        industry1.append(industry)
        country1.append(country)

        industry_text = "_".join(industry)
        print("Industry text:-",industry_text)
        country_text = "_".join(country)
        print("Country text:-",country_text)
        df = pd.read_excel("database.xlsx")

        # Creating an model
        model_build = Model_Buiding_approach(industry_text, country_text, model_name=model_name)

        print("Before Filtered DataFrame:- ", df.shape)
        # Filter the data based on Country and Industry
        df = model_build.dataframe_filter_return(df, industry1, country1, section1)

        print("Filtered DataFrame:- ", df.shape)
        # Training Dataset Vector
        if os.path.isfile("RFP/Embedding Models/corpus_embedding"+'-'+model_name+'-'+country_text+"-"+industry_text+".pt"):
            print("Need not to Run Training File as Model is Present")
        else:
            print("File is not present. Initializing the training command.")
            execution_time = model_build.training_dataset_vector(df, embedder = embedder)
            print("Execution of Training Time:-", execution_time)

        # Input the query
        # query = "Has your organisation been involved in any business-related litigation that may affect the performance of these Services in the past five years?"
        # query=Query

        # index_list, cos_sim_list = model_build.query_resolved(query=query, df=df, embedder=embedder)

        # print("Index List:-", index_list)
        # print("Cosine Similarity List", cos_sim_list)
        #---------------------------------------------------Search Engine Functionality(End) Extended-------------------------------------

        # Input the query

        # query = "Has your organisation been involved in any business-related litigation that may affect the performance of these Services in the past five years?"
        answer = []
        for q in list_questions:
            query = q
            index_list, cos_sim_list = model_build.query_resolved(
                query=query, df=df, embedder=embedder)
            if index_list:
                answer.append(df["Document_Link"][index_list[0]])

        #---------------------------------------------------Search Engine Functionality(End)-------------------------------------
        from fpdf import FPDF
        pdf = FPDF()
        pdf.add_page()
        # pdf.add_font('BB','',r'C:/Windows/Fonts/javatext.TTF', uni=True)
        pdf.set_font('Times', style='U', size=15)
        pdf.multi_cell(0, 5, 'Questions asked in uploaded RFP Document ')
        pdf.multi_cell(0, 5, '       ')
        count = 1

        for quest, ans in zip(list_questions, answer):
            if quest:
                quest = quest.encode('latin-1', 'replace').decode('latin-1')
            if ans:
                ans = ans.encode('latin-1', 'replace').decode('latin-1')
            pdf.set_font('Times', style='U', size=10)
            pdf.set_font('Times', size=9)
            pdf.multi_cell(0, 5, str(count)+': '+quest)
            pdf.multi_cell(0, 5, " ")
            pdf.multi_cell(0, 5, "Answer " + ': '+ans)
            pdf.multi_cell(0, 5, '' + '\n' + '')
            count = count+1

        from RFP.models import Document
        industry = request.session['industry']
        country = request.session['country']
        showname = request.session['showname']
        data = Question.objects.filter(country=country, industry=industry)
        Doc1 = Document.objects.filter(selected="on")
        pdf.output("./media/pickupdir/output1.pdf", 'F')

        # download combined file--------------start------------    content_type="application/pdf///vnd.openxmlformats-officedocument.wordprocessingml.document"
        import os
        from django.conf import settings
        from django.http import HttpResponse, Http404

        with open("./media/pickupdir/output1.pdf", 'rb') as f:
            data = f.read()
        response = HttpResponse(data, content_type="application/pdf")
        response['Content-Disposition'] = 'attachment; filename="RFPInjection.pdf"'
        return response
    except:
        industry = request.session['industry']
        country = request.session['country']
        showname = request.session['showname']
        from .models import Document
        p = Users.objects.filter(user=client_name)
        Doc1 = Document.objects.filter(selected="on", user__in=p)
        Userq = UQ.objects.filter(user=client_name)
        c = Userq.exists()

        User = askques.objects.filter(user=client_name, selected="on")
        Quest = User.exists()

        # display seleted index
        return render(request, 'firstpage.html', {"showname": showname, "country": country, "industry": industry, "Doc": Doc1, "Quest": Quest, "c": c})

    # download combined file--------------start------------------------content_type="application/pdf#vnd.openxmlformats-officedocument.wordprocessingml.document"

# drop_rfp_view


def drop_rfp_view(request):
    client_name = request.session['client_name']
    country = request.session['country']

    # files = os.listdir("./media/First_page_upload/Drop")
    # print(files)
    # from pathlib import Path
    # path1 = './media/pickupdir'
    # name_list = os.listdir(path1)
    # full_list = [os.path.join(path1, i) for i in name_list]
    # time_sorted_list = sorted(full_list, key=os.path.getmtime)
    # print(time_sorted_list)
    # lst_sorted = [str1.replace('\\', '/') for str1 in time_sorted_list]

    # filename1 = "C:/Users/shubhamjain35/Desktop/P Copy-801-showcase/KPM/media/pickupdir/Resolution_Life_Finance_Systems_RFP_Final.docx"
    # filename1 = lst_sorted[-1]
    FP = Product.objects.filter(user=client_name, clientgeo=country)
    l = len(FP)
    if l == 1:
        FP = Product.objects.get(user=client_name, clientgeo=country)
        s = FP.picup
    else:
        s = FP[l-1].picup
    # filename1 = "C:/Users/shubhamjain35/Desktop/P Copy-801-showcase/KPM/media/pickupdir/Resolution_Life_Finance_Systems_RFP_Final.docx"
    # file = filename1.split("/")[-1]
    # print("Final File Path============", file)
    filename1 = s
    file = s
    filename1 = "media/"+str(filename1)

    # tokens = ('is', 'does', 'do', 'what', 'when', 'where',
    #           'who', 'why', 'what', 'how', 'please', 'describe')

    # from docx import Document
    # import tqdm

    # document = Document(filename1)
    # #print("Questions in RFP "+filename1)
    # questions = []
    # sentences = []
    # number = 0
    # paragraphs = [para.text for para in document.paragraphs if para.text]

    # for index, paragraph in enumerate(paragraphs):
    #     text = paragraph.strip().lower()
    #     sentences.append(paragraph)
    #     if text.startswith(tokens) or text.endswith("?"):
    #         #         print(len(text))
    #         if len(text) > 100:
    #             number += 1
    #             print(str(number)+". ", paragraph)
    #             print(len(text))
    #             print()
    #             questions.append(paragraph)

    quest = Extraction()
    print("---"*10, "Start "+str(filename1)+"---"*10)
    df_questions,list_questions = quest.question_search(filename1)
    print("Questions in RFP "+filename1)
    print("Number of questions detected from the document is :", len(list_questions))
    print("++++++"*20)
    num = 0
    list_questions = [text.strip() for text in list_questions]
    for question in list_questions:
        num+=1
        print(str(num)+". ", question)
        print("------"*20)
    print("---"*10, "End "+str(filename1)+"---"*10)


    # Initializing Parameters and getting values from UI
    showname = request.session['showname']
    industry = request.session['industry']
    country = request.session['country']
    Query = request.GET.get("Query")
    # data=Question.objects.filter(country=country,industry=industry)
    # df = read_frame(data)
    # print(df.columns)
    # df.to_excel("Data.xlsx", index=False)

    # df = pd.read_excel("Question_Pairing_Final.xlsx",
    #                    sheet_name="Tablib Dataset")

    #---------------------------------------------------Search Engine Functionality(Start)-------------------------------------
    industry1= []
    country1=  []
    section1 = []
    # section=  ["Relevant Experience","Project Team Structure"]
    industry1.append(industry)
    country1.append(country)

    industry_text = "_".join(industry1)
    print("Industry text:-",industry_text)
    country_text = "_".join(country1)
    print("Country text:-",country_text)

    df = pd.read_excel("database.xlsx")

    # Creating an model
    model_build = Model_Buiding_approach(industry_text, country_text, model_name=model_name)

    print("Before Filtered DataFrame:- ", df.shape)
    # Filter the data based on Country and Industry
    df = model_build.dataframe_filter_return(df, industry1, country1, section1)



    print("Filtered DataFrame:- ", df.shape)
    # Training Dataset Vector
    if os.path.isfile("RFP/Embedding Models/corpus_embedding"+'-'+model_name+'-'+country_text+"-"+industry_text+".pt"):
        print("Need not to Run Training File as Model is Present")
    else:
        print("File is not present. Initializing the training command.")
        execution_time = model_build.training_dataset_vector(df, embedder = embedder)
        print("Execution of Training Time:-", execution_time)

    # Input the query

    # query = "Has your organisation been involved in any business-related litigation that may affect the performance of these Services in the past five years?"
    answer = []
    for q in list_questions:
        query = q
        index_list, cos_sim_list = model_build.query_resolved(
            query=query, df=df, embedder=embedder)
        if index_list:
            answer.append(df["Document_Link"][index_list[0]])
    id = 1
    for ques in list_questions:
        query = ques
        index_list, cos_sim_list = model_build.query_resolved(
            query=query, df=df, embedder=embedder)
        answer = [df["Document_Link"][index] for index in index_list]
        UserQuery = DropQuery(id=id, question=ques, user=showname,
                              answer1=answer[0], answer2=answer[1], answer3=answer[2])
        UserQuery.save()
        id = id+1

    Quest = DropQuery.objects.all()

    return render(request, "drop_rfp.html", {"questions": Quest, "showname": showname, "country": country, "industry": industry})


def drop_rfpquest_view1(request):
    industry = request.session['industry']
    country = request.session['country']
    showname = request.session['showname']
    client_name = request.session['client_name']
    if request.method == "POST":

        Tick1 = request.POST.get("1")
        Tick2 = request.POST.get("2")
        Tick3 = request.POST.get("3")
        Tick4 = request.POST.get("4")
        Tick5 = request.POST.get("5")
        Tick6 = request.POST.get("6")
        Tick7 = request.POST.get("7")
        Tick8 = request.POST.get("8")
        Tick9 = request.POST.get("9")
        Tick10 = request.POST.get("10")
        p = [Tick1, Tick2, Tick3, Tick4, Tick5,
             Tick6, Tick7, Tick8, Tick9, Tick10]
        answer = DropQuery.objects.filter(id__in=p)
        final_list = filter(None, p)
        ids = 1
        data = Question.objects.filter(country=country, industry=industry)
        # df = read_frame(data)
        # print(df.columns)
        # df.to_excel("Data.xlsx", index=False)
        # df = pd.read_excel("Question_Pairing_Final.xlsx",
        #                    sheet_name="Tablib Dataset")

        
        #---------------------------------------------------Search Engine Functionality(Start)-------------------------------------
        industry= ["Agnostic"]
        country=  ["AU"]
        # section=  ["Relevant Experience","Project Team Structure"]
        section = []

        industry_text = "_".join(industry)
        print("Industry text:-",industry_text)
        country_text = "_".join(country)
        print("Country text:-",country_text)
        df = pd.read_excel("database.xlsx")

        # Creating an model
        model_build = Model_Buiding_approach(industry_text, country_text, model_name=model_name)

        print("Before Filtered DataFrame:- ", df.shape)
        # Filter the data based on Country and Industry
        df = model_build.dataframe_filter_return(df, industry, country, section)



        print("Filtered DataFrame:- ", df.shape)
        # Training Dataset Vector
        if os.path.isfile("RFP/Embedding Models/corpus_embedding"+'-'+model_name+'-'+country_text+"-"+industry_text+".pt"):
            print("Need not to Run Training File as Model is Present")
        else:
            print("File is not present. Initializing the training command.")
            execution_time = model_build.training_dataset_vector(df, embedder = embedder)
            print("Execution of Training Time:-", execution_time)

        for c in answer:
            index_list, cos_sim_list = model_build.query_resolved(
                query=c.question, df=df, embedder=embedder)
            answer = [df["Document_Link"][index] for index in index_list]

            Select = SelectDropQuery(id=ids, user=showname, question=c.question,
                                     answer1=answer[0], answer2=answer[1], answer3=answer[2])
            Select.save()
            u = SelectDropQuery.objects.filter(
                user=showname).order_by('id').last()
            ids = u.id+1

        Select = SelectDropQuery(
            id=ids, user=showname, question="Please click on finish")
        Select.save()
        qt = SelectDropQuery.objects.all()
        # qt=DropQuery.objects.all()
        print(qt)
        try:
            p = SelectDropQuery.objects.filter(
                user=showname).order_by('id').first()
            # p=qt[0]
        except:
            p = "No Questions to display"
        print("PPPPPPPPPPPPPPPPPPPPPPPPPPP")
        print(p)
        print("PPPPPPPPPPPPPPPPPPPPPPPPPPP")
        return render(request, "drop_rfp_quest1.html", {"p": p, "showname": showname, "country": country, "industry": industry})

# drop_rfpquest_view2


def drop_rfpquest_view2(request, id):
    industry = request.session['industry']
    country = request.session['country']
    showname = request.session['showname']
    client_name = request.session['client_name']
    qt = DropQuery.objects.filter(select1="on") | DropQuery.objects.filter(
        select2="on") | DropQuery.objects.filter(select3="on")
    if request.method == "POST":

        p = SelectDropQuery.objects.filter(
            id__gt=id).exclude(id=id).order_by('id').first()
        # qt=DropQuery.objects.all()filter(optionselect="on").

        if p:
            answer1 = request.POST.get("1")
            if answer1:
                try:
                    obj1 = SelectDropQuery.objects.filter(
                        id=p.id-1).update(select1="on")
                except:
                    pass
            answer2 = request.POST.get("2")
            if answer2:
                try:
                    obj1 = SelectDropQuery.objects.filter(
                        id=p.id-1).update(select2="on")
                except:
                    pass
            answer3 = request.POST.get("3")
            if answer3:
                try:
                    obj1 = SelectDropQuery.objects.filter(
                        id=p.id-1).update(select3="on")
                except:
                    pass
            return render(request, "drop_rfp_quest2.html", {"p": p, "showname": showname, "country": country, "industry": industry})
        qt = SelectDropQuery.objects.all()
        return render(request, "drop_rfp_preview.html", {"qt": qt, "showname": showname, "country": country, "industry": industry})


# drop_rfp_preview_view
def drop_rfp_preview_view(request):
    industry = request.session['industry']
    country = request.session['country']
    showname = request.session['showname']
    client_name = request.session['client_name']

    qt = SelectDropQuery.objects.filter(user=showname).filter(select1="on") | SelectDropQuery.objects.filter(
        user=showname).filter(select2="on") | SelectDropQuery.objects.filter(user=showname).filter(select3="on")
    # qt=SelectDropQuery.objects.all()

    return render(request, "drop_rfp_preview.html", {"qt": qt, "showname": showname, "country": country, "industry": industry})


# drop_editans_view
def drop_editans_view(request, id):
    industry = request.session['industry']
    country = request.session['country']
    showname = request.session['showname']
    if request.method == "POST":
        pt = SelectDropQuery.objects.get(pk=id)

        fm = SelectDropQueryForm(request.POST, instance=pt)
        if fm.is_valid():
            fm.save()
    if request.method == "GET":
        pt = SelectDropQuery.objects.get(pk=id)
        fm = SelectDropQueryForm(instance=pt)

    return render(request, "editans_drop.html", {"form": fm, "showname": showname, "country": country, "industry": industry, "id": id})

# drop_drop_print_pdf


def drop_print_pdf(request):
    industry = request.session['industry']
    country = request.session['country']
    showname = request.session['showname']
    client_name = request.session['client_name']
    import os
    path = 'static/media/'+client_name+'/'+country+'/'+industry
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    ct = SelectDropQuery.objects.filter(user=showname).filter(select1="on") | SelectDropQuery.objects.filter(
        user=showname).filter(select2="on") | SelectDropQuery.objects.filter(user=showname).filter(select3="on")
    # pdf.add_font('BB','',r'C:/Windows/Fonts/javatext.TTF', uni=True)
    pdf.set_font('Times', style='U', size=15)
    pdf.multi_cell(
        0, 5, 'Autodetected Question and Responses from uploaded RFP' + '\n')
    pdf.multi_cell(0, 5, '\n')
# query = "Has your organisation been involved in any business-related litigation that may affect the performance of these Services in the past five years?"
    for i in ct:
        if i.question:
            a = i.question
            a = a.encode('latin-1', 'replace').decode('latin-1')
        if i.answer1:
            b = i.answer1
            b = b.encode('latin-1', 'replace').decode('latin-1')
        if i.answer2:
            c = i.answer2
            c = c.encode('latin-1', 'replace').decode('latin-1')
        if i.answer3:
            d = i.answer3
            d = d.encode('latin-1', 'replace').decode('latin-1')
        # pdf.set_font('BB',style ='U',size=10)
        pdf.set_font('Times', size=10)
        pdf.multi_cell(0, 5, 'Question:'+a + '\n')
        pdf.multi_cell(0, 5, " ")
        pdf.set_font('Times', size=9)
        if i.select1 == "on":
            pdf.multi_cell(0, 5, 'Answer:'+b)
            pdf.multi_cell(0, 5, '' + '\n' + '')
        if i.select2 == "on":
            pdf.multi_cell(0, 5, 'Answer:'+c)
            pdf.multi_cell(0, 5, '' + '\n' + '')
        if i.select3 == "on":
            pdf.multi_cell(0, 5, 'Answer:'+d)
            pdf.multi_cell(0, 5, '' + '\n' + '')
    import os
    path = 'static/media/'+client_name+'/'+country+'/'+industry
    try:
        os.makedirs(path)
        pdf.output("./"+path+"/uploaded_file.pdf", 'F')
    except:
        pdf.output("./"+path+"/uploaded_file.pdf", 'F')

    import os
    from django.conf import settings
    from django.http import HttpResponse, Http404

    with open("./"+path+"/uploaded_file.pdf", 'rb') as f:
        data = f.read()
    response = HttpResponse(data, content_type="application/pdf")
    response['Content-Disposition'] = 'attachment; filename="RFP_Uploaded.pdf"'

    return response
    # return render(request,"drop_rfp_preview.html",{"qt":ct})


# def user_upload_question_view(request):
#     industry=request.session['industry']
#     country=request.session['country']
#     showname=request.session['showname']
#     form=UserQuestionForm()
#     if request.method=="POST":
#         form=UserQuestionForm(request.POST)
#         if form.is_valid():
#             form.save()
#     form=UserQuestionForm()
#     return render(request,'UserQuestion.html',{'form':form,"showname":showname,"country":country,"industry":industry})


def user_upload_question_view(request, id):
    industry = request.session['industry']
    country = request.session['country']
    showname = request.session['showname']
    if request.method == "GET":
        id = id
        return render(request, 'UserQuestion.html', {"showname": showname, "country": country, "industry": industry, "id": id})
    if request.method == "POST":

        Rfp_Id = request.POST.get("Rfp_Id")
        Content_Type = request.POST.get("content_type")
        Country = request.POST.get("country")
        Industry = request.POST.get("industry")
        Section = request.POST.get("section")
        Sub_Section = request.POST.get("sub_section")
        Question = request.POST.get("question")
        file = request.FILES.get('file')
        print(file)
        c = UserQuestion.objects.create(rfp_id=Rfp_Id, content_type=Content_Type, country=Country,
                                        industry=Industry, Section=Section, Sub_Section=Sub_Section, question=Question, Document_link=file)
        c.save()
        # pt=SelectDropQuery.objects.get(pk=id)

    return render(request, 'UserQuestion.html', {"showname": showname, "country": country, "industry": industry})


def confirm_view(request, id):
    industry = request.session['industry']
    country = request.session['country']
    showname = request.session['showname']
    if request.method == "GET":
        UQL = UserQuery.objects.filter(id=id)
        UQL.update(sentapproval="on", viewed="off")
        return render(request, 'confirmation.html', {"showname": showname, "country": country, "industry": industry, "id": id})
    if request.method == "POST":

        return render(request, 'confirmation.html', {"showname": showname, "country": country, "industry": industry, "id": id})


def approve_view(request):
    #     industry=request.session['industry']
    #     country=request.session['country']
    #     showname=request.session['showname']
    # if request.method == "GET":
    #     SP = UserQuery.objects.filter(sentapproval="on").filter(viewed="off")
    #     return render(request, 'approve.html', {"SP": SP})
    # if request.method == "POST":
    #     SP = UserQuery.objects.filter(sentapproval="on").filter(viewed="off")
    #     return render(request, 'approve.html', {"SP": SP})
    if request.method == "GET":
        SP = documentapproval.objects.filter(approved="No")
        return render(request, 'approve.html', {"SP": SP})
    if request.method == "POST":
        SP = documentapproval.objects.filter(approved="No")
        return render(request, 'approve.html', {"SP": SP})


def approved_view(request, id):
    industry = request.session['industry']
    country = request.session['country']
    showname = request.session['showname']
    if request.method == "GET":
        SP = documentapproval.objects.filter(id=id)
        SP.update(approved="Yes")
        SP = documentapproval.objects.filter(approved="No")
        Yes = documentapproval.objects.filter(approved="Yes")
        return render(request, 'approve.html', {"showname": showname, "country": country, "industry": industry, "SP": SP, "Yes": Yes})
    if request.method == "POST":
        SP = documentapproval.objects.filter(id=id)
        SP.update(approved="Yes")
        SP = documentapproval.objects.filter(approved="No")
        Yes = documentapproval.objects.filter(approved="Yes")
        return render(request, 'approve.html', {"showname": showname, "country": country, "industry": industry, "SP": SP, "Yes": Yes})


def disapproved_view(request, id):
    industry = request.session['industry']
    country = request.session['country']
    showname = request.session['showname']
    if request.method == "GET":
        SP = documentapproval.objects.filter(id=id)
        SP.update(approved="No")
        SP = documentapproval.objects.filter(approved="No")
        Yes = documentapproval.objects.filter(approved="Yes")
        return render(request, 'approve.html', {"showname": showname, "country": country, "industry": industry, "SP": SP, "Yes": Yes})
    if request.method == "POST":
        SP = documentapproval.objects.filter(id=id)
        SP.update(approved="No")
        SP = documentapproval.objects.filter(approved="No")
        Yes = documentapproval.objects.filter(approved="Yes")
        return render(request, 'approve.html', {"showname": showname, "country": country, "industry": industry, "SP": SP, "Yes": Yes})


def login(request):
    if request.method == 'POST':
        username = request.POST['username']
        password = request.POST['password']
        user = auth.authenticate(username=username, password=password)
        if user is not None:
            auth.login(request, user)
            return redirect("/approve/request")
        else:
            messages.success(request, 'Incorrect Username or password')
            return redirect('login')
    else:
        return render(request, 'login.html')


def convertToBinaryData(filename):
    # Convert digital data to binary format
    with open(filename, 'rb') as file:
        binaryData = file.read()
    return binaryData


def writeTofile(data, filename):
    # Convert binary data to proper format and write it on Hard Disk
    with open(filename, 'wb') as file:
        file.write(data)


def Binaryfile_view(request):
    docfile = convertToBinaryData(
        "C:/Users/shubhamjain35/Desktop/RFP-Builder-New-main/SampleBinary.docx")

    c = Question.objects.create(id=310, Binaryword=docfile)
    c.save()

    fetch = Question.objects.get(id=310)
    v = fetch.Binaryword

    writeTofile(
        v, "C:/Users/shubhamjain35/Desktop/RFP-Builder-New-main/SampleBinary13.docx")

    return render(request, 'index.html')


def chatgpt_view(request):
    # openai.api_key = os.getenv("OPENAI_API_KEY")
    industry = request.session['industry']
    country = request.session['country']
    showname = request.session['showname']
    client_name = request.session['client_name']
    ap.api_key = "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"
    if request.method == 'POST':
        gtp_question = request.POST['gtp_question']

        response = ap.Completion.create(
            model="text-davinci-003",
            # prompt="I am a highly intelligent question answering bot. If you ask me a question that is rooted in truth, I will give you the answer. If you ask me a question that is nonsense, trickery, or has no clear answer, I will respond with \"Unknown\".\n\nQ: What is human life expectancy in the United States?\nA: Human life expectancy in the United States is 78 years.\n\nQ: Who was president of the United States in 1955?\nA: Dwight D. Eisenhower was president of the United States in 1955.\n\nQ: Which party did he belong to?\nA: He belonged to the Republican Party.\n\nQ: What is the square root of banana?\nA: Unknown\n\nQ: How does a telescope work?\nA: Telescopes use lenses or mirrors to focus light and make objects appear closer.\n\nQ: Where were the 1992 Olympics held?\nA: The 1992 Olympics were held in Barcelona, Spain.\n\nQ: How many squigs are in a bonk?\nA: Unknown\n\nQ: Where is the Valley of Kings?\nA:",
            # prompt="I am a highly intelligent question answering bot. If you ask me a question that is rooted in truth, I will give you the answer. If you ask me a question that is nonsense, trickery, or has no clear answer, I will respond with \"Unknown\".\n\nQ: What is human life expectancy in the United States?\nA: ",
            # prompt="I am a highly intelligent question answering bot. If you ask me a question that is rooted in truth, I will give you the answer. If you ask me a question that is nonsense, trickery, or has no clear answer, I will respond with \"Unknown\".\n\nQ: What are the different components of workday?\nA: ",
            prompt="I am a highly intelligent question answering bot. If you ask me a question that is rooted in truth, I will give you the answer. If you ask me a question that is nonsense, trickery, or has no clear answer, I will respond with \"Unknown\".\n\nQ:"+gtp_question+"?\nA: ",
            temperature=0,
            max_tokens=100,
            top_p=1,
            frequency_penalty=0.0,
            presence_penalty=0.0,
            stop=["\n"]
        )
        chat = response['choices'][0]['text'].strip()
        return render(request, 'chatgpt.html', {'c': chat, "gtp_question": gtp_question, "showname": showname, "country": country, "industry": industry})

    if request.method == 'GET':

        return render(request, 'chatgpt.html', {"showname": showname, "country": country, "industry": industry})
# question = "What are the different components of workday"
# answer = openai(question)
# answer = openai(question)


def SelectedIndex_view(request):

    industry = request.session['industry']
    country = request.session['country']
    showname = request.session['showname']
    client_name = request.session['client_name']
    ap.api_key = "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"
    if request.method == 'GET':
        p = Users.objects.get(user=client_name)
        Doccopy = Document.objects.filter(user=p)
        Document_usercopy.objects.filter(user=client_name).delete()
        for docu in Doccopy:
            c = Document_usercopy.objects.create(
                country=docu.country, industry=docu.industry, doc_index=docu.doc_index, user=client_name, File=docu.File)

        show = Document_usercopy.objects.filter(user=client_name).first()
        print("SelectedIndex_view_Get")
        print(show.id)
        print("SelectedIndex_view_Get")
        return render(request, 'SelectedIndex.html', {"showname": showname, "country": country, "industry": industry, "show": show})
    if request.method == 'POST':
        # CHECKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKK
        print('checkkkkkkkkkkk im here request post', request.POST)
        client_name = request.session['client_name']
        industry = request.session['industry']
        country = request.session['country']
        radio = request.POST.get("flexRadioDefault")
        Quest = request.POST.get("Quest")

        request_post_list = dict(request.POST).keys()
        print(request.POST, 'post dataaaa')
        print(request_post_list, 'post list')

        check_user = Users.objects.filter(user=client_name)

        print(check_user, 'check user')
        # exit(0)
        if check_user:
            d = check_user[0]

        else:
            d = Users(user=client_name)
            d.save()
            print(d, 'user saved')

        if radio:

            #u = Users.objects.get(user=client_name)

            Imagee = Image.objects.get(
                id=radio)

            l = Image.objects.filter(id=radio)
            t = Image.objects.filter(id=radio).update(selected="on")
            image_url = Image.objects.get(id=radio)

            print(l, 'lllll')
            print(image_url.image.url, 'urlll of image')
            if client_name in l:
                pass
            else:
                try:
                    c = Imagee.user.add(d)
                except:
                    pass

                l = Imagee.user.all()

        if Quest:
            print('inside ques')
            ck = askques.objects.filter(user=client_name)
            print(ck, 'ckkkk')
            if not ck:
                askquestion = askques(user=client_name, selected="on")
                askquestion.save()
            if ck:
                askquestion = askques.objects.filter(
                    user=client_name).update(selected="on")
        else:
            check = askques.objects.filter(user=client_name)
            if check:
                askque = askques.objects.filter(
                    user=client_name).update(selected=" ")
            else:
                print(client_name, 'inside else')
                askque = askques(user=client_name, selected=" ")
                askque.save()
        
        for i in request_post_list:
            try:
                print(int(i), 'integer')
                if i:

                    # docobj1 = RfpSection.objects.filter(id=i).update(selected="on")
                    docobj1 = RfpSection.objects.get(id=i).user.add(d)
                    print('updated user to the rfp section')
                else:
                    try:
                        docobj1 = RfpSection.objects.get(id=i, user=d)
                        c = docobj1.user.remove(d)
                    except:
                        pass
            except Exception as e:
                print(e, 'exception')

        
    # checkkkkkkkkkkkkkkkkkkkkkkkkkkkkkkkkkkkkkkkkkkkk

        p = Users.objects.get(user=client_name)
        Doccopy = RfpSection.objects.filter(user=p)

        print(Doccopy, 'doccopyy')
        
        # Document_usercopy.objects.filter(user=client_name).delete()
        for docu in Doccopy:

            if docu.document_link:
                file_path = f'https://rfpstoragecheck.blob.core.windows.net/data/{docu.industry}/{docu.country}/Content/{docu.document_link}'
            else:
                file_path = f'https://rfpstoragecheck.blob.core.windows.net/data/{docu.industry}/{docu.country}/Content/default.docx'

            c = Document_usercopy.objects.update_or_create(
                country=docu.country, industry=docu.industry, doc_index=docu.section_data, user=client_name, file_link=file_path)
            # c.save()
            
            print(c, 'created user copy')
        try:
            gtp_question = request.POST['gtp_question']

            response = ap.Completion.create(
                model="text-davinci-003",
                # prompt="I am a highly intelligent question answering bot. If you ask me a question that is rooted in truth, I will give you the answer. If you ask me a question that is nonsense, trickery, or has no clear answer, I will respond with \"Unknown\".\n\nQ: What is human life expectancy in the United States?\nA: Human life expectancy in the United States is 78 years.\n\nQ: Who was president of the United States in 1955?\nA: Dwight D. Eisenhower was president of the United States in 1955.\n\nQ: Which party did he belong to?\nA: He belonged to the Republican Party.\n\nQ: What is the square root of banana?\nA: Unknown\n\nQ: How does a telescope work?\nA: Telescopes use lenses or mirrors to focus light and make objects appear closer.\n\nQ: Where were the 1992 Olympics held?\nA: The 1992 Olympics were held in Barcelona, Spain.\n\nQ: How many squigs are in a bonk?\nA: Unknown\n\nQ: Where is the Valley of Kings?\nA:",
                # prompt="I am a highly intelligent question answering bot. If you ask me a question that is rooted in truth, I will give you the answer. If you ask me a question that is nonsense, trickery, or has no clear answer, I will respond with \"Unknown\".\n\nQ: What is human life expectancy in the United States?\nA: ",
                # prompt="I am a highly intelligent question answering bot. If you ask me a question that is rooted in truth, I will give you the answer. If you ask me a question that is nonsense, trickery, or has no clear answer, I will respond with \"Unknown\".\n\nQ: What are the different components of workday?\nA: ",
                prompt="I am a highly intelligent question answering bot. If you ask me a question that is rooted in truth, I will give you the answer. If you ask me a question that is nonsense, trickery, or has no clear answer, I will respond with \"Unknown\".\n\nQ:"+gtp_question+"?\nA: ",
                temperature=0,
                max_tokens=100,
                top_p=1,
                frequency_penalty=0.0,
                presence_penalty=0.0,
                stop=["\n"]
            )
            chat = response['choices'][0]['text'].strip()
            print(chat, 'response from chat gpt--------')
            show = Document_usercopy.objects.filter(user=client_name, doc_index='Executive Summary').first()

            print(show, 'all user data')
            
            print("SelectedIndex_view_POST")

            print("SelectedIndex_view_POST")
            print(show, 'show data')

            return render(request, 'SelectedIndex.html', {'c': chat, "gtp_question": gtp_question, "showname": showname, "country": country, "industry": industry, "show": show})
        except:
            show = Document_usercopy.objects.filter(user=client_name).first()
            print("SelectedIndex_view_POST_except")

            print("SelectedIndex_view_POST_except")
            # print(show, vars(show), showname, '-----------')
            return render(request, 'SelectedIndex.html', {"showname": showname, "country": country, "industry": industry, "show": show, "image_url": image_url})


def SelectedIndex2_view(request, id):
    industry = request.session['industry']
    country = request.session['country']
    showname = request.session['showname']
    client_name = request.session['client_name']
    if request.method == 'GET':
        p = Users.objects.get(user=client_name)
        Doccopy = Document.objects.filter(user=p)

        for docu in Doccopy:
            c = Document_usercopy.objects.create(country=docu.country, industry=docu.industry, doc_index=docu.doc_index,
                                                 File=docu.File)

        return render(request, 'SelectedIndex2.html', {"showname": showname, "country": country, "industry": industry, "Doccopy": Doccopy})
    if request.method == 'POST':
        Queryyy = request.POST.get("Queryyy")
        print("SelectedIndex2_view_POST")
        print(Queryyy)
        print("SelectedIndex2_view_POST")
        show2 = Document_usercopy.objects.filter(user=client_name).filter(
            id__gt=id).exclude(id=id).order_by('id').first()
        if not show2:
            # show2 = Document_usercopy.objects.filter(user=client_name).filter(
            #     id__gt=id).exclude(id=id).order_by('id').first()
            # print("SelectedIndex2_view_POST_not_Query")
            # print(show2.id)
            # print("SelectedIndex2_view_POST_not_Query")
            return render(request, 'SelectedIndexlastPage.html', {"showname": showname, "country": country, "industry": industry})
        if show2:
            showname = request.session['showname']
            industry = request.session['industry']
            country = request.session['country']
            client_name = request.session['client_name']
            Queryyy = request.POST.get("Queryyy")
            data = Question.objects.filter(country=country, industry=industry)
            show2 = Document_usercopy.objects.filter(user=client_name).filter(
                id__gt=id).exclude(id=id).order_by('id').first()
            
            if show2.doc_index == 'Executive Summary':
                return render(request, 'SelectedIndex.html', {"showname": showname, "country": country, "industry": industry, "show": show2})
            
            return render(request, 'SelectedIndex2.html', {'data': data, "showname": showname, "country": country, "industry": industry, "show2": show2})


def SelectedIndexlastPage_view(request):
    industry = request.session['industry']
    country = request.session['country']
    showname = request.session['showname']
    client_name = request.session['client_name']
    return render(request, 'SelectedIndexlastPage.html', {"showname": showname, "country": country, "industry": industry})


def Onscreenmcq_view(request, id):

    if request.method == "POST":
        showname = request.session['showname']
        industry = request.session['industry']
        country = request.session['country']
        client_name = request.session['client_name']
        Queryyy = request.POST.get("Queryyy")
        data = Question.objects.filter(country=country, industry=industry)
        print("$$$$$$$$$$$$$$$$$$$$")
        print(id)
        print("$$$$$$$$$$$$$$$$$$$$")
        # df = read_frame(data)
        # df.to_excel("Data.xlsx", index=False)
        # df = pd.read_excel("Question_Pairing_Final.xlsx",
        #                    sheet_name="Tablib Dataset")


        #---------------------------------------------------Search Engine Functionality(Start)-------------------------------------
        industry= ["Agnostic"]
        country=  ["AU"]
        # section=  ["Relevant Experience","Project Team Structure"]
        section = []

        industry_text = "_".join(industry)
        print("Industry text:-",industry_text)
        country_text = "_".join(country)
        print("Country text:-",country_text)
        df = pd.read_excel("database.xlsx")

        # Creating an model
        model_build = Model_Buiding_approach(industry_text, country_text, model_name=model_name)

        print("Before Filtered DataFrame:- ", df.shape)
        # Filter the data based on Country and Industry
        df = model_build.dataframe_filter_return(df, industry, country, section)



        print("Filtered DataFrame:- ", df.shape)
        # Training Dataset Vector
        if os.path.isfile("RFP/Embedding Models/corpus_embedding"+'-'+model_name+'-'+country_text+"-"+industry_text+".pt"):
            print("Need not to Run Training File as Model is Present")
        else:
            print("File is not present. Initializing the training command.")
            execution_time = model_build.training_dataset_vector(df, embedder = embedder)
            print("Execution of Training Time:-", execution_time)

        # Input the query
        # query = "Has your organisation been involved in any business-related litigation that may affect the performance of these Services in the past five years?"
        queryyy = Queryyy

        request.session['queryyy'] = queryyy
        non = ""
        index_list, cos_sim_list = model_build.query_resolved(
            query=queryyy, df=df, embedder=embedder)

        data = Question.objects.filter(id__in=index_list)
        data0 = ""
        data1 = ""
        data2 = ""
        id0 = ""
        id1 = ""
        id2 = ""
        if len(index_list) == 1:
            data0 = data[0].Model_Image_Link_BLOB
            id0 = data[0].id
        elif len(index_list) == 2:
            data0 = data[0].Model_Image_Link_BLOB
            id0 = data[0].id
            data1 = data[1].Model_Image_Link_BLOB
            id1 = data[1].id
        elif len(index_list) == 3:
            data0 = data[0].Model_Image_Link_BLOB
            id0 = data[0].id
            data1 = data[1].Model_Image_Link_BLOB
            id1 = data[1].id
            data2 = data[2].Model_Image_Link_BLOB
            id2 = data[2].id
        else:
            non = "No matching question found"
            # data0=data[0].Questions
            # id0=data[0].id
            # data1=data[1].Questions
            # id1=data[1].id
            # data2=data[2].Questions
           # id2=data[2].id
        # l = Question.objects.get(id=170)

        # f = l.File
        show2 = Document_usercopy.objects.filter(
            id=id)
        return render(request, 'SelectedIndex2.html', {'non': non, 'data': data,  "data0": data0, "data1": data1, "data2": data2, "id0": id0, "id1": id1, "id2": id2, "Query": Queryyy, "showname": showname, "country": country, "industry": industry, "show2": show2})


# documentapproval files
def documentapproval_view(request):
    client_name = request.session['client_name']
    country = request.session['country']
    fileapp = request.FILES.get('file')
    prod = documentapproval(
        user=client_name, documentapproval=fileapp, clientgeo=country)
    prod.save()

    return render(request, 'SelectedIndexlastPage.html')
