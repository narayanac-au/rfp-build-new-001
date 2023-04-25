# Approach 1 Updated code (for single file) : Testing
import warnings
import fitz
from unidecode import unidecode
import pandas as pd
import re
import numpy as np
import docx
from docx import Document
import re
import sys
import os
import io
import tqdm
import glob
import time
from doc2pdf import convert
import nltk
from nltk.tokenize import word_tokenize
nltk.download('averaged_perceptron_tagger')
nltk.download('punkt')
warnings.filterwarnings('ignore')


class Extraction:

    def detect_question_or_request(self, sentence):
        """ 
        This code defines a function detect_question_or_request that takes a sentence as input and returns a string indicationg whether
        the sentence is a question, request, or statement.

            Ex: Input 1: What is the annual revenue of your company ?
            Output: Question
            Input 2: Brief history of your company.
            Output: Request 
            """
        # Tokenize the sentence
        question_words = ['what', 'when', 'where', 'who', 'whom', 'which', 'whose', 'why', 'how',
                          'how long', 'how many', 'how much', 'how far', 'how often', 'have', 'has', 'do', 'did', 'does']
        verbs_and_modal_verbs = ["can", "could", "will", "would", "may", "might", "shall", "should", "must", 'please', 'define',
                                 'describe', 'outline', 'provide', 'brief', 'summarise', 'explain', 'submit', 'propose', 'write', 'detail', 'discuss']
        tokens = [token.lower() for token in word_tokenize(sentence)]

        # Check if sentence ends with a question mark or sentence contains any question words
        if sentence.strip().endswith("?") or any(sentence.lower().startswith(word) for word in question_words):
            return "Question"
        # Check if sentence starts with a verb or modal verb
        elif tokens[0].lower() in verbs_and_modal_verbs:
            return "Request"
        else:
            return "Statement"

    def question_search(self, file):
        """ 
        This code defines a function question_search that takes a file as input and returns two outputs. One is a dataframe indicationg whether
        the sentence is a question, request, or statement and other one is list which is having all the questions and request sentences

            Ex: Input 1: file which is doc/docx/pdf format
            Output 1: dataframe consists of two columns (Sentence, Sentence Type)
            Output 2: List consists of questions and request sentences 
            """

        questions = []
        results = {}
        unwanted_questions = [
            "please also refer to appendix", "please see appendix", "Appendix"]

        if file.endswith(".docx"):
            para_text = ''
            table_text = ''
            # Load the Word document
            doc = docx.Document(file)

            for paragraph in doc.paragraphs:
                para_text += f'{paragraph.text.strip()}|'

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        table_text += f'{cell.text.strip()}|'
                    table_text += '\n'
            text = para_text+table_text
            new_text = re.sub(r"\n([a-z])", r"\g<1>", text)
            new_text = new_text.split('|')
            new_text = np.concatenate([new_text]).tolist()
            cleaned_new_text = []
            for element in new_text:
                if '\n' in element:
                    new_elements = element.split('\n')
                    cleaned_new_text.append(new_elements)
                else:
                    cleaned_new_text.append(element)
            paragraphs = []
            for sub_list in cleaned_new_text:
                if type(sub_list) == list:
                    for item in sub_list:
                        paragraphs.append(item)
                else:
                    paragraphs.append(sub_list)
            paragraphs = [para for para in paragraphs if len(para) > 8]

        elif file.endswith(".pdf"):
            text = ""
            doc = fitz.open(file)
            for page in doc:
                text += page.get_text()
            new_text = re.sub(r"\n([a-z])", r"\g<1>", text)
            paragraphs = new_text.split('\n')
            paragraphs = [para.strip()
                          for para in paragraphs]  # To remove white space
            # To remove if sentence having only symbols/numbers/lessthan 8 characters
            paragraphs = [para for para in paragraphs if len(para) > 8]

        for paragraph in paragraphs:
            if re.search('[a-zA-Z]+', paragraph) != None:
                # To remove numbers or symbols at start of paragraph
                paragraph = paragraph[re.search(
                    r"[a-zA-Z]", paragraph).start():]
                # Looping through each paragraph to find whether it is a question or request or statement
                if self.detect_question_or_request(paragraph) == "Question" or self.detect_question_or_request(paragraph) == "Request":
                    results[paragraph] = self.detect_question_or_request(
                        paragraph)
                else:
                    results[paragraph] = "Statement"
                    questions.append(paragraph)

        # To see the result dataframe of all the sentences (question or request or statement)
        df_results = pd.DataFrame(list(results.items()), columns=[
                                  'Sentence', 'Sentence_type'])

        # To see the questions and requets in the form of list of sentences
        questions_list = df_results.loc[(df_results['Sentence_type'] == 'Question') | (
            df_results['Sentence_type'] == 'Request')]['Sentence'].values.tolist()
        questions_list_filtered = [ques for ques in questions_list if not any(re.search(
            un_ques, ques, re.IGNORECASE) for un_ques in unwanted_questions)]  # To remove unwanted sentences
        questions_list_filtered = [ques_1 for ques_1 in questions_list_filtered if len(
            ques_1.split()) > 2]  # To remove short questions (i.e. lessthan two words)
        return df_results, questions_list_filtered
