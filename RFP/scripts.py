import datetime
import wget
import os
from docx.shared import Pt
from docx.shared import Inches
from python_docx_replace import docx_replace, docx_blocks
from datetime import datetime
import aspose.words as aw
import pathlib
from docx import Document as Document_compose
from docxcompose.composer import Composer
from docx import Document
import platform

from django.conf import settings
from docxtpl import DocxTemplate, RichText
import shutil


def docx_template_replace(doc_path, doc_name, client_name='', title=''):
    doc = DocxTemplate(doc_path)
    dictionary = {}
    dictionary['client_name'] = client_name
    curr_date = datetime.today()
    dictionary['curr_date'] = curr_date.strftime("%B %d, %Y")
    dictionary['title'] = title

    doc.render(dictionary)

    temp_name = 'updated' + '_' + str(datetime.now()).replace(
        ' ', '_').replace('.', '_').replace(':', '_').replace('-', '_')
    temp_file = f'temporary/{temp_name}.docx'
    print(temp_file, 'temp file')
    doc.save(temp_file)
    os.remove(doc_path)
    media_path = f'media/{client_name}'
    isExist = os.path.exists(media_path)

    if not isExist:

        # Create a new directory because it does not exist
        os.makedirs(media_path)
        print("The new directory is created!")

    shutil.move(temp_file, f'{media_path}/{doc_name}')
    # os.rename(temp_file, f'{media_path}/{doc_name}')
    return f'{media_path}/{doc_name}'


# file_data = ["C:/Users/narayanac/Documents/RFP-Builder-Latest/Title_template2.docx","C:/Users/narayanac/Documents/RFP-Builder-Latest/Title_template3.docx"]

# file_data = ['Title_template2.docx', 'Title_template3.docx']

# def combine_word_documents(files):
#     print('entered')
#     merged_document = Document()

#     for index, file in enumerate(files):
#         print(index, file)
#         sub_doc = Document(file)
#         print(sub_doc, 'su')

#         # Don't add a page break if you've reached the last file.
#         if index < len(files)-1:
#            sub_doc.add_page_break()

#         for element in sub_doc.element.body:
#             merged_document.element.body.append(element)

#     merged_document.save('merged-test.docx')
#     return 'meged-test.docx'

# combine_word_documents(file_data)


# path = 'static/media/'+client_name+'/'+country+'/'+industry

# files_data = ["Agnostic_USA_Executive Summary.docx","Agnostic_USA_Implementation timeline.docx"]


def combine_all_docx(filename_master, files_list):
    print('inside combine', files_list)
    number_of_sections = len(files_list)
    master = Document_compose(filename_master)
    composer = Composer(master)
    for i in range(0, number_of_sections):
        print(settings.BASE_DIR, 'base')

        # file_path = os.path.join(settings.BASE_DIR, files_list[i])
        directory = os.getcwd()
        # print(directory, 'directoryhy')
        string_path = directory + files_list[i]
        print(string_path, 'path string')
        # file_path = pathlib.PurePath(directory, files_list[i])
        # print(file_path, 'file path')

        doc_temp = Document_compose(string_path)
        composer.append(doc_temp)
    composer.save("test_merge_1.docx")
    return 'test_merge_1.docx'

# path = 'C:/Users/narayanac/Documents/RFP-Builder-Latest/'
# filename_master = "C:/Users/narayanac/Documents/RFP-project/RFP-Builder-Latest/C.docx"
# combine_all_docx(filename_master, files_data)


def replace_word_document(client_name, replace_name, doc_path, doc_name):
    print(client_name, replace_name, 'arguments')
    doc = Document_compose(doc_path)
    Dictionary = {replace_name: client_name}
    for i in Dictionary:
        for p in doc.paragraphs:
            if p.text.find(i) >= 0:
                p.text = p.text.replace(i, Dictionary[i])

    # delete the existing updated document
    # os.remove('updated.docx')

    # save changed document
    doc.save('temporary/updated.docx')
    os.remove(doc_path)
    media_path = f'media/{client_name}'
    isExist = os.path.exists(media_path)

    if not isExist:

        # Create a new directory because it does not exist
        os.makedirs(media_path)
        print("The new directory is created!")

    shutil.move('temporary/updated.docx', f'{media_path}/{doc_name}')
    # os.rename('temporary/updated.docx', f'{media_path}/{doc_name}')
    return f'{media_path}/{doc_name}'


def replace_word_document_aspose(client_name, replace_name, doc_path):

    # load Word document
    # doc = aw.Document("document.docx")
    doc = aw.Document(doc_path)

    # replace text
    # doc.range.replace(cliend_name, "[CLIENT]", aw.replacing.FindReplaceOptions(aw.replacing.FindReplaceDirection.FORWARD))
    doc.range.replace(client_name, replace_name, aw.replacing.FindReplaceOptions(
        aw.replacing.FindReplaceDirection.FORWARD))

    # save the modified document
    # doc.save("updated.docx")
    doc.save('update.docx')
    print(doc, 'doc inside update')
    return doc

# import urllib

# def get_document(file_path, file_name):
#     print(file_path, file_name, 'file path')
#     document = urllib.request.urlretrieve (file_path, 'test_create_001.docx')
#     print(document, 'document output')
#     return document


def get_document(file_path):
    print(file_path, 'file path')

    file_name = wget.download(file_path)

    return file_name


# fileNames = [ "Input1.docx", "Input2.docx" ]

def merge_files(fileNames):
    output = aw.Document()
    # Remove all content from the destination document before appending.
    output.remove_all_children()

    for fileName in fileNames:

        directory = os.getcwd()
        # print(directory, 'directoryhy')
        string_path = directory + fileName
        print(string_path, 'path string')

        input = aw.Document(string_path)
        # Append the source document to the end of the destination document.
        output.append_document(
            input, aw.ImportFormatMode.KEEP_SOURCE_FORMATTING)

    output.save("aspose_merge.docx")
    return 'aspose_merge.docx'


def replace_word_title():
    # print(client_name, replace_name, 'arguments')
    doc = Document_compose('../Title_test.docx')
    # doc=Document_compose('../Healthcare_Australia_Executive Summary.docx')
    # Dictionary = {replace_name: client_name}
    print(doc, 'doc data')
    Dictionary = {"[CLIENT_NAME]": 'KPMG Test 001',
                  "[CURR_DATE]": datetime.today()}

    for i in range(0, 20):
        print(i, 'loop')

    for i in Dictionary:
        print(i, 'ii test')
        for p in doc.paragraphs:
            print(p, 'ppp')
            print(p.text, 'text')
            if p.text.find(i) >= 0:
                print('inside if')
                print(i, '--------', Dictionary[i], 'Dico')
                p.text = p.text.replace(i, Dictionary[i])

        for p in doc.title:
            print(p, 'ppp')
            print(p.text, 'text')
            if p.text.find(i) >= 0:
                print('inside if')
                print(i, '--------', Dictionary[i], 'Dico')
                p.text = p.text.replace(i, Dictionary[i])

    # delete the existing updated document
    # os.remove('updated.docx')

    # save changed document
    doc.save('replace_updated.docx')
    # os.remove(doc_path)
    # media_path = f'media/{client_name}'
    # isExist = os.path.exists(media_path)

    # if not isExist:

    #     # Create a new directory because it does not exist
    #     os.makedirs(media_path)
    #     print("The new directory is created!")

    # os.rename('temporary/updated.docx', f'{media_path}/{doc_name}')
    return 'replace_updated.docx'


def python_docx():
    # get your document using python-docx
    doc = Document("../Title_test.docx")

    Dictionary = {"[CLIENT_NAME]": 'KPMG Test 001',
                  "[CURR_DATE]": datetime.today()}

    # call the replace function with your key value pairs
    docx_replace(doc, **Dictionary)

    # call the blocks function with your sets
    docx_blocks(doc, signature=True, table_of_contents=False)

    # remove the first table in the Word document
    # docx_remove_table(doc, 0)

    # do whatever you want after that, usually save the document
    doc.save("replaced.docx")

    return 'replace_updated.docx'


def replace_aspose_word(doc_path, client_name):

    print(client_name, 'arguments')

    doc = Document_compose(doc_path)
    Dictionary = {
        'Evaluation Only. Created with Aspose.Words. Copyright 2003-2023 Aspose Pty Ltd.': '',
        'This document was truncated here because it was created in the Evaluation Mode.': '',
        'Proposed Plan on a Page': ''
    }

    for i in Dictionary:
        for p in doc.paragraphs:
            if p.text.find(i) >= 0:
                p.text = p.text.replace(i, Dictionary[i])

    # save changed document
    doc.save('temporary/updated_aspose.docx')
    os.remove(doc_path)
    media_path = f'media/{client_name}'
    isExist = os.path.exists(media_path)

    if not isExist:
        # Create a new directory because it does not exist
        os.makedirs(media_path)

    os.rename('temporary/updated_aspose.docx', f'{media_path}/merged_rfp.docx')
    return f'{media_path}/merged_rfp.docx'


def create_images_doc(image_object):
    print(image_object, 'images objectsss')
    print(vars(image_object), 'whats inside image object')

    docume = Document()
    docume.add_heading('IMAGE', 0)

    for i in image_object:
        image_url = f"https://rfpstoragecheck.blob.core.windows.net/rfpstorage/Section_Documents/{i.industry}/{i.country}/Images/{i.image_link}"
        get_image = get_document(image_url)

        p = docume.add_paragraph()
        runner = p.add_run("Images:")
        runner.bold = True

        docume.add_picture(get_image)
    docume.save('test_image_file.docx')
    return 'test_image_file.docx'


def write_header_footer(doc_path):
    # from docx import Document
    # from docx.shared import Pt

    # read  doc
    # doc = Document('../output-node-merger-v4.docx')
    doc = Document(doc_path)

    # define your style font and size for the header
    style_h = doc.styles['Header']
    font_h = style_h.font
    font_h.name = 'Arial'
    font_h.bold = True
    font_h.size = Pt(20)

    # define your style font and size for the footer
    style_f = doc.styles['Footer']
    font_f = style_f.font
    font_f.name = 'Arial'
    font_f.bold = False
    font_f.size = Pt(2)

    # for the header:
    # header = doc.sections[0].header
    # paragraph_h = header.paragraphs[0]
    # paragraph_h.text = 'January 2021 KPMG' # insert new value here.
    # paragraph_h.style = doc.styles['Header'] # this is what changes the style

    header = doc.sections[0].header
    paragraph_h = header.paragraphs[0]

    logo_run = paragraph_h.add_run()
    logo_run.add_picture("logo.png", width=Inches(1))

    # text_run = paragraph.add_run()
    # text_run.text = '\t' + "My Awesome Header" # For center align of text
    # text_run.style = "Heading 2 Char"

    # for the footer:
    footer = doc.sections[0].footer
    paragraph_f = footer.paragraphs[0]
    # insert new value here.
    paragraph_f.text = '© 2023 Copyright owned by one or more of the KPMG International entities. KPMG International entities provide no services to clients. All rights reserved. KPMG refers to the global organization or to one or more of the member firms of KPMG International Limited (“KPMG International”), each of which is a separate legal entity. KPMG International Limited is a private English company limited by guarantee and does not provide services to clients. For more detail about our structure please visit https://home.kpmg/governance'
    # paragraph_f.text = '© 2023 Copyright owned by one or more of the KPMG International entities.'

    paragraph_f.style = doc.styles['Footer']  # this is what changes the style

    doc.save('final_merged_document_v1.docx')
    return 'final_merged_document_v1.docx'


if __name__ == "__main__":
    # test = replace_word_title()
    # test = python_docx()
    # v = ['/TEST-MAY-04-2023-001_Healthcare_AU_Implementation Approach.docx', '/test_image_file_PuATo0v.docx']
    # test = merge_files(v)
    test = write_header_footer()
    print(test, 'testtt')
