from docx import Document
import os 
from datetime import datetime, date
# pip install azure-storage-blob
from azure.storage.blob import BlobServiceClient


def replace_word_doc(doc_path, client_full_name='', client_short_name='', client_geo='', client_address_line1='',
                        client_address_line2='', client_zip_code='', powered_industry='', 
                        kpmg_geo='', kpmg_address='', kpmg_lead_partner='', doc_name=''):
    """
    This function will replace words inside document based on parameters.It also saves output in document format.
    """

    ## Document Reading 
    doc=Document(doc_path)

    # Initializing parameters
    dictionary = {}

    if client_full_name:
        dictionary['[CLIENT_NAME]'] = client_full_name
    if client_short_name:
        dictionary['[CLIENT_SHORT_NAME]'] = client_short_name
    if client_geo:
        dictionary['[CLIENT_GEO]'] = client_geo
    if client_address_line1:
        dictionary['[CLIENT_ADDRESS_LINE1]'] = client_address_line1
    if client_address_line2:
        dictionary['[CLIENT_ADDRESS_LINE2]'] = client_address_line2
    if client_zip_code:
        dictionary['[CLIENT_ZIP_CODE]'] = client_zip_code
    if powered_industry:
        dictionary['[POWERED_INDUSTRY]'] = powered_industry
    if kpmg_geo:
        dictionary['[KPMG_GEO]'] = kpmg_geo
    if kpmg_address:
        dictionary['[KPMG_ADDRESS]'] = kpmg_address
    if kpmg_lead_partner:
        dictionary['[KPMG_LEAD_PARTNER]'] = kpmg_lead_partner

    curr_date = datetime.today()
    dictionary['[CURR_DATE]'] = curr_date.strftime("%B %d, %Y")

    print(dictionary, 'dicccc')
    # Replacing words with headingss
    # for i in dictionary:
    #     for p in doc.headings:
    #         if p.text.find(i)>=0:
    #             p.text=p.text.replace(i,dictionary[i])

    for p in doc.paragraphs:
        inline = p.runs
        for i in range(len(inline)):
            text = inline[i].text
            for key in dictionary.keys():
                if key in text:
                    text=text.replace(key,dictionary[key])
                    inline[i].text = text

    
    # # Replacing words with paragraphs
    # for i in dictionary:
    #     for p in doc.paragraphs:
    #         if p.text.find(i)>=0:
    #             p.text=p.text.replace(i,dictionary[i])

    # # Replacing words with paragraphs
    for i in dictionary:
        for p in doc.paragraphs:
            if p.text.find(i)>=0:
                p.text=p.text.replace(i,dictionary[i])

    # Replacing words inside table
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.find(i)>=0:
                    cell.text=cell.text.replace(i,dictionary[i])

    #save changed document
    # temp_name = str(datetime.now()).replace(':', '-').replace('.', '-')
    temp_name = 'updated'
    temp_file = f'temporary/{temp_name}.docx'
    print(temp_file, 'temp file')
    doc.save(temp_file)
    os.remove(doc_path)
    media_path = f'media/{client_full_name}'
    isExist = os.path.exists(media_path)

    if not isExist:

        # Create a new directory because it does not exist
        os.makedirs(media_path)
        print("The new directory is created!")

    os.rename(temp_file, f'{media_path}/{client_full_name}_{doc_name}')
    return f'{media_path}/{client_full_name}_{doc_name}'


def upload_blob_data(subfolder, filename, container_id):
    """
    This method upload the file with respect to sub-folder in azure blob storage.
    """
    # Initializing connection using connection string
    storage_connection_string = "DefaultEndpointsProtocol=https;AccountName=rfpstoragecheck;AccountKey=Dg56EQWUFqhnYfZR2YKnoRNb1LqwSF66+MD9xqZdH9da0B0SXZEcMQRwQhFO4Q+Tc4+QNnQiO/ws+AStVHBwDg==;EndpointSuffix=core.windows.net"
    blob_service_client = BlobServiceClient.from_connection_string(storage_connection_string)
    blob_service_client.get_container_client(container_id)
    
    overwrite = True
    blob_path = subfolder+"/"+filename
    blob_obj = blob_service_client.get_blob_client(container=container_id, blob=blob_path)
    with open(filename, mode='rb') as file_data:
        blob_obj.upload_blob(file_data, overwrite=overwrite)

    blob_url = "https://rfpstoragecheck.blob.core.windows.net/"+container_id+"/"+blob_path
    print(blob_url)
    return blob_url