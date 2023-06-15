import subprocess
from RFP.scripts import (
    replace_word_document,
    get_document,
    write_header_footer,
    merge_files,
    docx_template_replace,
    replace_aspose_word,
    create_images_doc,
)
from RFP.scripts import (
    replace_word_document,
    get_document,
    combine_all_docx,
    merge_files,
    docx_template_replace,
    replace_aspose_word,
    create_images_doc,
)
from wsgiref.util import FileWrapper
from django.http import HttpResponse, Http404
from django.conf import settings
from RFP.scripts import (
    replace_word_document,
    get_document,
    combine_all_docx,
    merge_files,
    docx_template_replace,
    replace_aspose_word,
)
import threading
from .replace_parameters_doc import replace_word_doc, upload_blob_data
from RFP.scripts import (
    replace_word_document,
    get_document,
    combine_all_docx,
    merge_files,
)
from django.core.files import File
from scripts import replace_word_document, get_document
from django.db.models import Q
import openai as ap
from dataclasses import dataclass
from email.charset import BASE64
from urllib import response
from django.shortcuts import render, redirect
from docx import Document
from RFP.models import *
from RFP.read_Data import read_Data
from RFP.models import Question as QP
from RFP.models import UserQuery as UQ
from PyPDF2 import PdfFileReader, PdfFileWriter, PdfFileMerger
from docx2pdf import convert
from RFP.Question_Extraction import Extraction
from docx import Document as CX
from django_pandas.io import read_frame
from RFP.Question_Similarity import Model_Buiding_approach
import pandas as pd
from .forms import UserQueryForm, SelectDropQueryForm
from .models import Image
from django.contrib.auth.models import User, auth
from django.contrib import messages
from sentence_transformers import SentenceTransformer, util
import os
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse
import json
from .models import KPMGgeo as kg
from .models import First_Page_upload as FPU
from .models import Product, ExtraImage
from .forms import ProjectForm
import nltk

from django.conf import settings
from django.http import HttpResponse, Http404, JsonResponse
from django.contrib.auth.models import User
from wsgiref.util import FileWrapper
nltk.download("averaged_perceptron_tagger")

nltk.download('punkt')
# Loading the Embedding Pretrained Model


def load_model():
    # Initializing the embedding model
    try:
        embedder = SentenceTransformer("./RFP/bert-base-uncased")
    except:
        embedder = SentenceTransformer("bert-base-uncased")
    return embedder


model_name = "bert-base-uncased"
embedder = load_model()

# # Reading the data directly from table using connection
# read = read_Data()
# df_rfpdata = read.read_dataframe("RFP_rfpdata")


# first page
def index_view(request):
    # showname=request.session['showname']

    # docobj1 = Document.objects.filter(
    # id__in=[1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 18, 19, 20, 21, 22, 23]).update(selected=" ")
    #     UQL=UQ.objects.all()
    #     UQL.delete()
    #     ud=SelectDropQuery.objects.all()
    #     ud.delete()
    # Dp=DropQuery.objects.filter(user=showname)
    # Dp.delete()
    return render(request, "index.html")


# second page details


def infodetails_view(request):
    if request.method == "GET":
        geo = KPMGgeo.objects.all()

        return render(request, "info.html", {"geo": geo})


@csrf_exempt
def geoadd_view(request):
    if request.method == "POST":
        id = request.POST["geo"]
        add = list(KPMGadd.objects.filter(KPMGgeo=id).values())

        return JsonResponse({"add": add}, safe=False)


# third page select standard RFP


def doc_content_view(request):
    if request.method == "POST":
        KPMGaddress1 = request.POST.get("address1")
        clientaddress_line1 = request.POST.get("address_line1")
        clientaddress_line2 = request.POST.get("address_line2")
        clientPostal_Code = request.POST.get("Postal_Code")
        client_name = request.POST.get("clientname")
        client_name = client_name.replace(" ", "_")
        username = request.POST.get("username")
        email = request.POST.get("email")
        showname = request.POST.get("showname")
        industry = request.POST.get("industry")
        country = request.POST.get("countries")
        address = request.POST.get("address")
        KPMGgeol = request.POST.get("geo")
        KPMGgeoo = kg.objects.filter(id=KPMGgeol)
        KPMGgeo = str(KPMGgeoo[0])
        kpmg_lead = request.POST.get("KPMGLEADPARTNER")
        kpmg_add = request.POST.get("kpmg_address")
        TitleforStyleSheetSelected = request.POST.get("TitleforStyleSheetSelected")
        request.session["TitleforStyleSheetSelected"] = TitleforStyleSheetSelected
        print("TitleforStyleSheetSelected",TitleforStyleSheetSelected)
        extrcount = "ABC"
        Acccount = "ABC"
        selfirst = "ABC"
        sellast = "ABC"
        log = "ABC"
        request.session["extrcount"] = extrcount
        request.session["Acccount"] = Acccount
        request.session["selfirst"] = selfirst
        request.session["sellast"] = sellast
        request.session["log"] = log
        print(industry, "industry")
        print(country, "country before session data")
        request.session["username"] = username
        request.session["client_name"] = client_name
        request.session["industry"] = industry
        request.session["country"] = country
        request.session["showname"] = showname
        request.session["address"] = address

        request.session["add_line_1"] = clientaddress_line1
        request.session["add_line_2"] = clientaddress_line2

        request.session["client_geo"] = country
        request.session["kpmg_geo"] = KPMGgeo
        request.session["client_zipcode"] = clientPostal_Code
        request.session["kpmg_lead"] = kpmg_lead
        request.session["kpmg_address"] = kpmg_add
        # formm = ImageForm(data=request.POST, files=request.FILES)

        # formm.save()

        client_name = request.session["client_name"]

        infos = info.objects.update_or_create(
            clientfullname=client_name,
            clientshortname=showname,
            clientindustry=industry,
            clientgeo=country,
            clientaddress_line1=clientaddress_line1,
            clientaddress_line2=clientaddress_line2,
            clientPostal_Code=clientPostal_Code,
            KPMGaddress1=KPMGaddress1,
            KPMGgeo=KPMGgeo,
        )
        # creating user for adminpanel
        try:
            username = username.upper()
            password = "RFP"+username+"@321"
            user = User.objects.create_user(
                username=username, password="RFP"+username+"@321", is_staff=False, email=email)
            user.save()
        except:
            pass

        # creating user for adminpanel
        # infos.save()
        import os

        path = "static/media/" + client_name + "/" + country + "/" + industry
        try:
            os.makedirs(path)

        except:
            pass
        """
        p = Users.objects.filter(user=client_name)
        # Doc = Document.objects.filter(
        #     country=country, industry=industry, user__in=p)
        Doc = Document.objects.filter(
            country=country, industry=industry)[:3]
        # Doc = list(Doc)
        print(Doc, 'top3_doc')
        UserDoc = Document.objects.filter(
            country=country, industry=industry, user__in=p).exclude(id__in=Doc)
        print(UserDoc, 'user_doc')
        default_user_doc = list(set(Doc) | set(UserDoc))

        print(default_user_doc, 'default docs')

        Poc = Document.objects.filter(
            country=country, industry=industry).exclude(id__in=Doc).exclude(id__in=UserDoc)

        """
        user = Users.objects.filter(user=client_name)

        default_rfp = (
            RfpSection.objects.filter(
                industry=industry, country=country, is_default=True
            )
            .order_by("order")
            .values_list("id", flat=True)
        )
        default_rfp = list(default_rfp)

        all_sections = RfpSection.objects.filter(
            industry=industry, country=country
        ).order_by("order")

        # Doc = RfpSection.objects.filter(industry=industry, country=country, is_default=True).order_by('order')

        # Doc = RfpSection.objects.filter(industry=industry, country=country)

        print(all_sections, country, industry, "doccc")

        user_doc = RfpSection.objects.filter(
            industry=industry, country=country, user__in=user
        ).values_list("id", flat=True)

        # UserDoc = RfpSection.objects.filter(
        #     Q(industry=industry) | Q(country=country), Q(user__in=user)).exclude(id__in=Doc).order_by('order')

        print(user_doc, "user_doc")
        # default_user_doc = list(set(Doc) | set(UserDoc))
        # default_user_doc = Doc | UserDoc

        # print(default_user_doc, 'default docs')

        # Poc = RfpSection.objects.filter(
        #     industry=industry, country=country).exclude(id__in=default_user_doc).order_by('order')

        # print(default_user_doc, 'default ,  user doc from rfpsection ')

        que = askques.objects.filter(user=client_name, selected="on")

        if que:
            que
        else:
            que = None
        from .models import Image

        industry = request.POST.get("industry")
        img = Image.objects.filter(caption=industry)
        print(Image, "imggg")
        print("CHEEEEEEEEEEEEEEEEEEKKKKKKKKKKKKKKKKKKKKKKKKKKK")
        print(img)
        for i in img:
            print(i.image_link)
        print("CHEEEEEEEEEEEEEEEEEEKKKKKKKKKKKKKKKKKKKKKKKKKKK")
        # print(Doc, 'Doc')
        print()
        print()
        print("-----")
        # print(Poc, 'poc')
        return render(
            request,
            "doc_content.html",
            {
                "all_sections": all_sections,
                "user_selected": user_doc,
                "showname": showname,
                "country": country,
                "industry": industry,
                "que": que,
                "Image": img,
                "default_sections": default_rfp,
                # "password": password,
                "username": username
            },
        )
    if request.method == "GET":
        client_name = request.session["client_name"]
        industry = request.session["industry"]
        country = request.session["country"]
        showname = request.session["showname"]

        user = Users.objects.filter(user=client_name)
        default_rfp = (
            RfpSection.objects.filter(
                industry=industry, country=country, is_default=True
            )
            .order_by("order")
            .values_list("id", flat=True)
        )
        default_rfp = list(default_rfp)
        all_sections = RfpSection.objects.filter(
            industry=industry, country=country
        ).order_by("order")

        print(all_sections, country, industry, "doccc")

        user_doc = RfpSection.objects.filter(
            industry=industry, country=country, user__in=user
        ).values_list("id", flat=True)

        print(user_doc, "user_doc")

        que = askques.objects.filter(user=client_name, selected="on")
        if que:
            que
        else:
            que = None
        from .models import Image

        # Image = Image.objects.filter(caption=industry)
        try:
            SelectedImage = Image.objects.filter(
                caption=industry, user__in=user)
        except:
            SelectedImage = None
        if SelectedImage:
            img = Image.objects.filter(caption=industry).exclude(user__in=user)

        else:
            img = Image.objects.filter(caption=industry)
        print("CHEEEEEEEEEEEEEEEEEEKKKKKKKKKKKKKKKKKKKKKKKKKKK")
        print(img)
        print("CHEEEEEEEEEEEEEEEEEEKKKKKKKKKKKKKKKKKKKKKKKKKKK")
        # print(Doc, 'Doc')
        print()
        print()
        print("-----")
        # print(Poc, 'poc')
        return render(
            request,
            "doc_content.html",
            {
                "all_sections": all_sections,
                "user_selected": user_doc,
                "showname": showname,
                "country": country,
                "industry": industry,
                "que": que,
                "Image": img,
                "SelectedImage": SelectedImage,
                "default_sections": default_rfp,
            },
        )


# fourth page type question RFP


def firstpage_view(request):
    if request.method == "POST":
        client_name = request.session["client_name"]
        industry = request.session["industry"]
        country = request.session["country"]
        extrcount = request.session["extrcount"]
        radio = request.POST.get("flexRadioDefault")
        Tick1 = request.POST.get("1")
        Tick2 = request.POST.get("2")
        Tick3 = request.POST.get("3")
        Tick4 = request.POST.get("4")
        Tick5 = request.POST.get("5")
        Tick6 = request.POST.get("6")
        Tick7 = request.POST.get("7")
        Tick8 = request.POST.get("8")
        Tick9 = request.POST.get("9")
        Tick10 = request.POST.get("10")
        Tick11 = request.POST.get("11")
        Tick12 = request.POST.get("12")
        Tick13 = request.POST.get("13")
        Tick14 = request.POST.get("14")
        Tick15 = request.POST.get("15")
        Tick16 = request.POST.get("16")
        Tick17 = request.POST.get("17")
        Tick18 = request.POST.get("18")
        Tick19 = request.POST.get("19")
        Tick20 = request.POST.get("20")
        Tick21 = request.POST.get("21")
        Tick22 = request.POST.get("22")
        Tick23 = request.POST.get("23")
        Quest = request.POST.get("Quest")

        # from .models import Image
        # p = Users.objects.filter(user=client_name)
        # print("NAMEEEEEEEEEEEEEEEEEEEEEEE")
        # print(p)
        # print("NAMEEEEEEEEEEEEEEEEEEEEEEE")
        # print("clear")
        # p.Image_set.all()
        # print("clear")
        # print("IMAAAAAAAAAAAAAGGEEE")
        # print(p.Image_set.all())
        # print("IMAAAAAAAAAAAAAGGEEE")

        check_user = Users.objects.filter(user=client_name)

        if check_user:
            d = check_user[0]

        else:
            d = Users(user=client_name)
            d.save()

        if radio:
            # u = Users.objects.get(user=client_name)

            Imagee = Image.objects.get(id=radio)

            l = Image.objects.filter(id=radio)
            t = Image.objects.filter(id=radio).update(selected="on")

            if client_name in l:
                pass
            else:
                try:
                    c = Imagee.user.add(d)
                except:
                    pass

                l = Imagee.user.all()

        if Quest:
            ck = askques.objects.filter(user=client_name)
            if not ck:
                askquestion = askques(user=client_name, selected="on")
                askquestion.save()
            if ck:
                askquestion = askques.objects.filter(user=client_name).update(
                    selected="on"
                )
        else:
            check = askques.objects.filter(user=client_name)
            if check:
                askque = askques.objects.filter(
                    user=client_name).update(selected=" ")
            else:
                askque = askques.objects.create(user=client_name, selected=" ")

        if Tick1:
            print("inside tick 111")

            docobj1 = Document.objects.filter(id=1).update(selected="on")
            docobj1 = Document.objects.get(id=1).user.add(d)

        else:
            print("inside tick 111 elseeee")
            try:
                docobj1 = Document.objects.get(id=1, user=d)
                c = docobj1.user.remove(d)
            except:
                pass

        if Tick2:
            docobj2 = Document.objects.filter(id=2).update(selected="on")
            docobj2 = Document.objects.get(id=2).user.add(d)

        else:
            try:
                docobj2 = Document.objects.get(id=2, user=d)
                c = docobj2.user.remove(d)
            except:
                pass

        if Tick3:
            docobj3 = Document.objects.filter(id=3).update(selected="on")
            docobj3 = Document.objects.get(id=3).user.add(d)

        else:
            try:
                docobj3 = Document.objects.get(id=3, user=d)
                c = docobj3.user.remove(d)
            except:
                pass

        if Tick4:
            docobj4 = Document.objects.filter(id=4).update(selected="on")
            docobj4 = Document.objects.get(id=4).user.add(d)

        else:
            try:
                docobj4 = Document.objects.get(id=4, user=d)
                c = docobj4.user.remove(d)
            except:
                pass

        if Tick5:
            docobj5 = Document.objects.filter(id=5).update(selected="on")
            docobj5 = Document.objects.get(id=5).user.add(d)

        else:
            try:
                docobj5 = Document.objects.get(id=5, user=d)
                c = docobj5.user.remove(d)
            except:
                pass

        if Tick6:
            docobj6 = Document.objects.filter(id=6).update(selected="on")
            docobj6 = Document.objects.get(id=6).user.add(d)

        else:
            try:
                docobj6 = Document.objects.get(id=6, user=d)
                c = docobj6.user.remove(d)
            except:
                pass

        if Tick7:
            docobj7 = Document.objects.filter(id=7).update(selected="on")
            docobj7 = Document.objects.get(id=7).user.add(d)

        else:
            try:
                docobj7 = Document.objects.get(id=7, user=d)
                c = docobj7.user.remove(d)
            except:
                pass

        if Tick8:
            docobj8 = Document.objects.filter(id=8).update(selected="on")
            docobj8 = Document.objects.get(id=8).user.add(d)

        else:
            try:
                docobj8 = Document.objects.get(id=8, user=d)
                c = docobj8.user.remove(d)
            except:
                pass

        if Tick9:
            docobj9 = Document.objects.filter(id=9).update(selected="on")
            docobj9 = Document.objects.get(id=9).user.add(d)

        else:
            try:
                docobj9 = Document.objects.get(id=9, user=d)
                c = docobj9.user.remove(d)
            except:
                pass

        if Tick10:
            docobj10 = Document.objects.filter(id=10).update(selected="on")
            docobj10 = Document.objects.get(id=10).user.add(d)

        else:
            try:
                docobj10 = Document.objects.get(id=10, user=d)
                c = docobj10.user.remove(d)
            except:
                pass

        if Tick11:
            docobj11 = Document.objects.filter(id=11).update(selected="on")
            docobj11 = Document.objects.get(id=11).user.add(d)

        else:
            try:
                docobj11 = Document.objects.get(id=11, user=d)
                c = docobj11.user.remove(d)
            except:
                pass

        if Tick12:
            docobj12 = Document.objects.filter(id=12).update(selected="on")
            docobj12 = Document.objects.get(id=12).user.add(d)

        else:
            try:
                docobj12 = Document.objects.get(id=12, user=d)
                c = docobj12.user.remove(d)
            except:
                pass

        if Tick13:
            docobj13 = Document.objects.filter(id=13).update(selected="on")
            docobj13 = Document.objects.get(id=13).user.add(d)

        else:
            try:
                docobj13 = Document.objects.get(id=13, user=d)
                c = docobj13.user.remove(d)
            except:
                pass

        if Tick14:
            docobj14 = Document.objects.filter(id=14).update(selected="on")
            docobj14 = Document.objects.get(id=14).user.add(d)

        else:
            try:
                docobj14 = Document.objects.get(id=14, user=d)
                c = docobj14.user.remove(d)
            except:
                pass

        if Tick15:
            docobj15 = Document.objects.filter(id=15).update(selected="on")
            docobj15 = Document.objects.get(id=15).user.add(d)

        else:
            try:
                docobj15 = Document.objects.get(id=15, user=d)
                c = docobj15.user.remove(d)
            except:
                pass

        if Tick16:
            docobj16 = Document.objects.filter(id=16).update(selected="on")
            docobj16 = Document.objects.get(id=16).user.add(d)

        else:
            try:
                docobj16 = Document.objects.get(id=16, user=d)
                c = docobj16.user.remove(d)
            except:
                pass

        if Tick17:
            docobj17 = Document.objects.filter(id=17).update(selected="on")
            docobj17 = Document.objects.get(id=17).user.add(d)

        else:
            try:
                docobj17 = Document.objects.get(id=17, user=d)
                c = docobj17.user.remove(d)
            except:
                pass

        if Tick18:
            docobj18 = Document.objects.filter(id=18).update(selected="on")
            docobj18 = Document.objects.get(id=18).user.add(d)

        else:
            try:
                docobj18 = Document.objects.get(id=18, user=d)
                c = docobj18.user.remove(d)
            except:
                pass

        if Tick19:
            docobj19 = Document.objects.filter(id=19).update(selected="on")
            docobj19 = Document.objects.get(id=19).user.add(d)

        else:
            try:
                docobj19 = Document.objects.get(id=19, user=d)
                c = docobj19.user.remove(d)
            except:
                pass

        if Tick20:
            docobj20 = Document.objects.filter(id=20).update(selected="on")
            docobj20 = Document.objects.get(id=20).user.add(d)

        else:
            try:
                docobj20 = Document.objects.get(id=20, user=d)
                c = docobj20.user.remove(d)
            except:
                pass

        if Tick21:
            docobj21 = Document.objects.filter(id=21).update(selected="on")
            docobj21 = Document.objects.get(id=21).user.add(d)

        else:
            try:
                docobj21 = Document.objects.get(id=21, user=d)
                c = docobj21.user.remove(d)
            except:
                pass

        if Tick22:
            docobj22 = Document.objects.filter(id=22).update(selected="on")
            docobj22 = Document.objects.get(id=22).user.add(d)

        else:
            try:
                docobj22 = Document.objects.get(id=22, user=d)
                c = docobj22.user.remove(d)
            except:
                pass

        if Tick23:
            docobj23 = Document.objects.filter(id=23).update(selected="on")
            docobj23 = Document.objects.get(id=23).user.add(d)

        else:
            try:
                docobj23 = Document.objects.get(id=23, user=d)
                c = docobj23.user.remove(d)
            except:
                pass

        client_name = request.session["client_name"]
        industry = request.session["industry"]
        country = request.session["country"]
        showname = request.session["showname"]
        data = Question.objects.filter(country=country, industry=industry)
        # data=Question.objects.all()
        # print("!!!!!!!!!!!!!!!!!!!!!!!!!")
        # print(Quest)
        # print("!!!!!!!!!!!!!!!!!!!!!!!!!")

        # print("final output",index_list,"_",cos_sim_list)

        # display seleted index
        p = Users.objects.filter(user=client_name)
        Doc1 = Document.objects.filter(selected="on", user__in=p)
        Userq = UQ.objects.filter(user=client_name)
        c = Userq.exists()

        User = askques.objects.filter(user=client_name, selected="on")
        Quest = User.exists()

        # display seleted index
        return render(
            request,
            "firstpage.html",
            {
                "Data": data,
                "showname": showname,
                "country": country,
                "industry": industry,
                "Doc": Doc1,
                "Quest": Quest,
                "c": c,
                "extrcount": extrcount,
                "Acccount": Acccount,
                "log": log,
            },
        )
    if request.method == "GET":
        client_name = request.session["client_name"]
        industry = request.session["industry"]
        country = request.session["country"]
        showname = request.session["showname"]
        try:
            extrcount = request.session["extrcount"]
            Acccount = request.session["Acccount"]
            selfirst = request.session["selfirst"]
            sellast = request.session["sellast"]
            log = request.session["log"]
        except:
            pass
        data = Question.objects.filter(country=country, industry=industry)
        p = Users.objects.filter(user=client_name)
        Doc1 = RfpSection.objects.filter(user__in=p)
        Userq = UQ.objects.filter(user=client_name)
        c = Userq.exists()

        User = askques.objects.filter(user=client_name, selected="on")
        Quest = User.exists()

        # display seleted index
        return render(
            request,
            "firstpage.html",
            {
                "Data": data,
                "showname": showname,
                "country": country,
                "industry": industry,
                "Doc": Doc1,
                "Quest": Quest,
                "c": c,
                "extrcount": extrcount,
                "Acccount": Acccount,
                "selfirst": selfirst,
                "sellast": sellast,
                "log": log,
            },
        )


# fifth page mcq options


def secondpage_view(request):
    if request.method == "GET":
        showname = request.session["showname"]
        industry = request.session["industry"]
        country = request.session["country"]
        if country == "AU":
            country2 = "AU"
        # data = request.data
        # print(data, 'dataaa')
        Query = request.GET.get("Query")
        print(Query, 'question from frontend')
        # data = Question.objects.filter(country=country, industry=industry)
        data = RfpData.objects.all()
        df_rfpdata = read_frame(data)
        print("Data Collected from Django Model.py")
        print(df_rfpdata.head())
        print(df_rfpdata.shape)
        print("End from Django Model.py")
        # df = read_frame(data)
        # df.to_excel("Data.xlsx", index=False)
        # df = pd.read_excel("Question_Pairing_Final.xlsx",
        #                    sheet_name="Tablib Dataset")
        # Time taking for passing multiple queries

        # ---------------------------------------------------Search Engine Functionality(Start)-------------------------------------
        industry1 = []
        country1 = []
        # section=  ["Relevant Experience","Project Team Structure"]
        section1 = []
        industry1.append(industry)
        country1.append(country2)

        industry_text = "_".join(industry1)
        print("Industry text:-", industry_text)
        country_text = "_".join(country1)
        print("Country text:-", country_text)
        # df = pd.read_excel("database.xlsx")
        # df = data
        # Creating an model
        model_build = Model_Buiding_approach(
            industry_text, country_text, model_name=model_name
        )

        print("Before Filtered DataFrame:- ", df_rfpdata.shape)
        # Filter the data based on Country and Industry
        df = model_build.dataframe_filter_return(
            df_rfpdata, industry1, country1, section1
        )

        print("Filtered DataFrame:- ", df.shape)
        # Training Dataset Vector
        if os.path.isfile(
            "RFP/Embedding Models/corpus_embedding"
            + "-"
            + model_name
            + "-"
            + country_text
            + "-"
            + industry_text
            + ".pt"
        ):
            print("Need not to Run Training File as Model is Present")
        else:
            print("File is not present. Initializing the training command.")
            execution_time = model_build.training_dataset_vector(
                df, embedder=embedder)
            print("Execution of Training Time:-", execution_time)

        # Input the query
        # query = "Has your organisation been involved in any business-related litigation that may affect the performance of these Services in the past five years?"
        query = Query

        index_list, cos_sim_list = model_build.query_resolved(
            query=query, df=df, embedder=embedder
        )
        print(index_list)
        index_list = [df["id"][index] for index in index_list]
        print(index_list)
        print("Index List:-", index_list)
        print("Cosine Similarity List", cos_sim_list)
        # ---------------------------------------------------Search Engine Functionality(End)-------------------------------------
        # Input the query
        # query = "Has your organisation been involved in any business-related litigation that may affect the performance of these Services in the past five years?"
        # query = Query

        request.session["query"] = query
        non = ""
        # index_list, cos_sim_list = model_build.query_resolved(
        #     query=query, df=df, embedder=embedder)
        print("INDEXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXX")
        print("INDEXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXX")
        print("INDEXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXX")
        print("INDEXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXX")
        print(index_list)
        print("INDEXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXX")
        data = RfpData.objects.filter(id__in=index_list)
        print("$$$$$$$$$$$$$$$$$$$$fileeeeeeeeeeeeeeeeeeee")
        print("$$$$$$$$$$$$$$$$$$$$fileeeeeeeeeeeeeeeeeeee")
        print("$$$$$$$$$$$$$$$$$$$$fileeeeeeeeeeeeeeeeeeee")
        print("$$$$$$$$$$$$$$$$$$$$fileeeeeeeeeeeeeeeeeeee")
        print(data)
        print("$$$$$$$$$$$$$$$$$$$$")
        # data =df.iloc[index_list]
        data0 = ""
        data1 = ""
        data2 = ""
        id0 = ""
        id1 = ""
        id2 = ""
        if len(index_list) == 1:
            data0 = data[0].document_link
            id0 = data[0].id
        elif len(index_list) == 2:
            data0 = data[0].document_link
            id0 = data[0].id
            data1 = data[1].document_link
            id1 = data[1].id
        elif len(index_list) == 3:
            data0 = data[0].document_link
            id0 = data[0].id
            data1 = data[1].document_link
            id1 = data[1].id
            data2 = data[2].document_link
            id2 = data[2].id
        else:
            non = "No matching question found"
            # data0=data[0].Questions
            # id0=data[0].id
            # data1=data[1].Questions
            # id1=data[1].id
            # data2=data[2].Questions
        # id2=data[2].id
        # l = Question.objects.get(id=170)

        # f = l.File
        data = {'non': non, 'index_list': index_list, "data0": data0, "data1": data1, "data2": data2, "id0": id0,
                "id1": id1, "id2": id2, "Query": Query, "showname": showname, "country": country, "industry": industry}
        print(data, 'final result')

        # return HttpResponse(data)
        # return render(request, 'mcq_modal.html', {'non': non,'index_list': index_list, "data0": data0, "data1": data1, "data2": data2, "id0": id0, "id1": id1, "id2": id2, "Query": Query, "showname": showname, "country": country, "industry": industry})
        return JsonResponse({"non": non, "data0": data0, "data1": data1, "data2": data2, "id0": id0, "id1": id1, "id2": id2, "Query": Query, "showname": showname, "country": country, "industry": industry})


#-------------------------------------------------------------------------------------------------------------
# Question Funcionality start
#-------------------------------------------------------------------------------------------------------------
def mcqquestionpage_view(request):
    if request.method == "GET":
        showname = request.session['showname']
        industry = request.session['industry']
        country = request.session['country']
        if country == "AU":
            country2 = "AU"
        Query = request.GET.get("Query")
        # data = Question.objects.filter(country=country, industry=industry)
        data = RfpData.objects.all()
        df_rfpdata = read_frame(data)
        print("Data Collected from Django Model.py")
        print(df_rfpdata.head())
        print(df_rfpdata.shape)
        print("End from Django Model.py")
        # df = read_frame(data)
        # df.to_excel("Data.xlsx", index=False)
        # df = pd.read_excel("Question_Pairing_Final.xlsx",
        #                    sheet_name="Tablib Dataset")
        # Time taking for passing multiple queries

        # ---------------------------------------------------Search Engine Functionality(Start)-------------------------------------
        industry1 = []
        country1 = []
        # section=  ["Relevant Experience","Project Team Structure"]
        section1 = []
        industry1.append(industry)
        country1.append(country2)

        industry_text = "_".join(industry1)
        print("Industry text:-", industry_text)
        country_text = "_".join(country1)
        print("Country text:-", country_text)
        # df = pd.read_excel("database.xlsx")
        # df = data
        # Creating an model
        model_build = Model_Buiding_approach(
            industry_text, country_text, model_name=model_name)

        print("Before Filtered DataFrame:- ", df_rfpdata.shape)
        # Filter the data based on Country and Industry
        df = model_build.dataframe_filter_return(
            df_rfpdata, industry1, country1, section1)

        print("Filtered DataFrame:- ", df.shape)
        # Training Dataset Vector
        if os.path.isfile("RFP/Embedding Models/corpus_embedding"+'-'+model_name+'-'+country_text+"-"+industry_text+".pt"):
            print("Need not to Run Training File as Model is Present")
        else:
            print("File is not present. Initializing the training command.")
            execution_time = model_build.training_dataset_vector(
                df, embedder=embedder)
            print("Execution of Training Time:-", execution_time)

        # Input the query
        # query = "Has your organisation been involved in any business-related litigation that may affect the performance of these Services in the past five years?"
        query = Query

        index_list, cos_sim_list = model_build.query_resolved(
            query=query, df=df, embedder=embedder)
        print(index_list)
        index_list = [df["id"][index] for index in index_list]
        print(index_list)
        print("Index List:-", index_list)
        print("Cosine Similarity List", cos_sim_list)
        # ---------------------------------------------------Search Engine Functionality(End)-------------------------------------
        # Input the query
        # query = "Has your organisation been involved in any business-related litigation that may affect the performance of these Services in the past five years?"
        # query = Query

        request.session['query'] = query
        non = ""
        # index_list, cos_sim_list = model_build.query_resolved(
        #     query=query, df=df, embedder=embedder)
        print("INDEXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXX")
        print("INDEXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXX")
        print("INDEXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXX")
        print("INDEXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXX")
        print(index_list)
        print("INDEXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXX")
        data = RfpData.objects.filter(id__in=index_list)
        print("$$$$$$$$$$$$$$$$$$$$fileeeeeeeeeeeeeeeeeeee")
        print("$$$$$$$$$$$$$$$$$$$$fileeeeeeeeeeeeeeeeeeee")
        print("$$$$$$$$$$$$$$$$$$$$fileeeeeeeeeeeeeeeeeeee")
        print("$$$$$$$$$$$$$$$$$$$$fileeeeeeeeeeeeeeeeeeee")
        print(data)
        print("$$$$$$$$$$$$$$$$$$$$")
        
        data0 = ""
        data1 = ""
        data2 = ""
        id0 = ""
        id1 = ""
        id2 = ""
        if len(index_list) == 1:
            data0 = data[0].document_link
            id0 = data[0].id
        elif len(index_list) == 2:
            data0 = data[0].document_link
            id0 = data[0].id
            data1 = data[1].document_link
            id1 = data[1].id
        elif len(index_list) == 3:
            data0 = data[0].document_link
            id0 = data[0].id
            data1 = data[1].document_link
            id1 = data[1].id
            data2 = data[2].document_link
            id2 = data[2].id
        else:
            non = "No matching question found"
            
        messages.success(request, "QUESTION ASKED WILL BE ADDED TO THE MAIN RFP DOCUMENT")
        return render(request, 'mcq.html', {'non': non, 'data': data, 'index_list': index_list, "data0": data0, "data1": data1, "data2": data2, "id0": id0, "id1": id1, "id2": id2, "Query": Query, "showname": showname, "country": country, "industry": industry})

#-------------------------------------------------------------------------------------------------------------
# Question Funcionality end
#-------------------------------------------------------------------------------------------------------------
 # sixth page preview


def add_ques_ans_selected_sections(request):
    # print(id, 'current page id')
    print(request.POST, 'post data')
    merge_docs = []
    node_command_string = "node doc-merger-individual.js"
    if request.POST.get('rfp_sec_id'):
        document = Document_usercopy.objects.filter(
            id=request.POST.get('rfp_sec_id')
        ).first()
        merge_docs.append(document.File.url)
        node_command_string += f' {document.File.url}'
        print(merge_docs, 'documents')
    if request.POST.get('K'):
        file_path = 'https://rfpstoragecheck.blob.core.windows.net/rfpstorage/Recommended_Documents/Documents/' + \
            request.POST.get('K')
        local_file_path = get_document(file_path)

        # merge_docs.append(file_path)
        node_command_string += f' /{local_file_path}'
    if request.POST.get('L'):
        file_path = 'https://rfpstoragecheck.blob.core.windows.net/rfpstorage/Recommended_Documents/Documents/' + \
            request.POST.get('L')
        local_file_path = get_document(file_path)

        # merge_docs.append(file_path)
        node_command_string += f' /{local_file_path}'
    if request.POST.get('M'):
        file_path = 'https://rfpstoragecheck.blob.core.windows.net/rfpstorage/Recommended_Documents/Documents/' + \
            request.POST.get('M')
        local_file_path = get_document(file_path)

        # merge_docs.append(file_path)
        node_command_string += f' /{local_file_path}'
    print(merge_docs, 'final list')
    print(node_command_string, 'final node command string')
    result = os.system(node_command_string)

    print(result, 'result of executed node file')

    if result == 0:
        subfolder = f"updated_documents/{request.session['client_name']}"
        container_id = "rfpstorage"
        updload_to_azure_blob = upload_blob_data(
            subfolder, 'output-individual.docx', container_id)
        print(updload_to_azure_blob, 'azure path')

        document.file_link = updload_to_azure_blob
        document.File.save('output-individual.docx', File(
            open('output-individual.docx', 'rb'))
        )
        document.save()
        # c = Document_usercopy.objects.update_or_create(
        #     rfp_section_id=docu.id, country=docu.country, industry=docu.industry, doc_index=docu.section_data, user=client_name, file_link=updload_to_azure_blob, matrix=matrix_value)

        # c[0].File.save(updated_doc, File(
        #     open(updated_doc, 'rb')))

    # return redirect('SelectedIndex2', id=request.POST.get('rfp_sec_id'))
    return JsonResponse({"result": "Successfully attached the file to current selected section."})


def pre_view(request):
    client_name = request.session["client_name"]
    if request.method == "GET":
        showname = request.session["showname"]
        industry = request.session["industry"]
        country = request.session["country"]
        UQL = UQ.objects.exclude(sentapproval="on").filter(user=client_name)
        return render(
            request,
            "preview.html",
            {
                "Question": Question,
                "showname": showname,
                "country": country,
                "industry": industry,
                "UQL": UQL,
            },
        )
    if request.method == "POST":
        industry = request.session["industry"]
        country = request.session["country"]
        showname = request.session["showname"]
        id0 = request.POST.get("K")
        id1 = request.POST.get("M")
        id2 = request.POST.get("L")
        p = [id0, id1, id2]
        answer = RfpData.objects.filter(id__in=p)
        query = request.session["query"]
        return render(
            request,
            "preview.html",
            {
                "answer": answer,
                "showname": showname,
                "country": country,
                "industry": industry,
                "query": query,
            },
        )


# preview page print standard pdf


def printpdf(request):
    if request.method == "GET":
        client_name = request.session["client_name"]
        industry = request.session["industry"]
        country = request.session["country"]
        # Once index selected  starts---------------------------------------------------------------------
        p = Users.objects.filter(user=client_name)

        files = Document.objects.filter(user__in=p)
        print("$$$$$$$$$$$$$$$$$$")
        print(files)
        print("$$$$$$$$$$$$$$$$$$")
        for file in files:
            file = files[0].File

        from docxcompose.composer import Composer
        from docx import Document as Document_compose

        path = "static/media/" + client_name + "/" + country + "/" + industry

        def combine_all_docx(filename_master, files_list):
            number_of_sections = len(files_list)
            master = Document_compose(filename_master)
            composer = Composer(master)
            for i in range(0, number_of_sections):
                doc_temp = Document_compose(files_list[i])
                composer.append(doc_temp)
            composer.save("./" + path + "/combined_file.docx")

        # For Example
        from docx import Document as DF

        docu = DF()

        docu.save("./" + path + "/C.docx")
        filename_master = "./" + path + "./C.docx"
        files_list = []
        for file in files:
            files_list.append(file.File)
        print("$$$$$$$$$$$$$$$$$$")
        print(files_list)
        print("$$$$$$$$$$$$$$$$$$")

        # Calling the function
        combine_all_docx(filename_master, files_list)
        # print standard starts----------------------------start-----------------------------------------

        # print standard starts--------------------------------end-------------------------------------

        # print fpdf file for question answer--------------------start-----------------------------------
        client_name = request.session["client_name"]
        industry = request.session["industry"]
        country = request.session["country"]
        import os

        path = "static/media/" + client_name + "/" + country + "/" + industry

        # download combined file--------------start------------    content_type="application/pdf///vnd.openxmlformats-officedocument.wordprocessingml.document"
        import os
        from django.conf import settings
        from django.http import HttpResponse, Http404

        client_name = request.session["client_name"]
        industry = request.session["industry"]
        country = request.session["country"]
        import os

        path = "static/media/" + client_name + "/" + country + "/" + industry
        with open("./" + path + "/combined_file.docx", "rb") as f:
            data = f.read()
        response = HttpResponse(
            data,
            content_type="vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = 'attachment; filename="RFP.docx"'

        return response
        # download combined file--------------start------------------------content_type="application/pdf--vnd.openxmlformats-officedocument.wordprocessingml.document"
        Question = request.GET.get("Question")

        return render(request, "preview.html", {"Question": Question})


# preview page erase data and again start RFP
def start_new_rfp(request):
    #     showname=request.session['showname']
    docobj1 = Document.objects.filter(
        id__in=[
            1,
            2,
            3,
            4,
            5,
            6,
            7,
            8,
            9,
            10,
            11,
            12,
            13,
            14,
            15,
            16,
            17,
            18,
            19,
            20,
            21,
            22,
            23,
        ]
    ).update(selected=" ")
    #     UQL=UQ.objects.filter(user=showname)
    #     UQL.delete()
    #     ud=SelectDropQuery.objects.all()
    #     ud.delete()
    return redirect("/RFP/index/")


# type question page print standard pdf


# upload files
def pic_upload_view(request):
    client_name = request.session["client_name"]
    country = request.session["country"]
    pic = request.FILES.get("file")
    prod = Product(user=client_name, picup=pic, clientgeo=country)
    prod.save()
    return render(request, "preview.html", {"Question": Question})


# editans_view
def editans_view(request, id):
    industry = request.session["industry"]
    country = request.session["country"]
    showname = request.session["showname"]
    id = id
    if request.method == "POST":
        pt = UQ.objects.get(pk=id)
        fm = UserQueryForm(request.POST, instance=pt)
        if fm.is_valid():
            fm.save()
    if request.method == "GET":
        pt = UQ.objects.get(pk=id)
        fm = UserQueryForm(instance=pt)

    return render(
        request,
        "editans.html",
        {
            "form": fm,
            "showname": showname,
            "country": country,
            "industry": industry,
            "id": id,
        },
    )


# pic_down_view


def pic_down_view(request):
    client_name = request.session["client_name"]
    country = request.session["country"]
    pic = request.FILES.get("file")
    First = First_Page_upload(user=client_name, picdown=pic, clientgeo=country)
    First.save()
    return render(request, "preview.html", {"Question": Question})


# file_injection_view-Direct download


def file_injection_view(request):
    client_name = request.session["client_name"]
    country = request.session["country"]
    industry = request.session["industry"]
    if country == "AU":
        country2 = "AU"
    data = RfpData.objects.all()
    df_rfpdata = read_frame(data)
    try:
        import os

        files = os.listdir("./media/First_page_upload/Drop")
        print(files)

        print(
            "---------------------------------Download Injection file using Python--------------------------------"
        )
        # from pathlib import Path
        # path1 = './media/pickupdir'
        # name_list = os.listdir(path1)
        # full_list = [os.path.join(path1, i) for i in name_list]
        # time_sorted_list = sorted(full_list, key=os.path.getmtime)

        # lst_sorted = [str1.replace('\\', '/') for str1 in time_sorted_list]
        FP = Product.objects.filter(user=client_name, clientgeo=country)
        l = len(FP)
        if l == 1:
            FP = Product.objects.get(user=client_name, clientgeo=country)
            s = FP.picup
        else:
            s = FP[l - 1].picup
        # filename1 = "C:/Users/shubhamjain35/Desktop/P Copy-801-showcase/KPM/media/pickupdir/Resolution_Life_Finance_Systems_RFP_Final.docx"
        # file = filename1.split("/")[-1]
        # print("Final File Path============", file)
        filename1 = s
        file = s
        filename1 = "media/" + str(filename1)
        print("Filename============", filename1)

        quest = Extraction()
        # filename=r"01 JHG_FMS SI Partner RFP Main Document final.docx"
        print("---" * 10, "Start " + filename1 + "---" * 10)
        df_questions, list_questions = quest.question_search(filename1)
        print("Questions in RFP " + filename1)
        print(
            "Number of questions detected from the document is :", len(
                list_questions)
        )
        print("++++++" * 20)
        num = 0
        list_questions = [text.strip() for text in list_questions]
        print(";;;;;;;;;;;;;;;;;;")
        print(list_questions)
        print(";;;;;;;;;;;;;;;;;;")
        # for question in list_questions:
        #     num += 1
        #     print(str(num)+". ", question)
        #     print("------"*20)
        #     print("1234567891234567")
        # print("---"*10, "End "+file+"---"*10)

        # ---------------------------------------------------Search Engine Functionality(Start)-------------------------------------
        industry1 = []
        country1 = []
        section1 = []
        # section=  ["Relevant Experience","Project Team Structure"]
        industry1.append(industry)
        country1.append(country2)
        print("abcdefg")
        industry_text = "_".join(industry)
        print("Industry text:-", industry_text)
        country_text = "_".join(country)
        print("Country text:-", country_text)
        # df = pd.read_excel("database.xlsx")

        # Creating an model
        model_build = Model_Buiding_approach(
            industry_text, country_text, model_name=model_name
        )

        print("Before Filtered DataFrame:- ", df_rfpdata.shape)
        # Filter the data based on Country and Industry
        df = model_build.dataframe_filter_return(
            df_rfpdata, industry1, country1, section1
        )
        print("LLLLLLLLLLLLLLLLLLLLLLLLLLLLLLLLLLLLLLLL")
        print(df.head())
        print("LLLLLLLLLLLLLLLLLLLLLLLLLLLLLLLLLLLLLLLL")
        print("Filtered DataFrame:- ", df.shape)
        # Training Dataset Vector
        if os.path.isfile(
            "RFP/Embedding Models/corpus_embedding"
            + "-"
            + model_name
            + "-"
            + country_text
            + "-"
            + industry_text
            + ".pt"
        ):
            print("Need not to Run Training File as Model is Present")
        else:
            print("File is not present. Initializing the training command.")
            execution_time = model_build.training_dataset_vector(
                df, embedder=embedder)
            print("Execution of Training Time:-", execution_time)

        # Input the query
        # query = "Has your organisation been involved in any business-related litigation that may affect the performance of these Services in the past five years?"
        # query=Query

        # index_list, cos_sim_list = model_build.query_resolved(query=query, df=df, embedder=embedder)

        # print("Index List:-", index_list)
        # print("Cosine Similarity List", cos_sim_list)
        # ---------------------------------------------------Search Engine Functionality(End) Extended-------------------------------------

        # Input the query

        # query = "Has your organisation been involved in any business-related litigation that may affect the performance of these Services in the past five years?"
        answer = []

        for q in list_questions:
            query = q
            index_list, cos_sim_list = model_build.query_resolved(
                query=query, df=df, embedder=embedder
            )
            index_list_updated = index_list
            index_list = [df["id"][index] for index in index_list]
            print("HHHHHHHHHHHHHHHHHHHHHHHHHHHHHHHHHHHHHHHHHHHHHH")
            print(index_list)
            print("HHHHHHHHHHHHHHHHHHHHHHHHHHHHHHHHHHHHHHHHHHHHHH")
            if index_list:
                answer.append(df["document_link"][index_list_updated[0]])
        print(answer)
        answ = answer[0]
        from scripts import get_document_update
        import os

        path = "temp/" + client_name + "/" + country + "/" + industry
        try:
            os.makedirs(path)

        except:
            pass

        outpath = path
        count = 1
        from docx import Document

        for ans, que in zip(answer, list_questions):
            doc = Document()
            p = doc.add_paragraph()
            runner = p.add_run("Question: " + que + "?")
            runner.bold = True
            doc.save(path + "/" + "question" + str(count) + ".docx")
            count += 1
            file_path = (
                r"https://rfpstoragecheck.blob.core.windows.net/rfpstorage/Recommended_Documents/Documents/"
                + ans
            )
            get_doc = get_document_update(file_path, outpath)
            print(get_doc, "get doc response")

        from docxcompose.composer import Composer
        from docx import Document as Document_compose
        import os

        path = "temp/" + client_name + "/" + country + "/" + industry

        try:
            os.rmdir(path)
            os.makedirs(path)

        except:
            pass

        def combine_all_docx(filename_master, files_list):
            number_of_sections = len(files_list)
            master = Document_compose(filename_master)
            composer = Composer(master)
            for i in range(0, number_of_sections):
                doc_temp = Document_compose(files_list[i])
                composer.append(doc_temp)
            composer.save("./" + path + "/combined_file.docx")

        # For Example
        from docx import Document as DF

        docu = DF()

        docu.save("./" + path + "/newdocscombine.docx")
        filename_master = "./" + path + "./newdocscombine.docx"
        file_list = []
        count = 1

        for ans1 in answer:
            file_list.append(path + "/" + "question" + str(count) + ".docx")
            file_list.append(
                "temp/" + client_name + "/" + country + "/" + industry + "/" + ans1
            )
            count += 1

        # Calling the function
        combine_all_docx(filename_master, file_list)

        # ---------------------------------------------------Search Engine Functionality(End)-------------------------------------

        # download combined file--------------start------------    content_type="application/pdf///vnd.openxmlformats-officedocument.wordprocessingml.document"
        import os
        from django.conf import settings
        from django.http import HttpResponse, Http404

        with open(path + "/combined_file.docx", "rb") as f:
            data = f.read()
        response = HttpResponse(
            data,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = 'attachment; filename="rfpdirectdoc.docx"'
        return response
    except Exception as e:
        print("Excepppppppppppppppppppptttttttttttttt")
        print(e)
        industry = request.session["industry"]
        country = request.session["country"]
        showname = request.session["showname"]
        from .models import Document

        p = Users.objects.filter(user=client_name)
        Doc1 = Document.objects.filter(selected="on", user__in=p)
        Userq = UQ.objects.filter(user=client_name)
        c = Userq.exists()

        User = askques.objects.filter(user=client_name, selected="on")
        Quest = User.exists()

        # display seleted index
        return render(
            request,
            "firstpage.html",
            {
                "showname": showname,
                "country": country,
                "industry": industry,
                "Doc": Doc1,
                "Quest": Quest,
                "c": c,
            },
        )

    # download combined file--------------start------------------------content_type="application/pdf#vnd.openxmlformats-officedocument.wordprocessingml.document"


# drop_rfp_view


def drop_rfp_view(request):
    client_name = request.session["client_name"]
    country = request.session["country"]
    if country == "AU":
        country1 = "AU"
    data = RfpData.objects.all()
    df_rfpdata = read_frame(data)

    # files = os.listdir("./media/First_page_upload/Drop")
    # print(files)
    # from pathlib import Path
    # path1 = './media/pickupdir'
    # name_list = os.listdir(path1)
    # full_list = [os.path.join(path1, i) for i in name_list]
    # time_sorted_list = sorted(full_list, key=os.path.getmtime)
    # print(time_sorted_list)
    # lst_sorted = [str1.replace('\\', '/') for str1 in time_sorted_list]

    # filename1 = "C:/Users/shubhamjain35/Desktop/P Copy-801-showcase/KPM/media/pickupdir/Resolution_Life_Finance_Systems_RFP_Final.docx"

    # filename1 = lst_sorted[-1]

    FP = Product.objects.filter(user=client_name, clientgeo=country)
    print("OOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOO")
    for p in FP:
        print(p)
    print("OOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOO")
    l = len(FP)
    if l == 1:
        FP = Product.objects.get(user=client_name, clientgeo=country)
        s = FP.picup
    else:
        s = FP[l - 1].picup
    # filename1 = "C:/Users/shubhamjain35/Desktop/P Copy-801-showcase/KPM/media/pickupdir/Resolution_Life_Finance_Systems_RFP_Final.docx"
    # file = filename1.split("/")[-1]
    # print("Final File Path============", file)
    filename1 = s
    file = s
    filename1 = "./media/" + str(filename1)
    print("$$$$$$$$$$$$$$$$$$$$fileeeeeeeeeeeeeeeeeeee")
    print("$$$$$$$$$$$$$$$$$$$$fileeeeeeeeeeeeeeeeeeee")
    print("$$$$$$$$$$$$$$$$$$$$fileeeeeeeeeeeeeeeeeeee")
    print("$$$$$$$$$$$$$$$$$$$$fileeeeeeeeeeeeeeeeeeee")
    print(filename1)
    print("$$$$$$$$$$$$$$$$$$$$")

    # tokens = ('is', 'does', 'do', 'what', 'when', 'where',
    #           'who', 'why', 'what', 'how', 'please', 'describe')

    # from docx import Document
    # import tqdm

    # document = Document(filename1)
    # #print("Questions in RFP "+filename1)
    # questions = []
    # sentences = []
    # number = 0
    # paragraphs = [para.text for para in document.paragraphs if para.text]

    # for index, paragraph in enumerate(paragraphs):
    #     text = paragraph.strip().lower()
    #     sentences.append(paragraph)
    #     if text.startswith(tokens) or text.endswith("?"):
    #         #         print(len(text))
    #         if len(text) > 100:
    #             number += 1
    #             print(str(number)+". ", paragraph)
    #             print(len(text))
    #             print()
    #             questions.append(paragraph)

    quest = Extraction()
    # filename1 = quest.input_file(filename1)
    # print("Updated File Path name:- ", filename1)
    print("---" * 10, "Start " + str(filename1) + "---" * 10)
    df_questions, list_questions = quest.question_search(filename1)
    print("Questions in RFP " + str(filename1))
    print("Number of questions detected from the document is :", len(list_questions))
    print("++++++" * 20)
    num = 0
    list_questions = [text.strip() for text in list_questions]
    for question in list_questions:
        num += 1
        print(str(num) + ". ", question)
        print("------" * 20)
    print("---" * 10, "End " + str(filename1) + "---" * 10)

    # Initializing Parameters and getting values from UI
    showname = request.session["showname"]
    industry = request.session["industry"]
    country = request.session["country"]
    # Query = request.GET.get("Query")
    # data=Question.objects.filter(country=country,industry=industry)
    # df = read_frame(data)
    # print(df.columns)
    # df.to_excel("Data.xlsx", index=False)

    # df = pd.read_excel("Question_Pairing_Final.xlsx",
    #                    sheet_name="Tablib Dataset")

    # ---------------------------------------------------Search Engine Functionality(Start)-------------------------------------
    industry1 = []
    country1 = []
    section1 = []
    # section=  ["Relevant Experience","Project Team Structure"]
    industry1.append(industry)
    country1.append(country)

    industry_text = "_".join(industry1)
    print("Industry text:-", industry_text)
    country_text = "_".join(country1)
    print("Country text:-", country_text)

    # df = pd.read_excel("database.xlsx")

    # Creating an model
    model_build = Model_Buiding_approach(
        industry_text, country_text, model_name=model_name
    )

    print("Before Filtered DataFrame:- ", df_rfpdata.shape)
    # Filter the data based on Country and Industry
    df = model_build.dataframe_filter_return(
        df_rfpdata, industry1, country1, section1)

    print("Filtered DataFrame:- ", df.shape)
    # Training Dataset Vector
    if os.path.isfile(
        "RFP/Embedding Models/corpus_embedding"
        + "-"
        + model_name
        + "-"
        + country_text
        + "-"
        + industry_text
        + ".pt"
    ):
        print("Need not to Run Training File as Model is Present")
    else:
        print("File is not present. Initializing the training command.")
        execution_time = model_build.training_dataset_vector(
            df, embedder=embedder)
        print("Execution of Training Time:-", execution_time)

    # Input the query

    # query = "Has your organisation been involved in any business-related litigation that may affect the performance of these Services in the past five years?"
    answer = []
    for q in list_questions:
        query = q
        index_list, cos_sim_list = model_build.query_resolved(
            query=query, df=df, embedder=embedder
        )
        if index_list:
            answer.append(df["document_link"][index_list[0]])
    id = 1
    for ques in list_questions:
        query = ques
        index_list, cos_sim_list = model_build.query_resolved(
            query=query, df=df, embedder=embedder
        )
        answer = [df["document_link"][index] for index in index_list]
        UserQuery = DropQuery(
            id=id,
            question=ques,
            user=showname,
            answer1=answer[0],
            answer2=answer[1],
            answer3=answer[2],
        )
        UserQuery.save()
        id = id + 1

    Quest = DropQuery.objects.all()
    messages.success(request, "QUESTION SELECTED WILL BE ADDED TO THE MAIN RFP DOCUMENT")
    return render(
        request,
        "drop_rfp.html",
        {
            "questions": Quest,
            "showname": showname,
            "country": country,
            "industry": industry,
        },
    )


def drop_rfpquest_view1(request):
    industry = request.session["industry"]
    country = request.session["country"]
    if country == "AU":
        country2 = "AU"
    print(industry, country, "POPOPOPOPOOP")
    showname = request.session["showname"]
    client_name = request.session["client_name"]
    data = RfpData.objects.all()
    df_rfpdata = read_frame(data)
    print("CHECKKKKKKKKKKKKKKKKKKKKKK")
    print(df_rfpdata["country"].unique())
    if request.method == "POST":
        Tick1 = request.POST.get("1")
        Tick2 = request.POST.get("2")
        Tick3 = request.POST.get("3")
        Tick4 = request.POST.get("4")
        Tick5 = request.POST.get("5")
        Tick6 = request.POST.get("6")
        Tick7 = request.POST.get("7")
        Tick8 = request.POST.get("8")
        Tick9 = request.POST.get("9")
        Tick10 = request.POST.get("10")
        Tick11 = request.POST.get("11")
        Tick12 = request.POST.get("12")
        Tick13 = request.POST.get("13")
        Tick14 = request.POST.get("14")
        Tick15 = request.POST.get("15")
        Tick16 = request.POST.get("16")
        Tick17 = request.POST.get("17")
        Tick18 = request.POST.get("18")
        Tick19 = request.POST.get("19")
        Tick20 = request.POST.get("20")
        Tick21 = request.POST.get("21")
        Tick22 = request.POST.get("22")
        Tick23 = request.POST.get("23")
        Tick24 = request.POST.get("24")
        Tick25 = request.POST.get("25")
        Tick26 = request.POST.get("26")
        Tick27 = request.POST.get("27")
        Tick28 = request.POST.get("28")
        Tick29 = request.POST.get("29")
        Tick30 = request.POST.get("30")
        Tick31 = request.POST.get("31")
        Tick32 = request.POST.get("32")
        Tick33 = request.POST.get("33")
        Tick34 = request.POST.get("34")
        Tick35 = request.POST.get("35")

        p = [
            Tick1,
            Tick2,
            Tick3,
            Tick4,
            Tick5,
            Tick6,
            Tick7,
            Tick8,
            Tick9,
            Tick10,
            Tick11,
            Tick12,
            Tick13,
            Tick14,
            Tick15,
            Tick16,
            Tick17,
            Tick18,
            Tick19,
            Tick20,
            Tick21,
            Tick22,
            Tick23,
            Tick24,
            Tick25,
            Tick26,
            Tick27,
            Tick28,
            Tick29,
            Tick30,
            Tick31,
            Tick32,
            Tick33,
            Tick34,
            Tick35,
        ]
        answer = DropQuery.objects.filter(id__in=p)
        # display = RfpData.objects.filter(id__in=p)
        print("***************************************")
        print("***************************************")
        print("***************************************")
        display = RfpData.objects.filter(id__in=p)
        print(answer)
        print("***************************************")
        final_list = filter(None, p)
        ids = 1
        data = Question.objects.filter(country=country, industry=industry)
        # df = read_frame(data)
        # print(df.columns)
        # df.to_excel("Data.xlsx", index=False)
        # df = pd.read_excel("Question_Pairing_Final.xlsx",
        #                    sheet_name="Tablib Dataset")

        # ---------------------------------------------------Search Engine Functionality(Start)-------------------------------------
        industry1 = []
        country1 = []

        # section=  ["Relevant Experience","Project Team Structure"]
        section1 = []
        industry1.append(industry)
        country1.append(country2)

        industry_text = "_".join(industry1)
        print("Industry text:-", industry_text)
        country_text = "_".join(country1)
        print("Country text:-", country_text)
        # df = pd.read_excel("database.xlsx")

        # Creating an model
        model_build = Model_Buiding_approach(
            industry_text, country_text, model_name=model_name
        )

        print("Before Filtered DataFrame:- ", df_rfpdata.shape)
        # Filter the data based on Country and Industry
        df = model_build.dataframe_filter_return(
            df_rfpdata, industry1, country1, section1
        )

        print("Filtered DataFrame:- ", df.shape)
        # Training Dataset Vector
        if os.path.isfile(
            "RFP/Embedding Models/corpus_embedding"
            + "-"
            + model_name
            + "-"
            + country_text
            + "-"
            + industry_text
            + ".pt"
        ):
            print("Need not to Run Training File as Model is Present")
        else:
            print("File is not present. Initializing the training command.")
            execution_time = model_build.training_dataset_vector(
                df, embedder=embedder)
            print("Execution of Training Time:-", execution_time)

        print("DF Countries unique", df["country"].unique())
        for c in answer:
            index_list, cos_sim_list = model_build.query_resolved(
                query=c.question, df=df, embedder=embedder
            )
            ans = [df["id"][index] for index in index_list]
            print("&&&&&&&&&&&")
            print(ans)
            print("&&&&&&&&&&&")
            k = 0
            answ = []
            for i in ans:
                answ.append(RfpData.objects.get(id=i))
                # answ[k] = RfpData.objects.filter(id=i)
                # k = k+1
            print("33333333333333333333333333333333333333333333333333333333333333")
            print(answ[0].document_link)
            print(answ[1].document_link)
            print(answ[2].document_link)
            print("33333333333333333333333333333333333333333333333333333333333333")
            Select = SelectDropQuery(
                id=ids,
                user=showname,
                question=c.question,
                answer1=answ[0].document_link,
                answer2=answ[1].document_link,
                answer3=answ[2].document_link,
            )
            Select.save()
            u = SelectDropQuery.objects.filter(
                user=showname).order_by("id").last()
            ids = u.id + 1

        Select = SelectDropQuery(
            id=ids, user=showname, question="Please click on finish"
        )
        Select.save()
        qt = SelectDropQuery.objects.all()
        # qt=DropQuery.objects.all()
        print(qt)
        try:
            p = SelectDropQuery.objects.filter(
                user=showname).order_by("id").first()
            # p=qt[0]
        except:
            p = "No Questions to display"

        return render(
            request,
            "drop_rfp_quest1.html",
            {"p": p, "showname": showname, "country": country, "industry": industry},
        )


# drop_rfpquest_view2


def drop_rfpquest_view2(request, id):
    industry = request.session["industry"]
    country = request.session["country"]
    if country == "AU":
        country = "AU"
    showname = request.session["showname"]
    client_name = request.session["client_name"]
    qt = (
        DropQuery.objects.filter(select1="on")
        | DropQuery.objects.filter(select2="on")
        | DropQuery.objects.filter(select3="on")
    )
    if request.method == "POST":
        p = (
            SelectDropQuery.objects.filter(id__gt=id)
            .exclude(id=id)
            .order_by("id")
            .first()
        )
        # qt=DropQuery.objects.all()filter(optionselect="on").

        if p:
            answer1 = request.POST.get("1")
            if answer1:
                try:
                    obj1 = SelectDropQuery.objects.filter(id=p.id - 1).update(
                        select1="on"
                    )
                except:
                    pass
            answer2 = request.POST.get("2")
            if answer2:
                try:
                    obj1 = SelectDropQuery.objects.filter(id=p.id - 1).update(
                        select2="on"
                    )
                except:
                    pass
            answer3 = request.POST.get("3")
            if answer3:
                try:
                    obj1 = SelectDropQuery.objects.filter(id=p.id - 1).update(
                        select3="on"
                    )
                except:
                    pass
            return render(
                request,
                "drop_rfp_quest2.html",
                {
                    "p": p,
                    "showname": showname,
                    "country": country,
                    "industry": industry,
                },
            )
        qt = SelectDropQuery.objects.all()
        return render(
            request,
            "drop_rfp_preview.html",
            {"qt": qt, "showname": showname,
                "country": country, "industry": industry},
        )


# drop_rfp_preview_view
def drop_rfp_preview_view(request):
    industry = request.session["industry"]
    country = request.session["country"]
    showname = request.session["showname"]
    client_name = request.session["client_name"]

    qt = (
        SelectDropQuery.objects.filter(user=showname).filter(select1="on")
        | SelectDropQuery.objects.filter(user=showname).filter(select2="on")
        | SelectDropQuery.objects.filter(user=showname).filter(select3="on")
    )
    # qt=SelectDropQuery.objects.all()

    return render(
        request,
        "drop_rfp_preview.html",
        {"qt": qt, "showname": showname, "country": country, "industry": industry},
    )


# drop_editans_view
def drop_editans_view(request, id):
    industry = request.session["industry"]
    country = request.session["country"]
    showname = request.session["showname"]
    if request.method == "POST":
        pt = SelectDropQuery.objects.get(pk=id)

        fm = SelectDropQueryForm(request.POST, instance=pt)
        if fm.is_valid():
            fm.save()
    if request.method == "GET":
        pt = SelectDropQuery.objects.get(pk=id)
        fm = SelectDropQueryForm(instance=pt)

    return render(
        request,
        "editans_drop.html",
        {
            "form": fm,
            "showname": showname,
            "country": country,
            "industry": industry,
            "id": id,
        },
    )


# drop_drop_print_pdf


def drop_print_pdf(request):
    industry = request.session["industry"]
    country = request.session["country"]
    showname = request.session["showname"]
    client_name = request.session["client_name"]
    import os

    path = "static/media/" + client_name + "/" + country + "/" + industry
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    ct = (
        SelectDropQuery.objects.filter(user=showname).filter(select1="on")
        | SelectDropQuery.objects.filter(user=showname).filter(select2="on")
        | SelectDropQuery.objects.filter(user=showname).filter(select3="on")
    )
    # pdf.add_font('BB','',r'C:/Windows/Fonts/javatext.TTF', uni=True)
    pdf.set_font("Times", style="U", size=15)
    pdf.multi_cell(
        0, 5, "Autodetected Question and Responses from uploaded RFP" + "\n")
    pdf.multi_cell(0, 5, "\n")
    # query = "Has your organisation been involved in any business-related litigation that may affect the performance of these Services in the past five years?"
    for i in ct:
        if i.question:
            a = i.question
            a = a.encode("latin-1", "replace").decode("latin-1")
        if i.answer1:
            b = i.answer1
            b = b.encode("latin-1", "replace").decode("latin-1")
        if i.answer2:
            c = i.answer2
            c = c.encode("latin-1", "replace").decode("latin-1")
        if i.answer3:
            d = i.answer3
            d = d.encode("latin-1", "replace").decode("latin-1")
        # pdf.set_font('BB',style ='U',size=10)
        pdf.set_font("Times", size=10)
        pdf.multi_cell(0, 5, "Question:" + a + "\n")
        pdf.multi_cell(0, 5, " ")
        pdf.set_font("Times", size=9)
        if i.select1 == "on":
            pdf.multi_cell(0, 5, "Answer:" + b)
            pdf.multi_cell(0, 5, "" + "\n" + "")
        if i.select2 == "on":
            pdf.multi_cell(0, 5, "Answer:" + c)
            pdf.multi_cell(0, 5, "" + "\n" + "")
        if i.select3 == "on":
            pdf.multi_cell(0, 5, "Answer:" + d)
            pdf.multi_cell(0, 5, "" + "\n" + "")
    import os

    path = "static/media/" + client_name + "/" + country + "/" + industry
    try:
        os.makedirs(path)
        pdf.output("./" + path + "/uploaded_file.pdf", "F")
    except:
        pdf.output("./" + path + "/uploaded_file.pdf", "F")

    import os
    from django.conf import settings
    from django.http import HttpResponse, Http404

    with open("./" + path + "/uploaded_file.pdf", "rb") as f:
        data = f.read()
    response = HttpResponse(data, content_type="application/pdf")
    response["Content-Disposition"] = 'attachment; filename="RFP_Uploaded.pdf"'

    return response
    # return render(request,"drop_rfp_preview.html",{"qt":ct})


# def user_upload_question_view(request):
#     industry=request.session['industry']
#     country=request.session['country']
#     showname=request.session['showname']
#     form=UserQuestionForm()
#     if request.method=="POST":
#         form=UserQuestionForm(request.POST)
#         if form.is_valid():
#             form.save()
#     form=UserQuestionForm()
#     return render(request,'UserQuestion.html',{'form':form,"showname":showname,"country":country,"industry":industry})


def user_upload_question_view(request, id):
    industry = request.session["industry"]
    country = request.session["country"]
    showname = request.session["showname"]
    if request.method == "GET":
        id = id
        return render(
            request,
            "UserQuestion.html",
            {"showname": showname, "country": country,
                "industry": industry, "id": id},
        )
    if request.method == "POST":
        Rfp_Id = request.POST.get("Rfp_Id")
        Content_Type = request.POST.get("content_type")
        Country = request.POST.get("country")
        Industry = request.POST.get("industry")
        Section = request.POST.get("section")
        Sub_Section = request.POST.get("sub_section")
        Question = request.POST.get("question")
        file = request.FILES.get("file")
        print(file)
        c = UserQuestion.objects.create(
            rfp_id=Rfp_Id,
            content_type=Content_Type,
            country=Country,
            industry=Industry,
            Section=Section,
            Sub_Section=Sub_Section,
            question=Question,
            Document_link=file,
        )
        c.save()
        # pt=SelectDropQuery.objects.get(pk=id)

    return render(
        request,
        "UserQuestion.html",
        {"showname": showname, "country": country, "industry": industry},
    )


def confirm_view(request, id):
    industry = request.session["industry"]
    country = request.session["country"]
    showname = request.session["showname"]
    if request.method == "GET":
        UQL = UserQuery.objects.filter(id=id)
        UQL.update(sentapproval="on", viewed="off")
        return render(
            request,
            "confirmation.html",
            {"showname": showname, "country": country,
                "industry": industry, "id": id},
        )
    if request.method == "POST":
        return render(
            request,
            "confirmation.html",
            {"showname": showname, "country": country,
                "industry": industry, "id": id},
        )


def approve_view(request):
    #     industry=request.session['industry']
    #     country=request.session['country']
    #     showname=request.session['showname']
    # if request.method == "GET":
    #     SP = UserQuery.objects.filter(sentapproval="on").filter(viewed="off")
    #     return render(request, 'approve.html', {"SP": SP})
    # if request.method == "POST":
    #     SP = UserQuery.objects.filter(sentapproval="on").filter(viewed="off")
    #     return render(request, 'approve.html', {"SP": SP})
    if request.method == "GET":
        SP = documentapproval.objects.filter(approved="No")
        loginusername = request.session["loginusername"]
        return render(
            request, "approve.html", {"SP": SP, "loginusername": loginusername}
        )
    if request.method == "POST":
        loginusername = request.session["loginusername"]
        SP = documentapproval.objects.filter(approved="No")
        return render(
            request, "approve.html", {"SP": SP, "loginusername": loginusername}
        )


def approvedocument_view(request):
    #     industry=request.session['industry']
    #     country=request.session['country']
    #     showname=request.session['showname']
    # if request.method == "GET":
    #     SP = UserQuery.objects.filter(sentapproval="on").filter(viewed="off")
    #     return render(request, 'approve.html', {"SP": SP})
    # if request.method == "POST":
    #     SP = UserQuery.objects.filter(sentapproval="on").filter(viewed="off")
    #     return render(request, 'approve.html', {"SP": SP})
    if request.method == "GET":
        SP = userstandardsection.objects.filter(approved="No")
        Yes = userstandardsection.objects.filter(approved="Yes")
        return render(request, "approvedocument.html", {"SP": SP, "Yes": Yes})
    if request.method == "POST":
        SP = userstandardsection.objects.filter(approved="No")
        Yes = userstandardsection.objects.filter(approved="Yes")
        return render(request, "approvedocument.html", {"SP": SP, "Yes": Yes})


def approved_view(request, id):
    # industry = request.session["industry"]
    # country = request.session["country"]
    # showname = request.session["showname"]
    if request.method == "GET":
        SP = userstandardsection.objects.filter(id=id)
        SP.update(approved="Yes")
        SP = userstandardsection.objects.filter(approved="No")
        Yes = userstandardsection.objects.filter(approved="Yes")
        return render(
            request,
            "approvedocument.html",
            {

                "SP": SP,
                "Yes": Yes,
            },
        )
    if request.method == "POST":
        SP = userstandardsection.objects.filter(id=id)
        SP.update(approved="Yes")
        SP = userstandardsection.objects.filter(approved="No")
        Yes = userstandardsection.objects.filter(approved="Yes")
        return render(
            request,
            "approvedocument.html",
            {

                "SP": SP,
                "Yes": Yes,
            },
        )


def disapproved_view(request, id):

    if request.method == "GET":
        SP = documentapproval.objects.filter(id=id)
        SP.update(approved="No")
        SP = documentapproval.objects.filter(approved="No")
        Yes = documentapproval.objects.filter(approved="Yes")
        return render(
            request,
            "approvedocument.html",
            {

                "SP": SP,
                "Yes": Yes,
            },
        )
    if request.method == "POST":
        SP = documentapproval.objects.filter(id=id)
        SP.update(approved="No")
        SP = documentapproval.objects.filter(approved="No")
        Yes = documentapproval.objects.filter(approved="Yes")
        return render(
            request,
            "approvedocument.html",
            {

                "SP": SP,
                "Yes": Yes,
            },
        )


def login(request):
    if request.method == "POST":
        loginusername = request.POST["username"]
        request.session["loginusername"] = loginusername
        password = request.POST["password"]
        user = auth.authenticate(username=loginusername, password=password)
        if user is not None:
            if user.groups.filter(name="Approvers").exists():
                auth.login(request, user)
                return redirect("/approveview")
            elif user.is_staff == False:
                return redirect("/userlogin")
            else:
                messages.success(request, "Incorrect Username or password")
                return redirect("/login")
        else:
            messages.success(request, "Incorrect Username or password")
            return redirect("/login")
    else:
        return render(request, "login.html")


def userlogin_view(request):
    if request.method == "GET":
        loginusername = request.session["loginusername"]
        SP = ImageUpload.objects.filter(approved="No")
        return render(
            request, "userlogin.html", {
                "SP": SP, "loginusername": loginusername}
        )
    if request.method == "POST":
        loginusername = request.session["loginusername"]
        SP = ImageUpload.objects.filter(approved="No")
        return render(
            request, "userlogin.html", {
                "SP": SP, "loginusername": loginusername}
        )

# signup


def usersignup(request):
    if request.method == "POST":
        username = request.POST["username"]
        password = request.POST["password"]
        email = request.POST["email"]
        username = username.upper()
        user = User.objects.create_user(
            username=username, password=password, is_staff=False, email=email)
        user.save()
        return redirect("/login")
    else:
        return render(request, "usersignup.html")

# signup

# approver login page


def loginapprover(request):
    if request.method == "POST":
        loginusername = request.POST["username"]
        request.session["loginusername"] = loginusername
        password = request.POST["password"]
        user = auth.authenticate(username=loginusername, password=password)
        from django.contrib.auth.models import Group
        if user is not None:
            if user.groups.filter(name="Approvers").exists():
                auth.login(request, user)
                return redirect("/approveview")

            return redirect("/loginapprover")
        else:
            messages.success(
                request, "Incorrect Username or password or check with admin")
            return redirect("/loginapprover")
    else:
        return render(request, "loginapprover.html")

# approver login page

# approver signup


def approversignup(request):
    if request.method == "POST":
        username = request.POST["username"]
        password = request.POST["password"]
        email = request.POST["email"]
        username = username.upper()
        user = User.objects.create_user(
            username=username, password=password, is_staff=False, email=email)
        user.save()
        from django.contrib.auth.models import Group
        my_group = Group.objects.get(name='Approvers_Request')
        idd = User.objects.get(username=username)
        my_group.user_set.add(idd.id)
        messages.success(
            request, "Approval account request has been sent to admin.Once approved you will be able to login")
        return redirect("/approversignup")
    else:
        return render(request, "approversignup.html")


# approver signup
def convertToBinaryData(filename):
    # Convert digital data to binary format
    with open(filename, "rb") as file:
        binaryData = file.read()
    return binaryData


def writeTofile(data, filename):
    # Convert binary data to proper format and write it on Hard Disk
    with open(filename, "wb") as file:
        file.write(data)


def Binaryfile_view(request):
    docfile = convertToBinaryData(
        "C:/Users/shubhamjain35/Desktop/RFP-Builder-New-main/SampleBinary.docx"
    )

    c = Question.objects.create(id=310, Binaryword=docfile)
    c.save()

    fetch = Question.objects.get(id=310)
    v = fetch.Binaryword

    writeTofile(
        v, "C:/Users/shubhamjain35/Desktop/RFP-Builder-New-main/SampleBinary13.docx"
    )

    return render(request, "index.html")


def chatgpt_view(request):
    # openai.api_key = os.getenv("OPENAI_API_KEY")
    industry = request.session["industry"]
    country = request.session["country"]
    showname = request.session["showname"]
    client_name = request.session["client_name"]
    ap.api_key = "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"
    if request.method == "POST":
        gtp_question = request.POST["gtp_question"]

        response = ap.Completion.create(
            model="text-davinci-003",
            # prompt="I am a highly intelligent question answering bot. If you ask me a question that is rooted in truth, I will give you the answer. If you ask me a question that is nonsense, trickery, or has no clear answer, I will respond with \"Unknown\".\n\nQ: What is human life expectancy in the United States?\nA: Human life expectancy in the United States is 78 years.\n\nQ: Who was president of the United States in 1955?\nA: Dwight D. Eisenhower was president of the United States in 1955.\n\nQ: Which party did he belong to?\nA: He belonged to the Republican Party.\n\nQ: What is the square root of banana?\nA: Unknown\n\nQ: How does a telescope work?\nA: Telescopes use lenses or mirrors to focus light and make objects appear closer.\n\nQ: Where were the 1992 Olympics held?\nA: The 1992 Olympics were held in Barcelona, Spain.\n\nQ: How many squigs are in a bonk?\nA: Unknown\n\nQ: Where is the Valley of Kings?\nA:",
            # prompt="I am a highly intelligent question answering bot. If you ask me a question that is rooted in truth, I will give you the answer. If you ask me a question that is nonsense, trickery, or has no clear answer, I will respond with \"Unknown\".\n\nQ: What is human life expectancy in the United States?\nA: ",
            # prompt="I am a highly intelligent question answering bot. If you ask me a question that is rooted in truth, I will give you the answer. If you ask me a question that is nonsense, trickery, or has no clear answer, I will respond with \"Unknown\".\n\nQ: What are the different components of workday?\nA: ",
            prompt='I am a highly intelligent question answering bot. If you ask me a question that is rooted in truth, I will give you the answer. If you ask me a question that is nonsense, trickery, or has no clear answer, I will respond with "Unknown".\n\nQ:'
            + gtp_question
            + "?\nA: ",
            temperature=0,
            max_tokens=100,
            top_p=1,
            frequency_penalty=0.0,
            presence_penalty=0.0,
            stop=["\n"],
        )
        chat = response["choices"][0]["text"].strip()
        return render(
            request,
            "chatgpt.html",
            {
                "c": chat,
                "gtp_question": gtp_question,
                "showname": showname,
                "country": country,
                "industry": industry,
            },
        )

    if request.method == "GET":
        return render(
            request,
            "chatgpt.html",
            {"showname": showname, "country": country, "industry": industry},
        )


# question = "What are the different components of workday"
# answer = openai(question)
# answer = openai(question)


def data_computation(request, i, d, standard_sections, client_name, image_url):
    subfolder = f"updated_documents/{client_name}"
    container_id = "rfpstorage"
    try:
        print(int(i), "integer")
        if i:
            # docobj1 = RfpSection.objects.filter(id=i).update(selected="on")
            docu = RfpSection.objects.get(id=int(i))
            docu.user.add(d)
            docu = RfpSection.objects.get(id=int(i))
            print(i, "i object")
            print(docu, "docobject")
            print("updated user to the rfp section")
            print()
            print("----")

            if docu.industry_matrix:
                matrix_value = docu.industry_matrix
            else:
                matrix_value = docu.country_matrix

            if docu.section_data == "Title Page":
                print("inside title")

                get_doc = get_document(image_url.cloud_link)
                print(get_doc, "get doc response")

                # update_doc = replace_word_document(client_name, '[CLIENT_NAME]', get_doc)
                if docu.document_link:
                    doc_name = docu.document_link
                else:
                    doc_name = "Title.docx"

                updated_doc = docx_template_replace(
                    get_doc, doc_name, client_name)
                # updated_doc = replace_word_doc(get_doc, client_name, request.session['showname'], request.session['client_geo'], request.session['add_line_1'],
                #                                 request.session['add_line_2'], request.session['client_zipcode'], request.session['industry'],
                #                                 request.session['kpmg_geo'], request.session['kpmg_address'], request.session['kpmg_lead'], doc_name)
                print(updated_doc, "updated version")

                updload_to_azure_blob = upload_blob_data(
                    subfolder, updated_doc, container_id
                )
                print(updload_to_azure_blob, "azure path")
                # exit(0)

                # c = Document_usercopy.objects.update_or_create(
                #     rfp_section_id=docu.id,country=docu.country, industry=docu.industry, doc_index=docu.section_data, user=client_name, file_link=image_url.cloud_link, matrix=matrix_value)
                c = Document_usercopy.objects.update_or_create(
                    rfp_section_id=docu.id,
                    country=docu.country,
                    industry=docu.industry,
                    doc_index=docu.section_data,
                    user=client_name,
                    file_link=updload_to_azure_blob,
                    matrix=matrix_value,
                )
                print(c, "c here")
                c[0].File.save(updated_doc, File(open(updated_doc, "rb")))
                # exit(0)

            else:
                if docu.document_link:
                    print("inside is document present")
                    # https://rfpstoragecheck.blob.core.windows.net/data/Healthcare/Australia/Healthcare_Australia_Executive Summary.docx
                    # file_path = f'https://rfpstoragecheck.blob.core.windows.net/data/{docu.industry}/{docu.country}/Content/{docu.document_link}'
                    file_path = f"https://rfpstoragecheck.blob.core.windows.net/rfpstorage/Section_Documents/{docu.industry}/{docu.country}/Content/{docu.document_link}"

                    get_doc = get_document(file_path)
                    print(get_doc, "get doc response")

                    # update_doc = replace_word_document(client_name, '[CLIENT_NAME]', get_doc)
                    # updated_doc = replace_word_document(client_name, '[CLIENT_NAME]', get_doc, docu.document_link)
                    updated_doc = replace_word_doc(
                        get_doc,
                        client_name,
                        request.session["showname"],
                        request.session["client_geo"],
                        request.session["add_line_1"],
                        request.session["add_line_2"],
                        request.session["client_zipcode"],
                        request.session["industry"],
                        request.session["kpmg_geo"],
                        request.session["kpmg_address"],
                        request.session["kpmg_lead"],
                        request.session["TitleforStyleSheetSelected"],
                        docu.document_link,
                    )

                    print(updated_doc, "update doc")
                    # exit(0)
                    updload_to_azure_blob = upload_blob_data(
                        subfolder, updated_doc, container_id
                    )
                    print(updload_to_azure_blob, "azure path")

                    # c = Document_usercopy.objects.update_or_create(
                    #     rfp_section_id=docu.id,country=docu.country, industry=docu.industry, doc_index=docu.section_data, user=client_name, file_link=file_path, matrix=matrix_value)
                    c = Document_usercopy.objects.update_or_create(
                        rfp_section_id=docu.id,
                        country=docu.country,
                        industry=docu.industry,
                        doc_index=docu.section_data,
                        user=client_name,
                        file_link=updload_to_azure_blob,
                        matrix=matrix_value,
                    )
                    print(c, "c here")

                    c[0].File.save(updated_doc, File(open(updated_doc, "rb")))

                    # exit(0)
                else:
                    try:
                        print("__________________**********___________")
                        print()
                        print()
                        other_document = RfpSection.objects.filter(
                            section_data=docu.section_data
                        ).exclude(document_link__in=["", None])[0]
                        if other_document:
                            file_path = f"https://rfpstoragecheck.blob.core.windows.net/rfpstorage/Section_Documents/{other_document.industry}/{other_document.country}/Content/{other_document.document_link}"

                            print(file_path, "file path to download")

                            get_doc = get_document(file_path)
                            print(get_doc, "get doc response")

                            updated_doc = replace_word_doc(
                                get_doc,
                                client_name,
                                request.session["showname"],
                                request.session["client_geo"],
                                request.session["add_line_1"],
                                request.session["add_line_2"],
                                request.session["client_zipcode"],
                                request.session["industry"],
                                request.session["kpmg_geo"],
                                request.session["kpmg_address"],
                                request.session["kpmg_lead"],
                                other_document.document_link,
                            )

                            print(updated_doc, "update doc")

                            updload_to_azure_blob = upload_blob_data(
                                subfolder, updated_doc, container_id
                            )
                            print(updload_to_azure_blob, 'azure path')

                            c = Document_usercopy.objects.update_or_create(
                                rfp_section_id=docu.id,
                                country=docu.country,
                                industry=docu.industry,
                                doc_index=docu.section_data,
                                user=client_name,
                                file_link=updload_to_azure_blob,
                                matrix=matrix_value,
                            )

                            c[0].File.save(updated_doc, File(
                                open(updated_doc, "rb")))
                    except Exception as ex:
                        print(ex, "exception")
                        file_path = "https://rfpstoragecheck.blob.core.windows.net/rfpstorage/Section_Documents/Blank_Documents.docx"

                        print(file_path, "file path to download")

                        get_doc = get_document(file_path)
                        print(get_doc, "get doc response")

                        updload_to_azure_blob = upload_blob_data(
                            subfolder, get_doc, container_id)
                        print(updload_to_azure_blob, 'azure path')

                        # c = Document_usercopy.objects.update_or_create(
                        #     rfp_section_id=docu.id,country=docu.country, industry=docu.industry, doc_index=docu.section_data, user=client_name, matrix=matrix_value)

                        c = Document_usercopy.objects.update_or_create(
                            rfp_section_id=docu.id,
                            country=docu.country,
                            industry=docu.industry,
                            doc_index=docu.section_data,
                            user=client_name,
                            file_link=updload_to_azure_blob,
                            matrix=matrix_value,
                        )

                        c[0].File.save(get_doc, File(open(get_doc, "rb")))

            if docu.country_matrix == "S":
                standard_sections.append(docu.section_data)

        else:
            try:
                docobj1 = RfpSection.objects.get(id=int(i), user=d)
                c = docobj1.user.remove(d)
            except:
                pass
    except Exception as e:
        print(e, "exception")


def SelectedIndex_view(request):
    industry = request.session["industry"]
    country = request.session["country"]
    showname = request.session["showname"]
    client_name = request.session["client_name"]
    ap.api_key = "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"
    if request.method == "GET":
        p = Users.objects.get(user=client_name)
        Doccopy = Document.objects.filter(user=p)
        Document_usercopy.objects.filter(user=client_name).delete()
        for docu in Doccopy:
            c = Document_usercopy.objects.create(
                country=docu.country,
                industry=docu.industry,
                doc_index=docu.doc_index,
                user=client_name,
                File=docu.File,
            )

        show = Document_usercopy.objects.filter(user=client_name).first()
        print("SelectedIndex_view_Get")
        print(show.id)
        print("SelectedIndex_view_Get")
        return render(
            request,
            "SelectedIndex.html",
            {
                "showname": showname,
                "country": country,
                "industry": industry,
                "show": show,
            },
        )
    if request.method == "POST":
        # CHECKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKK
        print("checkkkkkkkkkkk im here request post", request.POST)
        client_name = request.session["client_name"]
        industry = request.session["industry"]
        country = request.session["country"]
        radio = request.POST.get("flexRadioDefault")
        Quest = request.POST.get("Quest")

        request_post_list = dict(request.POST).keys()
        print(request_post_list, "post list")

        # request_post_list = sorted(request_post_list, key=int)

        print(request.POST, "post dataaaa")
        print(request_post_list, "post list")

        check_user = Users.objects.filter(user=client_name)

        print(check_user, "check user")
        # exit(0)
        if check_user:
            d = check_user[0]

        else:
            d = Users(user=client_name)
            d.save()
            print(d, "user saved")

        if radio:
            # u = Users.objects.get(user=client_name)

            Imagee = Image.objects.get(id=radio)

            l = Image.objects.filter(id=radio)
            t = Image.objects.filter(id=radio).update(selected="on")
            image_url = Image.objects.get(id=radio)

            print(l, "lllll")
            #print(image_url.image.url, "urlll of image")
            if client_name in l:
                pass
            else:
                try:
                    c = Imagee.user.add(d)
                except:
                    pass

                l = Imagee.user.all()

        if Quest:
            print("inside ques")
            ck = askques.objects.filter(user=client_name)
            print(ck, "ckkkk")
            if not ck:
                askquestion = askques(user=client_name, selected="on")
                askquestion.save()
            if ck:
                askquestion = askques.objects.filter(user=client_name).update(
                    selected="on"
                )
        else:
            check = askques.objects.filter(user=client_name)
            if check:
                askque = askques.objects.filter(
                    user=client_name).update(selected=" ")
            else:
                print(client_name, "inside else")
                askque = askques(user=client_name, selected=" ")
                askque.save()

        if request.POST.get("gtp_question"):
            gtp_question = request.POST["gtp_question"]
            selfirst = "DEF"
            request.session["selfirst"] = selfirst
            response = ap.Completion.create(
                model="text-davinci-003",
                # prompt="I am a highly intelligent question answering bot. If you ask me a question that is rooted in truth, I will give you the answer. If you ask me a question that is nonsense, trickery, or has no clear answer, I will respond with \"Unknown\".\n\nQ: What is human life expectancy in the United States?\nA: Human life expectancy in the United States is 78 years.\n\nQ: Who was president of the United States in 1955?\nA: Dwight D. Eisenhower was president of the United States in 1955.\n\nQ: Which party did he belong to?\nA: He belonged to the Republican Party.\n\nQ: What is the square root of banana?\nA: Unknown\n\nQ: How does a telescope work?\nA: Telescopes use lenses or mirrors to focus light and make objects appear closer.\n\nQ: Where were the 1992 Olympics held?\nA: The 1992 Olympics were held in Barcelona, Spain.\n\nQ: How many squigs are in a bonk?\nA: Unknown\n\nQ: Where is the Valley of Kings?\nA:",
                # prompt="I am a highly intelligent question answering bot. If you ask me a question that is rooted in truth, I will give you the answer. If you ask me a question that is nonsense, trickery, or has no clear answer, I will respond with \"Unknown\".\n\nQ: What is human life expectancy in the United States?\nA: ",
                # prompt="I am a highly intelligent question answering bot. If you ask me a question that is rooted in truth, I will give you the answer. If you ask me a question that is nonsense, trickery, or has no clear answer, I will respond with \"Unknown\".\n\nQ: What are the different components of workday?\nA: ",
                prompt='I am a highly intelligent question answering bot. If you ask me a question that is rooted in truth, I will give you the answer. If you ask me a question that is nonsense, trickery, or has no clear answer, I will respond with "Unknown".\n\nQ:'
                + gtp_question
                + "?\nA: ",
                temperature=0,
                max_tokens=100,
                top_p=1,
                frequency_penalty=0.0,
                presence_penalty=0.0,
                stop=["\n"],
            )
            chat = response["choices"][0]["text"].strip()
            print(chat, "response from chat gpt--------")
            # show = Document_usercopy.objects.filter(user=client_name).order_by('rfp_section__order').first()
            show = Document_usercopy.objects.get(doc_index="Executive Summary")

            print(show, "all user data")

            print("SelectedIndex_view_POST")

            print("SelectedIndex_view_POST")
            print(show, "show data")

            return render(
                request,
                "SelectedIndex.html",
                {
                    "c": chat,
                    "standard_sections": standard_sections,
                    "gtp_question": gtp_question,
                    "showname": showname,
                    "country": country,
                    "industry": industry,
                    "show": show,
                },
            )

        standard_sections = []
        # subfolder = f"updated_documents/{client_name}"
        # container_id = "rfpstorage"
        thread_list = []

        for i in range(0, len(request_post_list)):
            print(i, "ii - check")
            try:
                temp_var = f"t{i}"
                temp_var = threading.Thread(
                    target=data_computation,
                    args=(
                        request,
                        list(request_post_list)[i],
                        d,
                        standard_sections,
                        client_name,
                        image_url,
                    ),
                )
                temp_var.start()
                thread_list.append(temp_var)
                # compute_data = data_computation(request, list(request_post_list)[i], d, standard_sections, client_name, image_url)
            except Exception as e:
                print(e)
            # run_thread = data_computation()

        print(thread_list, "thread_list")
        for i in thread_list:
            i.join()

        # checkkkkkkkkkkkkkkkkkkkkkkkkkkkkkkkkkkkkkkkkkkkk

        # p = Users.objects.get(user=client_name)
        # Doccopy = RfpSection.objects.filter(user=p).order_by('order')

        # print(Doccopy, 'doccopyy')

        # Document_usercopy.objects.filter(user=client_name).delete()
        # standard_sections = []
        # for docu in Doccopy:
        #     if docu.industry_matrix:
        #         matrix_value = docu.industry_matrix
        #     else:
        #         matrix_value = docu.country_matrix

        #     if docu.document_link:
        #         file_path = f'https://rfpstoragecheck.blob.core.windows.net/data/{docu.industry}/{docu.country}/Content/{docu.document_link}'

        #         c = Document_usercopy.objects.update_or_create(
        #             country=docu.country, industry=docu.industry, doc_index=docu.section_data, user=client_name, file_link=file_path, matrix=matrix_value)
        #     else:
        #         c = Document_usercopy.objects.update_or_create(
        #             country=docu.country, industry=docu.industry, doc_index=docu.section_data, user=client_name, matrix=matrix_value)

        #     if docu.country_matrix == 'S':
        #         standard_sections.append(docu.section_data)

        #     print(c, 'created user copy')

        # try:
        #     gtp_question = request.POST['gtp_question']

        #     response = ap.Completion.create(
        #         model="text-davinci-003",
        #         # prompt="I am a highly intelligent question answering bot. If you ask me a question that is rooted in truth, I will give you the answer. If you ask me a question that is nonsense, trickery, or has no clear answer, I will respond with \"Unknown\".\n\nQ: What is human life expectancy in the United States?\nA: Human life expectancy in the United States is 78 years.\n\nQ: Who was president of the United States in 1955?\nA: Dwight D. Eisenhower was president of the United States in 1955.\n\nQ: Which party did he belong to?\nA: He belonged to the Republican Party.\n\nQ: What is the square root of banana?\nA: Unknown\n\nQ: How does a telescope work?\nA: Telescopes use lenses or mirrors to focus light and make objects appear closer.\n\nQ: Where were the 1992 Olympics held?\nA: The 1992 Olympics were held in Barcelona, Spain.\n\nQ: How many squigs are in a bonk?\nA: Unknown\n\nQ: Where is the Valley of Kings?\nA:",
        #         # prompt="I am a highly intelligent question answering bot. If you ask me a question that is rooted in truth, I will give you the answer. If you ask me a question that is nonsense, trickery, or has no clear answer, I will respond with \"Unknown\".\n\nQ: What is human life expectancy in the United States?\nA: ",
        #         # prompt="I am a highly intelligent question answering bot. If you ask me a question that is rooted in truth, I will give you the answer. If you ask me a question that is nonsense, trickery, or has no clear answer, I will respond with \"Unknown\".\n\nQ: What are the different components of workday?\nA: ",
        #         prompt="I am a highly intelligent question answering bot. If you ask me a question that is rooted in truth, I will give you the answer. If you ask me a question that is nonsense, trickery, or has no clear answer, I will respond with \"Unknown\".\n\nQ:"+gtp_question+"?\nA: ",
        #         temperature=0,
        #         max_tokens=100,
        #         top_p=1,
        #         frequency_penalty=0.0,
        #         presence_penalty=0.0,
        #         stop=["\n"]
        #     )
        #     chat = response['choices'][0]['text'].strip()
        #     print(chat, 'response from chat gpt--------')
        #     # show = Document_usercopy.objects.filter(user=client_name).order_by('rfp_section__order').first()
        #     show = Document_usercopy.objects.get(doc_index='Executive Summary')

        #     print(show, 'all user data')

        #     print("SelectedIndex_view_POST")

        #     print("SelectedIndex_view_POST")
        #     print(show, 'show data')

        #     return render(request, 'SelectedIndex.html', {'c': chat, "standard_sections": standard_sections, "gtp_question": gtp_question, "showname": showname, "country": country, "industry": industry, "show": show})
        # except:
        #     show = Document_usercopy.objects.filter(user=client_name).order_by('rfp_section__order').first()
        #     print("SelectedIndex_view_POST_except")

        #     print("SelectedIndex_view_POST_except")
        #     # print(show, vars(show), showname, '-----------')
        #     return render(request, 'SelectedIndex.html', {"showname": showname, "standard_sections": standard_sections, "country": country, "industry": industry, "show": show})
        show = (
            Document_usercopy.objects.filter(user=client_name)
            .order_by("rfp_section__order")
            .first()
        )
        print("SelectedIndex_view_POST_except")

        print("SelectedIndex_view_POST_except")
        # print(show, vars(show), showname, '-----------')
        return render(
            request,
            "SelectedIndex.html",
            {
                "showname": showname,
                "standard_sections": standard_sections,
                "country": country,
                "industry": industry,
                "show": show,
            },
        )


def SelectedIndex2_view(request, id):
    industry = request.session["industry"]
    country = request.session["country"]
    if country == "AU":
        country1 = "AU"
    showname = request.session["showname"]
    client_name = request.session["client_name"]
    if request.method == "GET":
        p = Users.objects.get(user=client_name)
        Doccopy = Document.objects.filter(user=p)

        for docu in Doccopy:
            c = Document_usercopy.objects.create(
                country=docu.country,
                industry=docu.industry,
                doc_index=docu.doc_index,
                File=docu.File,
            )

        return render(
            request,
            "SelectedIndex2.html",
            {
                "showname": showname,
                "country": country,
                "industry": industry,
                "Doccopy": Doccopy,
            },
        )
    if request.method == "POST":
        Queryyy = request.POST.get("Queryyy")
        print("SelectedIndex2_view_POST")
        print(Queryyy)
        print("SelectedIndex2_view_POST")
        editable_sections = ["I", "M", "E"]
        try:
            extraimg = request.POST.getlist("sectionextraimage")
            print(extraimg, "extraaaaaaaa")
            print(extraimg, "extraimg--------------------")
            if extraimg:
                extraselected = SectionExtraImage.objects.filter(
                    id__in=extraimg)
                delextraselected = ExtraImage.objects.exclude(id__in=extraimg)
                user = Users.objects.get(user=client_name)
                for e in extraselected:
                    c = e.user.add(user)
                for d in delextraselected:
                    c = d.user.remove(user)
                create_doc_of_images = create_images_doc(extraselected)
                user_copy = Document_usercopy.objects.get(id=id)
                save_image_doc = ImageDocumentUsercopy.objects.update_or_create(
                    doc_user_copy=user_copy
                )
                save_image_doc[0].image_doc.save(
                    create_doc_of_images, File(
                        open(create_doc_of_images, "rb"))
                )
                print("successfully added image document to section")
                # extra = ExtraImage.objects.filter(user=user)
                # extrano = ExtraImage.objects.exclude(user=user)
        except:
            pass

        pivot = Document_usercopy.objects.get(id=id)
        show2 = (
            Document_usercopy.objects.filter(user=client_name)
            .filter(rfp_section__order__gt=pivot.rfp_section.order)
            .exclude(id=pivot.id)
            .order_by("rfp_section__order")
            .first()
        )
        if not show2:
            # show2 = Document_usercopy.objects.filter(user=client_name).filter(
            #     id__gt=id).exclude(id=id).order_by('id').first()
            # print("SelectedIndex2_view_POST_not_Query")
            # print(show2.id)
            # print("SelectedIndex2_view_POST_not_Query")

            rfp_user_sections = Document_usercopy.objects.filter(
                user=client_name
            ).exclude(doc_index="Title Page")
            return render(
                request,
                "SelectedIndexlastPage.html",
                {
                    "client_name": client_name,
                    "showname": showname,
                    "country": country,
                    "industry": industry,
                    "filenames": rfp_user_sections,
                    "editable_sections": editable_sections,
                },
            )
        if show2:
            showname = request.session["showname"]
            industry = request.session["industry"]
            country = request.session["country"]
            client_name = request.session["client_name"]
            Queryyy = request.POST.get("Queryyy")
            data = Question.objects.filter(country=country, industry=industry)
            standard_sections = Document_usercopy.objects.filter(
                matrix="S"
            ).values_list("doc_index", flat=True)
            show2 = (
                Document_usercopy.objects.filter(user=client_name)
                .filter(rfp_section__order__gt=pivot.rfp_section.order)
                .exclude(id=pivot.id)
                .order_by("rfp_section__order")
                .first()
            )
            print(show2, "show2")
            IMGSEC = SectionExtraImage.objects.filter(
                country=country, industry=industry, section_data=show2.doc_index
            )
            for i in IMGSEC:
                print("JJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJ")
                print(i.id)
                print("JJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJJ")
            Noimage = "A"
            if not IMGSEC:
                Noimage = "No IMAGES FOUND"

            print(IMGSEC, "imgsec")
            if show2.doc_index == "Executive Summary":
                return render(
                    request,
                    "SelectedIndex.html",
                    {
                        "showname": showname,
                        "standard_sections": standard_sections,
                        "country": country,
                        "industry": industry,
                        "show": show2,
                        "IMGSEC": IMGSEC,
                        "Noimage": Noimage,
                    },
                )

            return render(
                request,
                "SelectedIndex2.html",
                {
                    "data": data,
                    "editable_sections": editable_sections,
                    "showname": showname,
                    "country": country,
                    "industry": industry,
                    "show2": show2,
                    "IMGSEC": IMGSEC,
                    "Noimage": Noimage,
                },
            )


def SelectedIndexlastPage_view(request):
    industry = request.session["industry"]
    country = request.session["country"]
    showname = request.session["showname"]
    client_name = request.session["client_name"]
    sellast = "DEF"
    request.session["sellast"] = sellast
    editable_sections = ["I", "M", "E"]
    rfp_user_sections = Document_usercopy.objects.filter(user=client_name).exclude(
        doc_index="Title Page"
    )
    return render(
        request,
        "SelectedIndexlastPage.html",
        {
            "client_name": client_name,
            "showname": showname,
            "country": country,
            "industry": industry,
            "filenames": rfp_user_sections,
            "editable_sections": editable_sections,
        },
    )
    # return render(request, 'SelectedIndexlastPage.html', {"showname": showname, "country": country, "industry": industry})


def Onscreenmcq_view(request, id):
    if request.method == "POST":
        showname = request.session["showname"]
        industry = request.session["industry"]
        country = request.session["country"]
        if country == "AU":
            country = "AU"
        client_name = request.session["client_name"]
        Queryyy = request.POST.get("Queryyy")
        data = RfpData.objects.all()
        df_rfpdata = read_frame(data)
        # data = Question.objects.filter(country=country, industry=industry)
        # print("$$$$$$$$$$$$$$$$$$$$")
        # print(id)
        # print("$$$$$$$$$$$$$$$$$$$$")
        # df = read_frame(data)
        # df.to_excel("Data.xlsx", index=False)
        # df = pd.read_excel("Question_Pairing_Final.xlsx",
        #                    sheet_name="Tablib Dataset")

        # ---------------------------------------------------Search Engine Functionality(Start)-------------------------------------
        industry1 = []
        country1 = []
        # section=  ["Relevant Experience","Project Team Structure"]
        section1 = []
        industry1.append(industry)
        country1.append(country)

        industry_text = "_".join(industry1)
        print("Industry text:-", industry_text)
        country_text = "_".join(country1)
        print("Country text:-", country_text)
        # df = pd.read_excel("database.xlsx")

        # Creating an model
        model_build = Model_Buiding_approach(
            industry_text, country_text, model_name=model_name
        )

        print("Before Filtered DataFrame:- ", df_rfpdata.shape)
        # Filter the data based on Country and Industry
        df = model_build.dataframe_filter_return(
            df_rfpdata, industry, country, section)

        print("Filtered DataFrame:- ", df.shape)
        # Training Dataset Vector
        if os.path.isfile(
            "RFP/Embedding Models/corpus_embedding"
            + "-"
            + model_name
            + "-"
            + country_text
            + "-"
            + industry_text
            + ".pt"
        ):
            print("Need not to Run Training File as Model is Present")
        else:
            print("File is not present. Initializing the training command.")
            execution_time = model_build.training_dataset_vector(
                df, embedder=embedder)
            print("Execution of Training Time:-", execution_time)

        # Input the query
        # query = "Has your organisation been involved in any business-related litigation that may affect the performance of these Services in the past five years?"
        queryyy = Queryyy

        request.session["queryyy"] = queryyy
        non = ""
        index_list, cos_sim_list = model_build.query_resolved(
            query=queryyy, df=df, embedder=embedder
        )

        data = Question.objects.filter(id__in=index_list)
        data0 = ""
        data1 = ""
        data2 = ""
        id0 = ""
        id1 = ""
        id2 = ""
        if len(index_list) == 1:
            data0 = data[0].Model_Image_Link_BLOB
            id0 = data[0].id
        elif len(index_list) == 2:
            data0 = data[0].Model_Image_Link_BLOB
            id0 = data[0].id
            data1 = data[1].Model_Image_Link_BLOB
            id1 = data[1].id
        elif len(index_list) == 3:
            data0 = data[0].Model_Image_Link_BLOB
            id0 = data[0].id
            data1 = data[1].Model_Image_Link_BLOB
            id1 = data[1].id
            data2 = data[2].Model_Image_Link_BLOB
            id2 = data[2].id
        else:
            non = "No matching question found"
            # data0=data[0].Questions
            # id0=data[0].id
            # data1=data[1].Questions
            # id1=data[1].id
            # data2=data[2].Questions
        # id2=data[2].id
        # l = Question.objects.get(id=170)

        # f = l.File
        show2 = Document_usercopy.objects.filter(id=id)
        return render(
            request,
            "SelectedIndex2.html",
            {
                "non": non,
                "data": data,
                "data0": data0,
                "data1": data1,
                "data2": data2,
                "id0": id0,
                "id1": id1,
                "id2": id2,
                "Query": Queryyy,
                "showname": showname,
                "country": country,
                "industry": industry,
                "show2": show2,
            },
        )


# documentapproval files


def download_merged_doc(request):
    filename = "whatever_in_absolute_path__or_not.pdf"
    content = FileWrapper(filename)
    response = HttpResponse(content, content_type="application/pdf")
    response["Content-Length"] = os.path.getsize(filename)
    response["Content-Disposition"] = (
        "attachment; filename=%s" % "whatever_name_will_appear_in_download.pdf"
    )
    return response


def download(request, path):
    file_path = os.path.join(settings.MEDIA_ROOT, path)
    if os.path.exists(file_path):
        with open(file_path, "rb") as fh:
            response = HttpResponse(
                fh.read(), content_type="application/vnd.ms-excel")
            response["Content-Disposition"] = "inline; filename=" + os.path.basename(
                file_path
            )
            return response
    raise Http404


def documentapproval_view(request):
    print("im here listening to dropzone")

    print(request.POST, "request post")
    print(request.FILES.get("file"), "file")
    # print(request.FILES, 'files')
    # exit(0)
    _file = request.FILES.get("file")
    print(_file.name, "name of the file")

    client_name = request.session["client_name"]
    country = request.session["country"]
    file_data = request.FILES.get("file")

    # documents = Document_usercopy.objects.filter(user=client_name)
    print(client_name, "client name")
    file_name = _file.name
    print(file_name, "file name to deformat")

    file_name = file_name.split(f"{client_name}_")
    print(file_name, "after splitting")
    # exit(0)
    find_document = Document_usercopy.objects.filter(
        user=client_name, rfp_section__document_file_name=file_name[1]
    )[0]
    print(find_document, "found")
    # find_document.File.save(_file, File(open(_file,'rb')))
    find_document.File = _file
    find_document.save()

    # exit(0)

    # replace_data_doc = replace_word_document(client_name, file_data)
    # prod = documentapproval(
    #     user=client_name, documentapproval=fileapp, clientgeo=country)
    # prod.save()
    return 'successfully uploaded'

import subprocess


def generate_rfp_document(request):
    print("im here inside the rfp document")
    client_name = request.session["client_name"]
    country = request.session["country"]
    industry = request.session["industry"]

    all_documents = (
        Document_usercopy.objects.filter(
            country=country, industry=industry, user=client_name
        )
        .exclude(File__in=["", None])
        .order_by("rfp_section__order")
    )

    print(all_documents, "all the documents")

    file_list = []
    node_command_string = "node doc-merger.js"
    for i in all_documents:
        print(i.File.url, i.id, 'urlllll')
        if i.file_link != 'https://rfpstoragecheck.blob.core.windows.net/rfpstorage/Section_Documents/Blank_Documents.docx':
            file_list.append(i.File.url)
        node_command_string += f' {i.File.url}'
        try:
            extra_image_file = ImageDocumentUsercopy.objects.filter(
                doc_user_copy_id=i.id
            )[0]
            print(extra_image_file, "imageeee docccc")
            if extra_image_file:
                file_list.append(extra_image_file.image_doc.url)
                node_command_string += f" {extra_image_file.image_doc.url}"
        except Exception as e:
            print(e, "exception at adding image document")

    print(file_list, "file list")
    print(node_command_string, "node command")

    # exit(0)

    # result = subprocess.run(["node", "doc-merget.js", file_list[0], file_list[1]], capture_output=True, text=True, check=True)

    # result = os.system("node doc-merger.js /media/files/media/KPMG_New_Testing_Node_001/KPMG_New_Testing_Node_001_Title.docx /media/files/media/KPMG_New_Testing_Node_001/KPMG_New_Testing_Node_001_Healthcare_AU_Executive_Summary.docx")
    result = os.system(node_command_string)
    print(result, "result of executed node file")
    # add_header_footer = write_header_footer('output-node-merger-v4.docx')
    # print(add_header_footer, 'adding header and footer to the final document')
    # exit(0)
    # combine = merge_files(file_list)
    # remove_aspose_wording = replace_aspose_word(combine, client_name)
    # print(remove_aspose_wording, 'remove aspose')
    # exit(0)
    create_udpate_user_rfp = RfpDocuments.objects.update_or_create(
        industry=industry, country=country, user=client_name
    )
    # create_udpate_user_rfp[0].rfp_file.save(
    #     remove_aspose_wording, File(open(remove_aspose_wording, 'rb')))
    create_udpate_user_rfp[0].rfp_file.save(
        "rfp_final_v4.docx", File(open("output-node-merger-v4.docx", "rb"))
    )

    file_path = create_udpate_user_rfp[0].rfp_file.url
    directory = os.getcwd()
    # print(directory, 'directoryhy')
    string_path = directory + file_path

    print(file_path, "file path opened with joinnnnn")
    if os.path.exists(string_path):
        with open(string_path, "rb") as fh:
            response = HttpResponse(
                fh.read(), content_type="application/vnd.ms-excel")
            response["Content-Disposition"] = "inline; filename=" + os.path.basename(
                file_path
            )
            return response
    raise Http404
    # return render(request, 'SelectedIndexlastPage.html')


def ExtraImage_view(request):
    showname = request.session["showname"]
    industry = request.session["industry"]
    country = request.session["country"]
    client_name = request.session["client_name"]
    if request.method == "POST":
        Extraimgsearch = request.POST.get("Extraimgsearch")
        request.session["Extraimgsearch"] = Extraimgsearch
        print("Extraimgsearch", Extraimgsearch)
        extraimg = request.POST.getlist("extraimage")
        checkon = ExtraImage.objects.filter(Industry=Extraimgsearch)
        checkonid = []
        for i in checkon:
            checkonid.append(i.id)
        print("checkonid", checkonid)
        print("extraimg", extraimg)
        extraimg = [int(x) for x in extraimg]
        print("extraimggggg", extraimg)
        on = list(set(checkonid).intersection(extraimg))
        print("on", on)
        if len(on):
            off = list(set(checkonid) - set(extraimg))
            print("off", off)
            extraselected = ExtraImage.objects.filter(id__in=on)
            delextraselected = ExtraImage.objects.filter(id__in=off)
            user = Users.objects.get(user=client_name)
            for e in extraselected:
                c = e.user.add(user)
            for d in delextraselected:
                c = d.user.remove(user)

        user = Users.objects.get(user=client_name)
        Extraimgsearch = request.session["Extraimgsearch"]
        extra = ExtraImage.objects.filter(Industry=Extraimgsearch, user=user)
        extrano = ExtraImage.objects.filter(
            Industry=Extraimgsearch).exclude(user=user)
        
        return render(
            request,
            "extraimage.html",
            {
                "showname": showname,
                "country": country,
                "industry": industry,
                "extra": extra,
                "extrano": extrano,
            },
        )
    if request.method == "GET":
        extrcount = "DEF"
        request.session["extrcount"] = extrcount
        user = Users.objects.get(user=client_name)
        extra = ExtraImage.objects.filter(user=user)
        extrano = ExtraImage.objects.exclude(user=user)
        indus = ExtraImage.objects.all()
        print("FFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFF")
        print(indus)
        for i in indus:
            print(i.Industry)
        print("FFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFF")
        messages.success(request, "SELECTED IMAGES WILL BE ADDED TO THE MAIN RFP DOCUMENT")
        return render(
            request,
            "extraimage.html",
            {
                "showname": showname,
                "country": country,
                "industry": industry,
                "extra": extra,
                "extrano": extrano,
            },
        )


def UploadExtraImage_view(request):
    showname = request.session["showname"]
    industry = request.session["industry"]
    country = request.session["country"]
    client_name = request.session["client_name"]
    if request.method == "POST":
        return render(
            request,
            "UploadExtraImage.html",
            {"showname": showname, "country": country, "industry": industry},
        )
    if request.method == "GET":
        return render(
            request,
            "UploadExtraImage.html",
            {"showname": showname, "country": country, "industry": industry},
        )


def Uploadassumprisk_view(request):
    showname = request.session["showname"]
    industry = request.session["industry"]
    country = request.session["country"]
    client_name = request.session["client_name"]
    if request.method == "POST":
        return render(
            request,
            "Uploadassumprisk.html",
            {"showname": showname, "country": country, "industry": industry},
        )
    if request.method == "GET":
        return render(
            request,
            "Uploadassumprisk.html",
            {"showname": showname, "country": country, "industry": industry},
        )


# upload Image files
def image_upload_view(request):
    username = request.session["username"]
    client_name = request.session["client_name"]
    loginusername = request.session["loginusername"]
    country = request.session["country"]
    pic = request.FILES.get("file")
    prod = ImageUpload(user=loginusername, picup=pic, clientgeo=country)
    prod.save()
    return render(request, "UploadExtraImage.html")


# upload Assumption Risk files
def Assumption_upload_view(request):
    username = request.session["username"]
    country = request.POST["countries"]
    industry = request.POST["industry"]
    section = request.POST["Section"]
    loginusername = request.session["loginusername"]
    description = request.POST["description"]
    assump = userriskandassumption.objects.create(
        user=username,
        country=country,
        industry=industry,
        section=section,
        description=description,
    )
    assump.save()
    return render(request, "Uploadassumprisk.html")


def AssuptionAndRisk_view(request):
    showname = request.session["showname"]
    industry = request.session["industry"]
    country = request.session["country"]
    client_name = request.session["client_name"]
    # Assuption_And_Risk = AssuptionAndRisk.objects.filter(
    #     Topic="Assuption_And_Risk")
    # Key_consideration_and_risk = AssuptionAndRisk.objects.filter(
    #     Topic="Key_consideration_and_risk")

    if request.method == "POST":
        check = request.POST.getlist("check")

        print("OOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOO")
        print("check", check)

        print("OOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOO")
        extraselected = AssuptionAndRisk.objects.filter(id__in=check)
        print("OOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOO")
        print("extraselected", extraselected)

        print("OOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOOO")
        delextraselected = AssuptionAndRisk.objects.exclude(id__in=check)

        user = Users.objects.get(user=client_name)
        for e in extraselected:
            c = e.user.add(user)
        for d in delextraselected:
            c = d.user.remove(user)
        Generalcheck = AssuptionAndRisk.objects.filter(
            category="General", user=user,country=country)

        Generalnotcheck = AssuptionAndRisk.objects.filter(category="General",country=country ).exclude(
            user=user
        )

        Resourcescheck = AssuptionAndRisk.objects.filter(
            category="Resources", user=user,country=country
        )

        Resourcesnotcheck = AssuptionAndRisk.objects.filter(
            category="Resources",country=country
        ).exclude(user=user)

        Workdaycheck = AssuptionAndRisk.objects.filter(
            category="Workday", user=user,country=country)

        Workdaynotcheck = AssuptionAndRisk.objects.filter(category="Workday",country=country).exclude(
            user=user
        )

        Softwarecheck = AssuptionAndRisk.objects.filter(
            category="Software", user=user,country=country)

        Softwarenotcheck = AssuptionAndRisk.objects.filter(category="Software",country=country).exclude(
            user=user
        )

        Integrationcheck = AssuptionAndRisk.objects.filter(
            category="Integration", user=user,country=country
        )

        Integrationnotcheck = AssuptionAndRisk.objects.filter(
            category="Integration",country=country
        ).exclude(user=user)

        Datacheck = AssuptionAndRisk.objects.filter(
            category="Data Migration", user=user,country=country
        )

        Datanotcheck = AssuptionAndRisk.objects.filter(
            category="Data Migration",country=country
        ).exclude(user=user)

        Testingcheck = AssuptionAndRisk.objects.filter(
            category="Testing", user=user,country=country)

        Testingnotcheck = AssuptionAndRisk.objects.filter(category="Testing",country=country).exclude(
            user=user
        )

        Changecheck = AssuptionAndRisk.objects.filter(
            category="Change Management", user=user,country=country
        )

        Changenotcheck = AssuptionAndRisk.objects.filter(
            category="Change Management",country=country
        ).exclude(user=user)

        Deploymentcheck = AssuptionAndRisk.objects.filter(
            category="Deployment and Support", user=user,country=country
        )

        Deploymentnotcheck = AssuptionAndRisk.objects.filter(
            category="Deployment and Support",country=country
        ).exclude(user=user)

        Covidcheck = AssuptionAndRisk.objects.filter(
            category="Covid-19", user=user,country=country)

        Covidnotcheck = AssuptionAndRisk.objects.filter(category="Covid-19",country=country).exclude(
            user=user
        )

        return render(
            request,
            "AssuptionAndRisk.html",
            {
                "showname": showname,
                "country": country,
                "industry": industry,
                "Generalcheck": Generalcheck,
                "Generalnotcheck": Generalnotcheck,
                "Resourcescheck": Resourcescheck,
                "Resourcesnotcheck": Resourcesnotcheck,
                "Workdaycheck": Workdaycheck,
                "Workdaynotcheck": Workdaynotcheck,
                "Softwarecheck": Softwarecheck,
                "Softwarenotcheck": Softwarenotcheck,
                "Integrationcheck": Integrationcheck,
                "Integrationnotcheck": Integrationnotcheck,
                "Datacheck": Datacheck,
                "Datanotcheck": Datanotcheck,
                "Testingcheck": Testingcheck,
                "Testingnotcheck": Testingnotcheck,
                "Changecheck": Changecheck,
                "Changenotcheck": Changenotcheck,
                "Deploymentcheck": Deploymentcheck,
                "Deploymentnotcheck": Deploymentnotcheck,
                "Covidcheck": Covidcheck,
                "Covidnotcheck": Covidnotcheck,
            },
        )

    if request.method == "GET":
        Acccount = "DEF"
        request.session["Acccount"] = Acccount
        user = Users.objects.get(user=client_name)
        Generalcheck = AssuptionAndRisk.objects.filter(
            category="General", user=user,country=country)

        Generalnotcheck = AssuptionAndRisk.objects.filter(category="General",country=country).exclude(
            user=user
        )

        Resourcescheck = AssuptionAndRisk.objects.filter(
            category="Resources", user=user,country=country
        )

        Resourcesnotcheck = AssuptionAndRisk.objects.filter(
            category="Resources",country=country
        ).exclude(user=user)

        Workdaycheck = AssuptionAndRisk.objects.filter(
            category="Workday", user=user,country=country)

        Workdaynotcheck = AssuptionAndRisk.objects.filter(category="Workday",country=country).exclude(
            user=user
        )

        Softwarecheck = AssuptionAndRisk.objects.filter(
            category="Software", user=user,country=country)

        Softwarenotcheck = AssuptionAndRisk.objects.filter(category="Software",country=country).exclude(
            user=user
        )

        Integrationcheck = AssuptionAndRisk.objects.filter(
            category="Integration", user=user,country=country
        )

        Integrationnotcheck = AssuptionAndRisk.objects.filter(
            category="Integration",country=country
        ).exclude(user=user)

        Datacheck = AssuptionAndRisk.objects.filter(
            category="Data Migration", user=user,country=country
        )

        Datanotcheck = AssuptionAndRisk.objects.filter(
            category="Data Migration",country=country
        ).exclude(user=user)

        Testingcheck = AssuptionAndRisk.objects.filter(
            category="Testing", user=user,country=country)

        Testingnotcheck = AssuptionAndRisk.objects.filter(category="Testing",country=country).exclude(
            user=user
        )

        Changecheck = AssuptionAndRisk.objects.filter(
            category="Change Management", user=user,country=country
        )

        Changenotcheck = AssuptionAndRisk.objects.filter(
            category="Change Management",country=country
        ).exclude(user=user)

        Deploymentcheck = AssuptionAndRisk.objects.filter(
            category="Deployment and Support", user=user,country=country
        )

        Deploymentnotcheck = AssuptionAndRisk.objects.filter(
            category="Deployment and Support",country=country
        ).exclude(user=user)

        Covidcheck = AssuptionAndRisk.objects.filter(
            category="Covid-19", user=user,country=country)

        Covidnotcheck = AssuptionAndRisk.objects.filter(category="Covid-19",country=country).exclude(
            user=user
        )
        messages.success(request, "SELECTED ASSUMPTION AND RISKS WILL BE ADDED TO THE MAIN RFP DOCUMENT")
        return render(
            request,
            "AssuptionAndRisk.html",
            {
                "showname": showname,
                "country": country,
                "industry": industry,
                "Generalcheck": Generalcheck,
                "Generalnotcheck": Generalnotcheck,
                "Resourcescheck": Resourcescheck,
                "Resourcesnotcheck": Resourcesnotcheck,
                "Workdaycheck": Workdaycheck,
                "Workdaynotcheck": Workdaynotcheck,
                "Softwarecheck": Softwarecheck,
                "Softwarenotcheck": Softwarenotcheck,
                "Integrationcheck": Integrationcheck,
                "Integrationnotcheck": Integrationnotcheck,
                "Datacheck": Datacheck,
                "Datanotcheck": Datanotcheck,
                "Testingcheck": Testingcheck,
                "Testingnotcheck": Testingnotcheck,
                "Changecheck": Changecheck,
                "Changenotcheck": Changenotcheck,
                "Deploymentcheck": Deploymentcheck,
                "Deploymentnotcheck": Deploymentnotcheck,
                "Covidcheck": Covidcheck,
                "Covidnotcheck": Covidnotcheck,
            },
        )


def notsatisfieddoc_view(request):
    client_name = request.session["client_name"]
    country = request.session["country"]
    queries = request.POST.get("Query")
    print("queries", queries)
    docfile = request.FILES.get("file")
    doc = notsatisfieddoc(
        user=client_name, docup=docfile, clientgeo=country, query=queries
    )
    doc.save()
    return render(request, "notsatisfieddoc.html")


def clientlogo_view(request):
    showname = request.session["showname"]
    industry = request.session["industry"]
    country = request.session["country"]
    client_name = request.session["client_name"]
    if request.method == "POST":
        Extraimgsearch = request.POST.get("Extraimgsearch")
        request.session["Extraimgsearch"] = Extraimgsearch
        print("Extraimgsearch", Extraimgsearch)
        extraimg = request.POST.getlist("extraimage")
        checkon = clientlogo.objects.filter(Industry=Extraimgsearch)
        checkonid = []
        for i in checkon:
            checkonid.append(i.id)
        print("checkonid", checkonid)
        print("extraimg", extraimg)
        extraimg = [int(x) for x in extraimg]
        print("extraimggggg", extraimg)
        on = list(set(checkonid).intersection(extraimg))
        print("on", on)
        if len(on):
            off = list(set(checkonid) - set(extraimg))
            print("off", off)
            extraselected = clientlogo.objects.filter(id__in=on)
            delextraselected = clientlogo.objects.filter(id__in=off)
            user = Users.objects.get(user=client_name)
            for e in extraselected:
                c = e.user.add(user)
            for d in delextraselected:
                c = d.user.remove(user)

        user = Users.objects.get(user=client_name)
        Extraimgsearch = request.session["Extraimgsearch"]
        print("KKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKK")
        print("Extraimgsearch", Extraimgsearch)
        print("KKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKKK")
        extra = clientlogo.objects.filter(Industry=Extraimgsearch, user=user)
        extrano = clientlogo.objects.filter(
            Industry=Extraimgsearch).exclude(user=user)
        print("extra", extra)
        print("extrano", extrano)
        return render(
            request,
            "clientlogo.html",
            {
                "showname": showname,
                "country": country,
                "industry": industry,
                "extra": extra,
                "extrano": extrano,
            },
        )
    if request.method == "GET":
        log = "DEF"
        request.session['log'] = log
        user = Users.objects.get(user=client_name)
        extra = clientlogo.objects.filter(user=user)
        extrano = clientlogo.objects.exclude(user=user)
        indus = clientlogo.objects.all()
        print("FFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFF")
        print(indus)
        for i in extrano:
            print(i.Industry)
        print("FFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFFF")
        messages.success(request, "SELECTED LOGO WILL BE ADDED TO THE MAIN RFP DOCUMENT")
        return render(request, 'clientlogo.html', {"showname": showname, "country": country, "industry": industry, "extra": extra, "extrano": extrano})


# upload Image files
def logo_upload_view(request):
    showname = request.session["showname"]
    industry = request.session["industry"]
    username = request.session["username"]
    client_name = request.session["client_name"]
    country = request.session["country"]
    pic = request.FILES.get("file")
    prod = logoUpload(user=username, picup=pic, clientgeo=country)
    prod.save()
    return render(
        request,
        "UploadClientlogo.html",
        {"showname": showname, "country": country, "industry": industry},
    )


def approveimage_view(request):
    if request.method == "GET":
        SP = userextraimage.objects.filter(approved="No")
        Yes = userextraimage.objects.filter(approved="Yes")
        return render(request, "approveimage.html", {"SP": SP, "Yes": Yes})
    if request.method == "POST":
        SP = userextraimage.objects.filter(approved="No")
        Yes = userextraimage.objects.filter(approved="Yes")
        return render(request, "approveimage.html", {"SP": SP, "Yes": Yes})


def approvedimage_view(request, id):
    
    if request.method == "GET":
        SP = userextraimage.objects.filter(id=id)
        SP.update(approved="Yes")
        SP = userextraimage.objects.filter(approved="No")
        Yes = userextraimage.objects.filter(approved="Yes")
        return render(
            request,
            "approveimage.html",
            {
                
                "SP": SP,
                "Yes": Yes,
            },
        )
    if request.method == "POST":
        SP = userextraimage.objects.filter(id=id)
        SP.update(approved="Yes")
        SP = userextraimage.objects.filter(approved="No")
        Yes = userextraimage.objects.filter(approved="Yes")
        return render(
            request,
            "approveimage.html",
            {
                
                "SP": SP,
                "Yes": Yes,
            },
        )


def disapprovedimage_view(request, id):
    
    if request.method == "GET":
        SP = userextraimage.objects.filter(id=id)
        SP.update(approved="No")
        SP = userextraimage.objects.filter(approved="No")
        Yes = userextraimage.objects.filter(approved="Yes")
        return render(
            request,
            "approveimage.html",
            {
                
                "SP": SP,
                "Yes": Yes,
            },
        )
    if request.method == "POST":
        SP = userextraimage.objects.filter(id=id)
        SP.update(approved="No")
        SP = userextraimage.objects.filter(approved="No")
        Yes = userextraimage.objects.filter(approved="Yes")
        return render(
            request,
            "approveimage.html",
            {
                
                "SP": SP,
                "Yes": Yes,
            },
        )


def approvelogo_view(request):
    if request.method == "GET":
        SP = logoUpload.objects.filter(approved="No")
        Yes = logoUpload.objects.filter(approved="Yes")
        return render(request, "approvelogo.html", {"SP": SP, "Yes": Yes})
    if request.method == "POST":
        SP = logoUpload.objects.filter(approved="No")
        Yes = logoUpload.objects.filter(approved="Yes")
        return render(request, "approvelogo.html", {"SP": SP, "Yes": Yes})


def approvedlogo_view(request, id):
    
    if request.method == "GET":
        SP = logoUpload.objects.filter(id=id)
        SP.update(approved="Yes")
        SP = logoUpload.objects.filter(approved="No")
        Yes = logoUpload.objects.filter(approved="Yes")
        return render(
            request,
            "approvelogo.html",
            {
                
                "SP": SP,
                "Yes": Yes,
            },
        )
    if request.method == "POST":
        SP = logoUpload.objects.filter(id=id)
        SP.update(approved="Yes")
        SP = logoUpload.objects.filter(approved="No")
        Yes = logoUpload.objects.filter(approved="Yes")
        return render(
            request,
            "approvelogo.html",
            {
                
                "SP": SP,
                "Yes": Yes,
            },
        )


def disapprovedlogo_view(request, id):
    
    if request.method == "GET":
        SP = logoUpload.objects.filter(id=id)
        SP.update(approved="No")
        SP = logoUpload.objects.filter(approved="No")
        Yes = logoUpload.objects.filter(approved="Yes")
        return render(
            request,
            "approvelogo.html",
            {
                
                "SP": SP,
                "Yes": Yes,
            },
        )
    if request.method == "POST":
        SP = logoUpload.objects.filter(id=id)
        SP.update(approved="No")
        SP = logoUpload.objects.filter(approved="No")
        Yes = logoUpload.objects.filter(approved="Yes")
        return render(
            request,
            "approvelogo.html",
            {
                
                "SP": SP,
                "Yes": Yes,
            },
        )


def approvequestionans_view(request):
    if request.method == "GET":
        SP = userquestionans.objects.filter(approved="No")
        Yes = userquestionans.objects.filter(approved="Yes")
        return render(request, "approvequestionans.html", {"SP": SP, "Yes": Yes})
    if request.method == "POST":
        SP = userquestionans.objects.filter(approved="No")
        Yes = userquestionans.objects.filter(approved="Yes")
        return render(request, "approvequestionans.html", {"SP": SP, "Yes": Yes})


def approvedquestionans_view(request, id):
    
    if request.method == "GET":
        SP = userquestionans.objects.filter(id=id)
        SP.update(approved="Yes")
        SP = userquestionans.objects.filter(approved="No")
        Yes = userquestionans.objects.filter(approved="Yes")
        return render(
            request,
            "approvequestionans.html",
            {
                
                "SP": SP,
                "Yes": Yes,
            },
        )
    if request.method == "POST":
        SP = userquestionans.objects.filter(id=id)
        SP.update(approved="Yes")
        SP = userquestionans.objects.filter(approved="No")
        Yes = userquestionans.objects.filter(approved="Yes")
        return render(
            request,
            "approvequestionans.html",
            {
                
                "SP": SP,
                "Yes": Yes,
            },
        )


def disapprovequestionans_view(request, id):
    
    if request.method == "GET":
        SP = userquestionans.objects.filter(id=id)
        SP.update(approved="No")
        SP = userquestionans.objects.filter(approved="No")
        Yes = userquestionans.objects.filter(approved="Yes")
        return render(
            request,
            "approvequestionans.html",
            {
                
                "SP": SP,
                "Yes": Yes,
            },
        )
    if request.method == "POST":
        SP = userquestionans.objects.filter(id=id)
        SP.update(approved="No")
        SP = userquestionans.objects.filter(approved="No")
        Yes = userquestionans.objects.filter(approved="Yes")
        return render(
            request,
            "approvequestionans.html",
            {
                
                "SP": SP,
                "Yes": Yes,
            },
        )


def approveassumptionrisk_view(request):
    if request.method == "GET":
        SP = userriskandassumption.objects.filter(approved="No")
        Yes = userriskandassumption.objects.filter(approved="Yes")
        return render(request, "userassumptionrisk.html", {"SP": SP, "Yes": Yes})
    if request.method == "POST":
        SP = userriskandassumption.objects.filter(approved="No")
        Yes = userriskandassumption.objects.filter(approved="Yes")
        return render(request, "userassumptionrisk.html", {"SP": SP, "Yes": Yes})


def approvedassumptionrisk_view(request, id):
    if request.method == "GET":
        SP = userriskandassumption.objects.filter(id=id)
        SP.update(approved="Yes")
        SP = userriskandassumption.objects.filter(approved="No")
        Yes = userriskandassumption.objects.filter(approved="Yes")
        return render(
            request,
            "userassumptionrisk.html",
            {
                
                "SP": SP,
                "Yes": Yes,
            },
        )
    if request.method == "POST":
        SP = userriskandassumption.objects.filter(id=id)
        SP.update(approved="Yes")
        SP = userriskandassumption.objects.filter(approved="No")
        Yes = userriskandassumption.objects.filter(approved="Yes")
        return render(
            request,
            "userassumptionrisk.html",
            {
                
                "SP": SP,
                "Yes": Yes,
            },
        )


def disapprovedassumptionrisk_view(request, id):
    
    if request.method == "GET":
        SP = userriskandassumption.objects.filter(id=id)
        SP.update(approved="No")
        SP = userriskandassumption.objects.filter(approved="No")
        Yes = userriskandassumption.objects.filter(approved="Yes")
        return render(
            request,
            "userassumptionrisk.html",
            {
                
                "SP": SP,
                "Yes": Yes,
            },
        )
    if request.method == "POST":
        SP = userriskandassumption.objects.filter(id=id)
        SP.update(approved="No")
        SP = userriskandassumption.objects.filter(approved="No")
        Yes = userriskandassumption.objects.filter(approved="Yes")
        return render(
            request,
            "userassumptionrisk.html",
            {
                
                "SP": SP,
                "Yes": Yes,
            },
        )


def userquestionans_view(request):
    loginusername = request.session["loginusername"]
    if request.method == "GET":
        return render(request, "userquestionans.html")

    if request.method == "POST":
        user = request.POST.get("client_name")
        rfpid = request.POST.get("Rfp_Id")
        country = request.POST.get("countries")
        industry = request.POST.get("industry")
        section = request.POST.get("Section")
        subsection = request.POST.get("sub_section")
        question = request.POST.get("question")
        document = request.FILES["docfile"]
        image = request.FILES.get("imgfile")
        docu = userquestionans.objects.create(
            user=loginusername,
            rfpid=rfpid,
            country=country,
            industry=industry,
            section=section,
            subsection=subsection,
            question=question,
            document=document,
            image=image,
        )
        docu.save()
        return render(request, "userquestionans.html")


def userstandardsection_view(request):
    if request.method == "GET":
        return render(request, "userstandardsection.html")

    if request.method == "POST":
        loginusername = request.session["loginusername"]
        user = loginusername
        rfpid = request.POST.get("Rfp_Id")
        country = request.POST.get("countries")
        industry = request.POST.get("industry")
        section = request.POST.get("Section")
        subsection = request.POST.get("sub_section")
        question = request.POST.get("question")
        document = request.FILES["docfile"]
        image = request.FILES["Imgfile"]
        docu = userstandardsection.objects.create(
            user=user,
            rfpid=rfpid,
            country=country,
            industry=industry,
            section=section,
            subsection=subsection,
            question=question,
            document=document,
            image=image,
        )
        docu.save()
        return render(request, "userstandardsection.html")


def useraddextraimages_view(request):
    if request.method == "GET":
        return render(request, "useraddextraimages.html")
    if request.method == "POST":
        country = request.POST["countries"]
        industry = request.POST["industry"]
        image = request.FILES.get("Imgfile")
        loginusername = request.session["loginusername"]
        img = userextraimage.objects.create(
            user=loginusername, country=country, industry=industry, image=image
        )
        img.save()
        return render(request, "useraddextraimages.html")


def userriskandassumption_view(request):
    if request.method == "GET":
        return render(request, "userriskandassumption.html")
    if request.method == "POST":
        country = request.POST["countries"]
        industry = request.POST["industry"]
        section = request.POST["Section"]
        loginusername = request.session["loginusername"]
        description = request.POST["description"]
        assump = userriskandassumption.objects.create(
            user=loginusername, country=country, industry=industry, section=section, description=description)
        assump.save()
        return render(request, "userriskandassumption.html")


def useraddlogo_view(request):
    if request.method == "GET":
        return render(request, "useraddlogo.html")
    if request.method == "POST":
        country = request.POST["countries"]
        industry = request.POST["industry"]
        image = request.FILES.get("Imgfile")
        loginusername = request.session["loginusername"]
        img = logoUpload.objects.create(
            user=loginusername, clientgeo=country, industry=industry, picup=image
        )
        img.save()
        return render(request, "useraddlogo.html")


def rejectfeedbackform_view(request, id):
    if request.method == "GET":
        return render(request, "rejectfeedback.html", {"id": id})
    if request.method == "POST":
        feedback = request.POST.get("feedback")
        try:
            assump = userriskandassumption.objects.filter(
                id=id).update(feedback=feedback, approved="Rejected")
            assump.save()
            SP = userriskandassumption.objects.filter(approved="No")
            Yes = userriskandassumption.objects.filter(approved="Yes")
            return render(request, "userassumptionrisk.html", {"SP": SP, "Yes": Yes})
        except:
            pass

        try:
            assump = logoUpload.objects.filter(
                id=id).update(feedback=feedback, approved="Rejected")
            assump.save()
        except:
            pass

        try:
            assump = userquestionans.objects.filter(
                id=id).update(feedback=feedback, approved="Rejected")
            assump.save()
        except:
            pass

        try:
            assump = userstandardsection.objects.filter(
                id=id).update(feedback=feedback, approved="Rejected")
            assump.save()
        except:
            pass

        try:
            assump = userextraimage.objects.filter(
                id=id).update(feedback=feedback, approved="Rejected")
            assump.save()
        except:
            pass

        return render(request, "rejectfeedback.html", {"id": id})


def usersummerytable_view(request):
    loginusername = request.session["loginusername"]
    if request.method == "GET":
        SP = userstandardsection.objects.filter(
            user=loginusername).order_by('id')
        assumption = userriskandassumption.objects.filter(
            user=loginusername).order_by('id')
        logo = logoUpload.objects.filter(user=loginusername).order_by('id')
        extraimage = userextraimage.objects.filter(
            user=loginusername).order_by('id')
        questionans = userquestionans.objects.filter(
            user=loginusername).order_by('id')

        return render(
            request,
            "usersummarytable.html",
            {
                "SP": SP,
                "assumption": assumption,
                "logo": logo,
                "extraimage": extraimage,
                "questionans": questionans,
            },
        )
    if request.method == "POST":
        SP = userstandardsection.objects.all().order_by('id')
        assumption = userriskandassumption.objects.all().order_by('id')
        logo = logoUpload.objects.all().order_by('id')
        extraimage = userextraimage.objects.all().order_by('id')
        questionans = userquestionans.objects.all().order_by('id')
        return render(
            request,
            "usersummarytable.html",
            {
                "SP": SP,
                "assumption": assumption,
                "logo": logo,
                "extraimage": extraimage,
                "questionans": questionans,
            },
        )

def user_dashboard(request):
    print('inside dashoboard')
    return render(request, 'user_dashboard.html')