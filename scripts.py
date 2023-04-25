# from docx import Document

# file_data = ["C:/Users/narayanac/Documents/RFP-Builder-Latest/Title_template2.docx","C:/Users/narayanac/Documents/RFP-Builder-Latest/Title_template3.docx"]

# file_data = ['Title_template2.docx', 'Title_template3.docx']
# def combine_word_documents(files=file_data):
#     print('entered')
#     merged_document = Document()

#     for index, file in enumerate(files):
#         print(index, file)
#         sub_doc = Document(file)
#         print(sub_doc, 'su')

#         # Don't add a page break if you've reached the last file.
#         if index < len(files)-1:
#            sub_doc.add_page_break()

#         for element in sub_doc.element.body:
#             merged_document.element.body.append(element)

#     merged_document.save('merged-test.docx')

# combine_word_documents(file_data)

import wget
import os
import aspose.words as aw
from docxcompose.composer import Composer
from docx import Document as Document_compose


# path = 'static/media/'+client_name+'/'+country+'/'+industry

files_data = ["Agnostic_USA_Executive Summary.docx",
              "Agnostic_USA_Implementation timeline.docx"]


def combine_all_docx(filename_master, files_list):
    number_of_sections = len(files_list)
    master = Document_compose(filename_master)
    composer = Composer(master)
    for i in range(0, number_of_sections):
        doc_temp = Document_compose(files_list[i])
        composer.append(doc_temp)
    composer.save("test_merge.docx")
    return 'test_merge.docx'

# path = 'C:/Users/narayanac/Documents/RFP-Builder-Latest/'
# filename_master = "C:/Users/narayanac/Documents/RFP-project/RFP-Builder-Latest/C.docx"
# combine_all_docx(filename_master, files_data)


def replace_word_document(client_name, replace_name, doc_path, doc_name):
    print(client_name, replace_name, 'arguments')
    doc = Document_compose(doc_path)
    Dictionary = {replace_name: client_name}
    for i in Dictionary:
        for p in doc.paragraphs:
            if p.text.find(i) >= 0:
                p.text = p.text.replace(i, Dictionary[i])

    # delete the existing updated document
    # os.remove('updated.docx')

    # save changed document
    doc.save('temporary/updated.docx')
    os.remove(doc_path)
    media_path = f'media/{client_name}'
    isExist = os.path.exists(media_path)

    if not isExist:

        # Create a new directory because it does not exist
        os.makedirs(media_path)
        print("The new directory is created!")

    os.rename('temporary/updated.docx', f'{media_path}/{doc_name}')
    return f'{media_path}/{doc_name}'


def replace_word_document_aspose(client_name, replace_name, doc_path):

    # load Word document
    # doc = aw.Document("document.docx")
    doc = aw.Document(doc_path)

    # replace text
    # doc.range.replace(cliend_name, "[CLIENT]", aw.replacing.FindReplaceOptions(aw.replacing.FindReplaceDirection.FORWARD))
    doc.range.replace(client_name, replace_name, aw.replacing.FindReplaceOptions(
        aw.replacing.FindReplaceDirection.FORWARD))

    # save the modified document
    # doc.save("updated.docx")
    doc.save('update.docx')
    print(doc, 'doc inside update')
    return doc

# import urllib

# def get_document(file_path, file_name):
#     print(file_path, file_name, 'file path')
#     document = urllib.request.urlretrieve (file_path, 'test_create_001.docx')
#     print(document, 'document output')
#     return document


def get_document(file_path):
    print(file_path, 'file path')

    file_name = wget.download(file_path)

    return file_name


def get_document_update(file_path, outpath):
    print(file_path, 'file path')

    file_name = wget.download(file_path, out=outpath)

    return file_name
